#!/usr/bin/env python3
"""
Gera tabela Word com estrutura completa do Regulamento 2023/0447
Coluna 1: Referência (Art. X.º, n.º Y, al. Z) | Anexo X, n.º Y, al. Z)
Coluna 2: Texto verbatim (Calibri Light 8)
Coluna 3: Em branco

V.2: Inclui estrutura completa dos Anexos (números, sub-números, alíneas)
"""

import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONTE = 'Calibri Light'
TAMANHO_TEXTO = 8
TAMANHO_REF = 8

def set_cell_font(cell, text, bold=False, size=TAMANHO_TEXTO):
    for para in cell.paragraphs:
        para.clear()
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.name = FONTE
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONTE)
    rFonts.set(qn('w:hAnsi'), FONTE)
    rPr.insert(0, rFonts)

def limpar(t):
    return t.replace('▌', '').replace('\u200b', '').strip()

# ── Extrair estrutura ──────────────────────────────────────────────────────
doc_src = Document('pe00002.pt26.PB.aftermeeting 2.docx')

rows = []
in_scope = False
current_art_num = None
current_art_titulo = None
current_n = None
current_al = None
in_annex = None
current_annex_n = None    # número principal do Anexo (ex: "1", "2")
current_annex_subn = None # sub-número do Anexo (ex: "1.1", "2.3")

RE_CAP    = re.compile(r'^CAPÍTULO\s+(II|III|IV|V|VI|VII|VIII|IX|X)\b', re.IGNORECASE)
RE_ART    = re.compile(r'^Artigo\s+(\d+)\.º\s*\n?(.+)?', re.IGNORECASE | re.DOTALL)
RE_NUM    = re.compile(r'^(\d+)\.\t(.+)', re.DOTALL)
RE_NUM2   = re.compile(r'^(\d+)\.\s{2,}(.+)', re.DOTALL)
RE_ALIN   = re.compile(r'^([a-z])\)\t(.+)', re.DOTALL)
RE_ALIN2  = re.compile(r'^([a-z])\)\s{2,}(.+)', re.DOTALL)
RE_SUB    = re.compile(r'^(i{1,3}|iv|vi{0,3}|ix|x{0,3}i{0,3})\)\t(.+)', re.DOTALL | re.IGNORECASE)
RE_ANNEX  = re.compile(r'^ANEXO\s+(I{1,3}|IV|V)\b', re.IGNORECASE)
RE_ANNEX_NUM    = re.compile(r'^(\d+)\.\t(.+)', re.DOTALL)           # Ex: "1. Alimentação..."
RE_ANNEX_SUBN   = re.compile(r'^(\d+\.\d+)\.\t(.+)', re.DOTALL)      # Ex: "1.1. Os cães..."
RE_ANNEX_ALIN   = re.compile(r'^([a-z])\)\t(.+)', re.DOTALL)         # Ex: "a) 22°C..."
RE_ANNEX_ALIN2  = re.compile(r'^([a-z])\)\s{2,}(.+)', re.DOTALL)

paragraphs = list(doc_src.paragraphs)
i = 0

while i < len(paragraphs):
    p = paragraphs[i]
    t = p.text.strip()
    i += 1

    if not t or t == '▌' or re.match(r'^[▌\s]+$', t):
        continue

    t_clean = limpar(t)

    # Detectar início (Cap II)
    if RE_CAP.match(t) and 'II' in t:
        in_scope = True
        continue

    if not in_scope:
        if RE_ANNEX.match(t):
            in_scope = True
        else:
            continue

    # Detectar Capítulos/Secções
    if RE_CAP.match(t):
        current_art_num = None
        current_n = None
        current_al = None
        continue

    if re.match(r'^SECÇÃO', t, re.IGNORECASE):
        continue

    # Detectar Anexo
    m_annex = RE_ANNEX.match(t)
    if m_annex:
        in_annex = m_annex.group(1).upper()
        current_art_num = None
        current_n = None
        current_al = None
        current_annex_n = None
        current_annex_subn = None
        rows.append((f'Anexo {in_annex}', limpar(t)))
        continue

    # Se estamos num Anexo, processar estrutura de Anexo
    if in_annex:
        # Detectar sub-número do Anexo (ex: 1.1, 2.3)
        m_subn = RE_ANNEX_SUBN.match(t)
        if m_subn:
            current_annex_subn = m_subn.group(1)
            texto = limpar(m_subn.group(2))
            ref = f'Anexo {in_annex}, {current_annex_subn}'
            rows.append((ref, texto))
            continue

        # Detectar número do Anexo (ex: 1, 2)
        m_num = RE_ANNEX_NUM.match(t)
        if m_num and not re.match(r'^\d+\.\d+', m_num.group(1)):  # Não é sub-número
            current_annex_n = m_num.group(1)
            current_annex_subn = None
            texto = limpar(m_num.group(2))
            ref = f'Anexo {in_annex}, {current_annex_n}'
            rows.append((ref, texto))
            continue

        # Detectar alínea do Anexo
        m_alin = RE_ANNEX_ALIN.match(t) or RE_ANNEX_ALIN2.match(t)
        if m_alin:
            al = m_alin.group(1)
            texto = limpar(m_alin.group(2))
            if current_annex_subn:
                ref = f'Anexo {in_annex}, {current_annex_subn}, al. {al})'
            elif current_annex_n:
                ref = f'Anexo {in_annex}, {current_annex_n}, al. {al})'
            else:
                ref = f'Anexo {in_annex}, al. {al})'
            rows.append((ref, texto))
            continue

    # Detectar Artigo (fora de Anexo)
    if not in_annex:
        m_art = re.match(r'^(Artigo\s+(\d+)\.º)\s*\n?(.*)', t, re.IGNORECASE | re.DOTALL)
        if m_art:
            current_art_num = m_art.group(2)
            titulo_art = m_art.group(3).strip()
            titulo_art = limpar(titulo_art)
            current_art_titulo = titulo_art
            current_n = None
            current_al = None
            in_annex = None
            ref = f'Art. {current_art_num}.º'
            rows.append((ref, titulo_art if titulo_art else ''))
            continue

        # Detectar número (Art)
        m_num = RE_NUM.match(t) or RE_NUM2.match(t)
        if m_num and current_art_num:
            current_n = m_num.group(1)
            current_al = None
            texto = limpar(m_num.group(2))
            ref = f'Art. {current_art_num}.º, n.º {current_n}'
            rows.append((ref, texto))
            continue

        # Detectar alínea (Art)
        m_al = RE_ALIN.match(t) or RE_ALIN2.match(t)
        if m_al and current_art_num:
            current_al = m_al.group(1)
            texto = limpar(m_al.group(2))
            if current_n:
                ref = f'Art. {current_art_num}.º, n.º {current_n}, al. {current_al})'
            else:
                ref = f'Art. {current_art_num}.º, al. {current_al})'
            rows.append((ref, texto))
            continue

        # Detectar sub-alínea (Art)
        m_sub = RE_SUB.match(t)
        if m_sub and current_art_num and current_al:
            sub = m_sub.group(1).lower()
            texto = limpar(m_sub.group(2))
            ref = f'Art. {current_art_num}.º, n.º {current_n}, al. {current_al}), {sub})'
            rows.append((ref, texto))
            continue

        # Parágrafo solto em artigo
        if current_art_num and not re.match(r'^(CAPÍTULO|SECÇÃO|Feito|Pelo Parlamento|Anexo)', t, re.IGNORECASE):
            texto = limpar(t)
            if texto and len(texto) > 3:
                if current_n:
                    ref = f'Art. {current_art_num}.º, n.º {current_n}'
                else:
                    ref = f'Art. {current_art_num}.º'
                if not rows or rows[-1][1] != texto:
                    rows.append((ref, texto))

print(f"Total de linhas extraídas: {len(rows)}")
for ref, txt in rows[-20:]:
    print(f"  [{ref}] {txt[:60]}")

# ── Criar documento Word ─────────────────────────────────────────────────────
out = Document()

section = out.sections[0]
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(1.5)

# Título
titulo_p = out.add_paragraph()
titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = titulo_p.add_run('Regulamento (UE) 2023/0447 — Estrutura Completa v.2 (Cap. II ao Anexo III)')
run_t.font.name = FONTE
run_t.font.size = Pt(11)
run_t.font.bold = True

out.add_paragraph()

# Tabela
table = out.add_table(rows=1, cols=3)
table.style = 'Table Grid'

col_widths = [Cm(4.0), Cm(11.0), Cm(3.0)]
for j, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[j].width = width

# Cabeçalho
hdr = table.rows[0].cells
set_cell_font(hdr[0], 'Referência', bold=True, size=8)
set_cell_font(hdr[1], 'Texto (verbatim)', bold=True, size=8)
set_cell_font(hdr[2], '', bold=False, size=8)

# Dados
for ref, texto in rows:
    row = table.add_row()
    row.cells[0].width = col_widths[0]
    row.cells[1].width = col_widths[1]
    row.cells[2].width = col_widths[2]
    set_cell_font(row.cells[0], ref, bold=False, size=TAMANHO_REF)
    set_cell_font(row.cells[1], texto, bold=False, size=TAMANHO_TEXTO)
    set_cell_font(row.cells[2], '', bold=False, size=TAMANHO_TEXTO)

filename = 'Reg2023_0447_Tabela_Estrutura_v2.docx'
out.save(filename)
print(f"\n✅ Documento criado: {filename}")
print(f"   {len(rows)} linhas (artigos + n.ºs + alíneas + anexos + sub-números)")
