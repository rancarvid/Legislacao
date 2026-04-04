"""
Gera tabela Word de prazos de implementação do Regulamento 2023/0447 (cães e gatos).
Organização temática nas linhas; colunas cronológicas de prazo.
Output: tabela_prazos_regulamento.docx
"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Cores
# ---------------------------------------------------------------------------
C_HEADER_BG    = "1F4E79"   # azul petróleo escuro — cabeçalho
C_HEADER_TEXT  = "FFFFFF"   # branco
C_GROUP_BG     = "D9E8F5"   # azul claro — separador de grupo
C_GROUP_TEXT   = "1F4E79"   # azul petróleo escuro
C_ROW_EVEN     = "EBF3FB"   # azul muito claro
C_ROW_ODD      = "FFFFFF"   # branco
C_CHECK        = "1F4E79"   # azul para ✓
C_REF          = "1F4E79"   # azul para referência do artigo
C_TEXT         = "404040"   # cinzento escuro para verbatim
C_BORDER       = "B8CCE4"   # azul claro para bordas

# ---------------------------------------------------------------------------
# Dimensões
# ---------------------------------------------------------------------------
PAGE_W_CM    = 29.7
MARGIN_CM    = 1.5
USABLE_CM    = PAGE_W_CM - 2 * MARGIN_CM   # 26.7 cm
COL1_CM      = 10.8
N_PRAZO_COLS = 12
PRAZO_CM     = (USABLE_CM - COL1_CM) / N_PRAZO_COLS   # ~1.325 cm

# ---------------------------------------------------------------------------
# Colunas de prazo (cronológicas)
# ---------------------------------------------------------------------------
COLUNAS = [
    {"key": "2a",       "label": "2\nanos"},
    {"key": "3a",       "label": "3\nanos"},
    {"key": "4a",       "label": "4\nanos"},
    {"key": "5a",       "label": "5\nanos"},
    {"key": "6a",       "label": "6\nanos"},
    {"key": "7a",       "label": "7\nanos"},
    {"key": "8a",       "label": "8\nanos"},
    {"key": "10a",      "label": "10\nanos"},
    {"key": "14a",      "label": "14\nanos"},
    {"key": "15a",      "label": "15\nanos"},
    {"key": "1Jul2030", "label": "1 Jul.\n2030"},
    {"key": "1Jul2036", "label": "1 Jul.\n2036"},
]

# ---------------------------------------------------------------------------
# Dados — grupos temáticos
# ---------------------------------------------------------------------------
GRUPOS = [
    {
        "titulo": "Grupo 1 — Bem-estar e Reprodução",
        "normas": [
            {
                "ref": "Art. 7(3)",
                "texto": '"The delegated acts concerning the excessive conformational traits shall be adopted by 1 July 2030."',
                "prazos": ["1Jul2030"],
            },
            {
                "ref": "Art. 7(3)",
                "texto": '"The delegated acts concerning the genotypes shall be adopted by 1 July 2036."',
                "prazos": ["1Jul2036"],
            },
            {
                "ref": "Art. 8(2)\n[via Art. 33(2)(iii)]",
                "texto": '"Article 8(2) shall apply from 1 July 2030."',
                "prazos": ["1Jul2030"],
            },
            {
                "ref": "Art. 8(1)\n[via Art. 33(2)(iii)]",
                "texto": '"Article 8(1) shall apply from 1 July 2036."',
                "prazos": ["1Jul2036"],
            },
            {
                "ref": "Art. 12(2)(3)\n[via Art. 33(2)(v)]",
                "texto": '"Article 12(2) and (3) shall apply from … seven years from the date of entry into force of this Regulation."',
                "prazos": ["7a"],
            },
            {
                "ref": "Art. 12(4)",
                "texto": '"The implementing act concerning the training courses referred to in paragraph 3 shall be adopted by … [3 years from entry into force of this Regulation]."',
                "prazos": ["3a"],
            },
            {
                "ref": "Art. 13(1)(a)",
                "texto": '"… establishments … receive a visit by a veterinarian … initially by … [three years after the date of entry into force of this Regulation] or one year following the notification of the new establishment."',
                "prazos": ["3a"],
            },
            {
                "ref": "Art. 13(2)",
                "texto": '"By … [24 months from the date of entry into force of this Regulation], the Commission shall adopt delegated acts … laying down the minimum criteria to be assessed by the veterinarian during the advisory welfare visit."',
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 16\n[via Art. 33(2)(i)]",
                "texto": '"Article 16 shall apply from … [three years from the date of entry into force of this Regulation]."',
                "prazos": ["3a"],
            },
        ],
    },
    {
        "titulo": "Grupo 2 — Identificação e Rastreabilidade",
        "normas": [
            {
                "ref": "Art. 20(2)",
                "texto": '"Where dogs and cats have been individually identified … before … [two years after the date of entry into force of this Regulation] they shall be considered to be compliant with the requirements in paragraph 1…"',
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 20(7)(a)",
                "texto": '"… for operators and natural or legal persons placing dogs and cats on the market from … [4 years from the entry into force of this Regulation]."',
                "prazos": ["4a"],
            },
            {
                "ref": "Art. 20(7)(b)",
                "texto": '"… for pet owners … who do not place dogs on the market: from … [10 years from entry into force of this Regulation]."',
                "prazos": ["10a"],
            },
            {
                "ref": "Art. 20(7)(c)",
                "texto": '"… for pet owners … who do not place cats on the market: from … [15 years from entry into force of this Regulation]."',
                "prazos": ["15a"],
            },
        ],
    },
    {
        "titulo": "Grupo 3 — Bases de Dados e Publicidade Online",
        "normas": [
            {
                "ref": "Art. 21(3)\n[via Art. 33(2)(ii)]",
                "texto": '"Article 21(3) … shall apply from … [four years from entry into force of this Regulation]."',
                "prazos": ["4a"],
            },
            {
                "ref": "Art. 21(6)(a)",
                "texto": '"The implementing acts referred to in point (a) shall be adopted by … [two years after the date of entry into force of this Regulation]…"',
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 21(6)(b)(c)",
                "texto": '"… the implementing act referred to in points (b) and (c) shall be adopted by … [three years from date of entry into force of this Regulation]."',
                "prazos": ["3a"],
            },
            {
                "ref": "Art. 23(1)\n[via Art. 33(2)(ii)]",
                "texto": '"Article 23(1) shall apply from … [four years from entry into force of this Regulation]."',
                "prazos": ["4a"],
            },
            {
                "ref": "Art. 23(7)(a)(c)",
                "texto": '"The Commission shall adopt the implementing acts referred to … points (a) and (c) by … [two years after the date of entry into force of this Regulation]."',
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 23(7)(b)(d)(e)",
                "texto": '"It shall adopt the implementing acts referred to … points (b), (d) and (e) by … [three years from the date of entry into force of this Regulation]."',
                "prazos": ["3a"],
            },
        ],
    },
    {
        "titulo": "Grupo 4 — Detenção Responsável e Formação",
        "normas": [
            {
                "ref": "Art. 15\n[via Art. 33(2)(iv)]",
                "texto": '"Article 15 … shall apply from … [five years from the date of entry into force of this Regulation]."',
                "prazos": ["5a"],
            },
            {
                "ref": "Art. 21(3) 2.ª parte\n+ 21(4)(5)\n[via Art. 33(2)(iv)]",
                "texto": '"Article 21(3), second subparagraph, (4) and (5) … shall apply from … [five years from the date of entry into force of this Regulation]."',
                "prazos": ["5a"],
            },
            {
                "ref": "Art. 22(1)(a)(b)(c)\n[via Art. 33(2)(iv)]",
                "texto": '"Article 22(1), points (a), (b) and (c) … shall apply from … [five years from the date of entry into force of this Regulation]."',
                "prazos": ["5a"],
            },
            {
                "ref": "Art. 23(3)(4)\n[via Art. 33(2)(iv)]",
                "texto": '"Article 23(3) and (4) … shall apply from … [five years from the date of entry into force of this Regulation]."',
                "prazos": ["5a"],
            },
        ],
    },
    {
        "titulo": "Grupo 5 — Importação e Movimento",
        "normas": [
            {
                "ref": "Art. 10\n[via Art. 33(2)(vi)]",
                "texto": '"Article 10 shall apply from … [eight years from the date of entry into force of this Regulation]."',
                "prazos": ["8a"],
            },
            {
                "ref": "Art. 26(1)(2)(3)\n[via Art. 33(2)(iv)]",
                "texto": '"Article 26(1), (2) and (3) shall apply from … [five years from the date of entry into force of this Regulation]."',
                "prazos": ["5a"],
            },
            {
                "ref": "Art. 26(4)\n[via Art. 33(2)(vii)]",
                "texto": '"Article 26(4) shall apply from … [10 years from the entry into force of the Regulation]."',
                "prazos": ["10a"],
            },
            {
                "ref": "Art. 26(5)",
                "texto": '"The Commission shall by … [8 years after the date of entry into force of this Regulation] adopt implementing acts laying down detailed arrangements for the Union pet travellers\' database."',
                "prazos": ["8a"],
            },
        ],
    },
    {
        "titulo": "Grupo 6 — Relatórios e Avaliação",
        "normas": [
            {
                "ref": "Art. 24(2)",
                "texto": '"The first such report shall be drawn up and transmitted to the Commission by … [6 years from the date of entry into force of this Regulation]."',
                "prazos": ["6a"],
            },
            {
                "ref": "Art. 31(1)",
                "texto": '"… the Commission shall publish, by … [7 years from the date of entry into force of this Regulation] … a monitoring report on the welfare of dogs and cats placed on the market in the Union."',
                "prazos": ["7a"],
            },
            {
                "ref": "Art. 31(2)",
                "texto": '"By … [14 years from the date of entry into force of this Regulation], the Commission shall carry out an evaluation of this Regulation…"',
                "prazos": ["14a"],
            },
        ],
    },
    {
        "titulo": "Grupo 7 — Disposições Institucionais",
        "normas": [
            {
                "ref": "Art. 30(2)",
                "texto": '"Member States shall, by … [two years after the date of entry into force of this Regulation], inform the Commission about any existing stricter national rules…"',
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 33(2)",
                "texto": '"It shall apply from … [two years from the date of entry into force of this Regulation]. However, Articles 8, 10, 12(2)(3), 15, 16, 21, 22, 23, 26 apply from later dates as specified herein."',
                "prazos": ["2a"],
            },
        ],
    },
]


# ---------------------------------------------------------------------------
# Auxiliares XML / python-docx
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn("w:shd")):
        tcPr.remove(s)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_width(cell, cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:tcW")):
        tcPr.remove(el)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm / 2.54 * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_cell_valign(cell, val="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(el)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)


def set_cell_margins(cell, top=40, bottom=40, left=72, right=72):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_paragraph_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn("w:spacing")):
        pPr.remove(s)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    pPr.append(sp)


def add_run_break(run):
    br = OxmlElement("w:br")
    run._r.append(br)


def disable_autofit(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)


def set_table_width(table, cm):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(cm / 2.54 * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), C_BORDER)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def clear_cell(cell):
    """Remove todos os parágrafos de uma célula e devolve célula limpa."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    p.clear()
    return p


# ---------------------------------------------------------------------------
# Construtores de linhas
# ---------------------------------------------------------------------------

def build_header_row(row):
    """Linha de cabeçalho: 'Norma / Disposição' + nomes das colunas de prazo."""
    cells = row.cells

    # Coluna 1
    c = cells[0]
    set_cell_shading(c, C_HEADER_BG)
    set_cell_width(c, COL1_CM)
    set_cell_valign(c, "center")
    set_cell_margins(c)
    p = clear_cell(c)
    set_paragraph_spacing(p, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Norma / Disposição")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Colunas de prazo
    for i, col in enumerate(COLUNAS):
        c = cells[i + 1]
        set_cell_shading(c, C_HEADER_BG)
        set_cell_width(c, PRAZO_CM)
        set_cell_valign(c, "center")
        set_cell_margins(c, top=40, bottom=40, left=20, right=20)
        p = clear_cell(c)
        set_paragraph_spacing(p, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = col["label"].split("\n")
        run = p.add_run(parts[0])
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if len(parts) > 1:
            add_run_break(run)
            run2 = p.add_run(parts[1])
            run2.bold = True
            run2.font.size = Pt(8)
            run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def build_group_row(row, titulo):
    """Linha separadora de grupo temático (fundo azul claro, texto itálico)."""
    cells = row.cells
    # Merge all cells
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)
    set_cell_shading(merged, C_GROUP_BG)
    set_cell_valign(merged, "center")
    set_cell_margins(merged, top=50, bottom=50, left=100, right=80)
    p = clear_cell(merged)
    set_paragraph_spacing(p, 0, 0)
    run = p.add_run(titulo)
    run.bold = True
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def build_norma_row(row, norma, even):
    """Linha de norma: referência + verbatim + ✓ nas colunas certas."""
    bg = C_ROW_EVEN if even else C_ROW_ODD
    cells = row.cells

    # -- Coluna 1: referência + verbatim --
    c = cells[0]
    set_cell_shading(c, bg)
    set_cell_width(c, COL1_CM)
    set_cell_valign(c, "center")
    set_cell_margins(c, top=40, bottom=40, left=80, right=60)

    # Parágrafo 1: referência (negrito, azul)
    p1 = clear_cell(c)
    set_paragraph_spacing(p1, 0, 2)
    ref_parts = norma["ref"].split("\n")
    run_ref = p1.add_run(ref_parts[0])
    run_ref.bold = True
    run_ref.font.size = Pt(8)
    run_ref.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    for part in ref_parts[1:]:
        add_run_break(run_ref)
        run_via = p1.add_run(part)
        run_via.bold = False
        run_via.italic = True
        run_via.font.size = Pt(7)
        run_via.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # Parágrafo 2: verbatim (normal, cinzento)
    p2 = c.add_paragraph()
    set_paragraph_spacing(p2, 0, 0)
    run_txt = p2.add_run(norma["texto"])
    run_txt.font.size = Pt(7.5)
    run_txt.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # -- Colunas de prazo --
    for i, col in enumerate(COLUNAS):
        cell = cells[i + 1]
        set_cell_shading(cell, bg)
        set_cell_width(cell, PRAZO_CM)
        set_cell_valign(cell, "center")
        set_cell_margins(cell, top=40, bottom=40, left=10, right=10)
        p = clear_cell(cell)
        set_paragraph_spacing(p, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col["key"] in norma["prazos"]:
            run = p.add_run("✓")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


# ---------------------------------------------------------------------------
# Documento principal
# ---------------------------------------------------------------------------

def gerar():
    doc = Document()

    # Página landscape A4
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = int(29.7 / 2.54 * 914400)   # EMU (1 inch = 914400 EMU)
    sec.page_height = int(21.0 / 2.54 * 914400)
    sec.left_margin   = int(MARGIN_CM / 2.54 * 914400)
    sec.right_margin  = int(MARGIN_CM / 2.54 * 914400)
    sec.top_margin    = int(1.5 / 2.54 * 914400)
    sec.bottom_margin = int(1.5 / 2.54 * 914400)

    # Estilo normal sem espaçamento extra
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # Título
    h = doc.add_paragraph()
    set_paragraph_spacing(h, 0, 100)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run("Prazos de Implementação — Regulamento 2023/0447 (cães e gatos)")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Subtítulo
    sub = doc.add_paragraph()
    set_paragraph_spacing(sub, 0, 160)
    r2 = sub.add_run(
        "Análise cruzada por categoria temporal — organização temática | "
        "Prazos contados a partir da data de entrada em vigor (vigésimo dia após publicação no JOUE)"
    )
    r2.italic = True
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Contar linhas totais: 1 cabeçalho + N grupos + N normas
    total_normas = sum(len(g["normas"]) for g in GRUPOS)
    total_rows = 1 + len(GRUPOS) + total_normas

    # Criar tabela
    table = doc.add_table(rows=total_rows, cols=1 + N_PRAZO_COLS)
    disable_autofit(table)
    set_table_width(table, USABLE_CM)
    set_table_borders(table)

    # Linha de cabeçalho
    build_header_row(table.rows[0])

    # Preencher grupos e normas
    row_idx = 1
    norma_counter = 0   # para alternar even/odd dentro de cada grupo
    for grupo in GRUPOS:
        build_group_row(table.rows[row_idx], grupo["titulo"])
        row_idx += 1
        norma_counter = 0
        for norma in grupo["normas"]:
            build_norma_row(table.rows[row_idx], norma, even=(norma_counter % 2 == 0))
            row_idx += 1
            norma_counter += 1

    # Nota de rodapé
    nota = doc.add_paragraph()
    set_paragraph_spacing(nota, 200, 0)
    rn = nota.add_run(
        "Nota: As datas específicas de 1 Jul. 2030 e 1 Jul. 2036 não dependem da data de entrada em vigor. "
        "Os prazos em anos são contados a partir da entrada em vigor, que ocorre no vigésimo dia após publicação no Jornal Oficial da UE. "
        "Texto verbatim em inglês conforme Art. 33 e artigos respectivos do Regulamento 2023/0447."
    )
    rn.italic = True
    rn.font.size = Pt(7.5)
    rn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out = "tabela_prazos_regulamento.docx"
    doc.save(out)
    print(f"Documento gerado: {out}")
    print(f"  Linhas na tabela: {total_rows} ({len(GRUPOS)} grupos, {total_normas} normas)")
    print(f"  Colunas: 1 (norma) + {N_PRAZO_COLS} (prazos) = {1 + N_PRAZO_COLS}")


if __name__ == "__main__":
    gerar()
