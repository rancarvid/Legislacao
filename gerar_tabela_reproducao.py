#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera tabela comparativa simples (sem cores) de critérios de reprodução:
Regulamento 2023/0447 × @codigo (DL 214/2013) × @rgbeac (proposta jun. 2025)
Output: tabela_reproducao_comparativa.docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ---------------------------------------------------------------------------
# DADOS DA TABELA
# Estrutura: (criterio, regulamento, codigo, rgbeac)
# ---------------------------------------------------------------------------

TITULO = "Quadro comparativo — Critérios de reprodução"
SUBTITULO = "Regulamento (UE) 2023/0447 × @codigo (proposta DL 214/2013, Anexo I) × @rgbeac (proposta jun. 2025, Art.º 70.º)"

SECOES = [
    {
        "titulo": "1. Idades mínimas de reprodução",
        "linhas": [
            (
                "Cadelas (fêmeas canídeos)",
                "A partir do 2.º estro\n(Anexo I, Pt. 3, n.º 2.2)",
                "Raças pequeno/médio porte: 3.º cio\nRaças grande porte: 2 anos\n(Art.º 8.º, n.º 2, al. a) + Anexo I)",
                "18 meses\n(Art.º 70.º, n.º 2)",
            ),
            (
                "Gatas (fêmeas felídeos)",
                "A partir dos 10 meses\n(Anexo I, Pt. 3, n.º 1.1)",
                "1 ano\n(Anexo I)",
                "12 meses\n(Art.º 70.º, n.º 3)",
            ),
            (
                "Cães (machos canídeos)",
                "Não fixado\n(Art.º 8.º não estabelece limite mínimo para machos)",
                "Raças pequeno/médio porte: 1,5 anos\nRaças grande porte: 2 anos\n(Anexo I)",
                "12 meses\n(Art.º 70.º, n.º 4)",
            ),
            (
                "Gatos (machos felídeos)",
                "Não fixado",
                "1 ano\n(Anexo I)",
                "12 meses\n(Art.º 70.º, n.º 4)",
            ),
        ],
    },
    {
        "titulo": "2. Idades máximas de reprodução",
        "linhas": [
            (
                "Cadelas",
                "Exame veterinário escrito obrigatório a partir dos 8 anos (não proibição)\n(Anexo I, Pt. 3, n.º 6.6)",
                "8 anos; ou 5 anos se for a 1.ª gestação\n(Anexo I)",
                "6 anos\n(Art.º 70.º, n.º 2)",
            ),
            (
                "Gatas",
                "Exame veterinário escrito obrigatório a partir dos 6 anos (não proibição)\n(Anexo I, Pt. 3, n.º 6.6)",
                "12 anos\n(Anexo I)",
                "6 anos\n(Art.º 70.º, n.º 3)",
            ),
            (
                "Cães machos",
                "Não fixado",
                "12 anos, ou mais se o estatuto sanitário o permitir\n(Anexo I)",
                "7 anos\n(Art.º 70.º, n.º 4)",
            ),
            (
                "Gatos machos",
                "Não fixado",
                "12 anos\n(Anexo I)",
                "7 anos\n(Art.º 70.º, n.º 4)",
            ),
        ],
    },
    {
        "titulo": "3. Limites de ninhadas e intervalos",
        "linhas": [
            (
                "N.º máximo de ninhadas por período",
                "Máx. 3 ninhadas em 2 anos (cadelas e gatas)\n(Anexo I, Pt. 3, n.º 3.3)",
                "Cadelas: 1 ninhada/ano; máx. 4–6 na vida\nGatas: 1 ninhada/ano (máx. 3 em 2 anos com aprovação MV); máx. 8–10 na vida\n(Anexo I)",
                "Máx. 4 ninhadas na vida reprodutiva\n(Art.º 70.º, n.ºs 2 e 3)",
            ),
            (
                "Intervalo mínimo entre ninhadas",
                "Período de recuperação de, pelo menos, 1 ano após atingir o limite de 3 ninhadas em 2 anos\n(Anexo I, Pt. 3, n.º 4.4)",
                "Cadelas: 1 ciclo reprodutivo\nGatas: 3–6 meses\n(Anexo I)",
                "Mínimo de 12 meses entre cada ninhada (cadelas e gatas)\n(Art.º 70.º, n.ºs 2 e 3)",
            ),
            (
                "Cios sucessivos (fêmeas)",
                "Não regulado explicitamente",
                "Proibido acasalamento em cios sucessivos\n(Art.º 8.º, n.º 2, al. a))",
                "Não regulado explicitamente (o intervalo de 12 meses implica proibição)",
            ),
        ],
    },
    {
        "titulo": "4. Cesarianas",
        "linhas": [
            (
                "Após 1.ª cesariana",
                "Não regulado",
                "Não regulado",
                "Exige atestado do MV assistente confirmando ausência de risco\n(Art.º 70.º, n.º 10)",
            ),
            (
                "Após 2 cesarianas",
                "Proibição de usar para reprodução\n(Anexo I, Pt. 3, n.º 5.5)",
                "Não regulado explicitamente",
                "Proibição absoluta de usar para reprodução\n(Art.º 70.º, n.º 12)",
            ),
        ],
    },
    {
        "titulo": "5. Consanguinidade e híbridos",
        "linhas": [
            (
                "Cruzamentos consanguíneos (pai/filho, irmãos, avós/netos)",
                "Proibidos, salvo autorização da autoridade competente para preservação de raças locais com reserva genética limitada\n(Art.º 8.º, n.º 6, al. a))",
                "Não regulado explicitamente; exclui animais com defeitos genéticos e malformações\n(Art.º 8.º, n.º 2, al. c))",
                "Sujeitos a portaria (coeficiente de consanguinidade e limites definidos por portaria)\n(Art.º 70.º, n.º 7)",
            ),
            (
                "Produção de híbridos (cruzamento interespecífico/inter-racial para fins comerciais)",
                "Proibida\n(Art.º 8.º, n.º 6, al. b))",
                "Proibida, exceto cruzamentos autorizados para melhoramento de raça ou investigação, ou cruzamentos acidentais sem fins comerciais\n(Art.º 8.º, n.º 3)",
                "Não regulado explicitamente",
            ),
        ],
    },
    {
        "titulo": "6. Conformação e genótipo",
        "linhas": [
            (
                "Traços conformacionais excessivos",
                "Proibida a reprodução; exige consulta prévia a MV; atos delegados até julho de 2030 definirão critérios\n(Art.º 8.º, n.ºs 2 e 3, al. b))",
                "Excluídos da reprodução animais com defeitos genéticos e malformações (ex: displasia da anca, rim poliquístico)\n(Art.º 8.º, n.º 2, al. c))",
                "Reprodução condicionada a parecer do MV que ateste ausência de efeito prejudicial com base em genótipo/fenótipo; restrição de raças por portaria\n(Art.º 70.º, n.ºs 8 e 9)",
            ),
            (
                "Testes genéticos e rastreios de saúde",
                "Atos delegados da Comissão definirão critérios (até julho de 2036 para genótipos)\n(Art.º 8.º, n.ºs 1 e 3, al. a))",
                "Excluídos animais com defeitos genéticos específicos\n(Art.º 8.º, n.º 2, al. c))",
                "Obrigatórios; critérios definidos por portaria específicos para cada raça\n(Art.º 70.º, n.º 7)",
            ),
            (
                "Porte semelhante dos progenitores (prevenção de distócia)",
                "Não regulado explicitamente",
                "Obrigatório\n(Art.º 8.º, n.º 2, al. b))",
                "Não regulado explicitamente",
            ),
        ],
    },
    {
        "titulo": "7. Reprodução assistida",
        "linhas": [
            (
                "Inseminação artificial (IA)",
                "Não regulado",
                "Não regulado",
                "Restrita a situações excecionais, previstas em portaria\n(Art.º 70.º, n.º 11)",
            ),
            (
                "Electroejaculação",
                "Não regulado",
                "Não regulado",
                "Proibida\n(Art.º 70.º, n.º 12)",
            ),
        ],
    },
    {
        "titulo": "8. Temperamento e comportamento",
        "linhas": [
            (
                "Critérios comportamentais para seleção de reprodutores",
                "Não regulado",
                "Excluídos animais com alterações comportamentais\n(Art.º 8.º, n.º 2, al. c))",
                "Obrigatório temperamento amistoso e confiante; proibida reprodução de animais perigosos ou com comportamento agressivo ou excessivamente tímido\n(Art.º 70.º, n.º 6)",
            ),
        ],
    },
    {
        "titulo": "9. Registo e obrigações administrativas",
        "linhas": [
            (
                "Registo de reprodutores",
                "Registo geral no sistema nacional (art.ºs 17.º e ss.)",
                "Não regulado especificamente neste artigo",
                "Inscrição obrigatória no SIAC como animal reprodutor\n(Art.º 70.º, n.º 6)",
            ),
            (
                "Obrigação de esterilização após fim da vida reprodutiva",
                "Não regulado",
                "Não regulado",
                "Obrigação do detentor de esterilizar o animal e garantir o seu bem-estar até ao final da vida\n(Art.º 70.º, n.º 13)",
            ),
        ],
    },
    {
        "titulo": "10. Separação das crias da progenitora",
        "linhas": [
            (
                "Cachorros (cães) — idade mínima de separação",
                "8 semanas (estabelecimentos de criação, abrigos e famílias de acolhimento)\nDerrogação por razões médicas com parecer escrito do MV\n(Anexo I, Pt. 4, n.º 3.3)",
                "Não regulado no Art.º 8.º nem no Anexo I reprodutivo",
                "10.ª semana de idade\n(com período de desmame gradual)\n(Art.º 70.º, n.º 5)",
            ),
            (
                "Gatinhos — idade mínima de separação",
                "Abrigos e famílias de acolhimento: 8 semanas\nEstabelecimentos de criação: 12 semanas\nDerrogação por razões médicas com parecer escrito do MV\n(Anexo I, Pt. 4, n.º 3.3)",
                "Não regulado no Art.º 8.º nem no Anexo I reprodutivo",
                "12.ª semana de idade\n(com período de desmame gradual)\n(Art.º 70.º, n.º 5)",
            ),
            (
                "Registo do parecer de derrogação",
                "O operador conserva o registo até ao último cachorro/gatinho da ninhada ser colocado no mercado\n(Anexo I, Pt. 4, n.º 3.3)",
                "Não regulado",
                "Não regulado explicitamente",
            ),
        ],
    },
    {
        "titulo": "11. Desmame (processo e critérios)",
        "linhas": [
            (
                "Processo de desmame — duração e método",
                "Introdução gradual de alimentos sólidos ao longo de um período não inferior a 7 dias\n(Anexo I, Pt. 1, n.º 4.4)",
                "Não regulado especificamente\n(o @codigo prevê programa de alimentação adequado à fase fisiológica — Art.º 8.º, n.º 2, al. a) por referência ao Anexo I)",
                "«Período de desmame gradual» salvaguardado antes da separação\n(Art.º 70.º, n.º 5 — não especifica duração)",
            ),
            (
                "Idade mínima para conclusão do desmame",
                "Não antes das 6 semanas de idade (tanto para cachorros como para gatinhos)\n(Anexo I, Pt. 1, n.º 4.4)",
                "Não regulado",
                "Não regulado explicitamente (a separação só é permitida a partir da 10.ª/12.ª semana, o que implica desmame anterior)",
            ),
            (
                "Alimentação de cachorros/gatinhos não desmamados",
                "Colostro nos primeiros 2 dias; depois leite da progenitora ou de cadela/gata lactante; se impossível, sucedâneo de leite específico para cachorros/gatinhos\n(Anexo I, Pt. 1, n.º 2.2)",
                "Não regulado no Anexo I reprodutivo\n(o Art.º 8.º, n.º 3, al. f) do @codigo prevê alojamento de fêmeas com crias em condições de bem-estar)",
                "Não regulado especificamente no Art.º 70.º",
            ),
            (
                "Monitorização do crescimento de não desmamados",
                "Todos os cachorros/gatinhos não desmamados devem receber leite, sucedâneo ou combinação em quantidade suficiente para ganho de peso constante\n(Anexo I, Pt. 1, n.º 3.3)",
                "Não regulado",
                "Não regulado",
            ),
        ],
    },
    {
        "titulo": "12. Socialização das crias",
        "linhas": [
            (
                "Idade de início e tipo de socialização",
                "A partir das 3 semanas de idade: oportunidades diárias e progressivas de contacto social com congéneres, seres humanos e, sempre que possível, outras espécies\n(Anexo I, Pt. 4, n.º 1.1)",
                "Não regulado especificamente na secção de reprodução\n(o @codigo exige enriquecimento ambiental — Art.º 13.º, n.º 5)",
                "Todos os alojamentos de criação de cães e gatos devem dispor de um plano de socialização e enriquecimento ambiental\n(Art.º 52.º, n.º 7 do @rgbeac)",
            ),
            (
                "Separação por comportamento agressivo",
                "Animais que representem ameaça ou causem stress excessivo devem ser mantidos separados\n(Anexo I, Pt. 4, n.º 1.1)",
                "Não regulado especificamente",
                "Não regulado especificamente",
            ),
        ],
    },
]

# ---------------------------------------------------------------------------
# GERAÇÃO DO DOCUMENTO
# ---------------------------------------------------------------------------

def set_cell_margins(cell, top=50, start=100, bottom=50, end=100):
    """Define margens internas da célula em twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def set_repeat_header(row):
    """Faz com que a linha repita como cabeçalho em cada página."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)

def make_document():
    doc = Document()

    # --- Margens da página ---
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # --- Título ---
    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_p.add_run(TITULO)
    run.bold = True
    run.font.size = Pt(14)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(SUBTITULO)
    run2.font.size = Pt(9)
    run2.italic = True

    data_p = doc.add_paragraph()
    data_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = data_p.add_run(f"Data: {datetime.date.today().strftime('%d/%m/%Y')}")
    run3.font.size = Pt(9)
    run3.italic = True

    doc.add_paragraph()  # espaço

    COLS = ["Critério", "Regulamento (UE) 2023/0447", "@codigo\n(proposta DL 214/2013)", "@rgbeac\n(proposta jun. 2025)"]
    COL_WIDTHS = [Cm(4.0), Cm(5.0), Cm(5.5), Cm(5.5)]

    for secao in SECOES:
        # --- Cabeçalho de secção ---
        h = doc.add_paragraph()
        run_h = h.add_run(secao["titulo"])
        run_h.bold = True
        run_h.font.size = Pt(10)
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(2)

        # --- Tabela ---
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Larguras das colunas
        for i, col in enumerate(table.columns):
            col.width = COL_WIDTHS[i]

        # Linha de cabeçalho
        hrow = table.rows[0]
        set_repeat_header(hrow)
        for i, col_title in enumerate(COLS):
            cell = hrow.cells[i]
            cell.width = COL_WIDTHS[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(col_title)
            run.bold = True
            run.font.size = Pt(8.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)

        # Linhas de dados
        for linha in secao["linhas"]:
            row = table.add_row()
            for i, texto in enumerate(linha):
                cell = row.cells[i]
                cell.width = COL_WIDTHS[i]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(texto)
                run.font.size = Pt(8)
                if i == 0:
                    run.bold = True
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                set_cell_margins(cell)

        doc.add_paragraph()  # espaço entre secções

    # --- Nota de rodapé ---
    nota = doc.add_paragraph()
    nota.paragraph_format.space_before = Pt(4)
    run_nota = nota.add_run(
        "Nota: @codigo e @rgbeac são propostas de legislação nacional, não legislação vigente. "
        "O Regulamento (UE) 2023/0447 é de aplicação direta nos Estados-Membros. "
        "A legislação vigente é o DL n.º 276/2001, de 17 de outubro, e o DL n.º 82/2019, de 27 de junho."
    )
    run_nota.font.size = Pt(7.5)
    run_nota.italic = True

    output_path = "tabela_reproducao_comparativa.docx"
    doc.save(output_path)
    print(f"Documento gerado: {output_path}")
    return output_path

if __name__ == "__main__":
    make_document()
