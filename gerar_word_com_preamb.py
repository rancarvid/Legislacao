#!/usr/bin/env python3
"""
Gera documento Word comparativo COM PREÂMBULO
Copia EXATAMENTE a configuração de gerar_word.py
Adiciona preâmbulo após os artigos mantendo estética idêntica
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm

# Importar do script original
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from gerar_word import (
    add_page_title, add_legenda, add_article_section,
    configurar_documento, cell_header, cell_body,
    set_table_width, remove_table_borders,
    set_cell_bg, set_cell_borders, add_run_styled,
    RGBColor
)
from gerar_comparativo_reuniao import ARTIGOS, COR
from gerar_html_com_preamb_v2 import extrair_preamb, TEMA_CONSIDERANDOS

# Extrair dados
ARTIGOS_DATA = ARTIGOS
preamb, preamb_por_tema = extrair_preamb()
TEMAS_ORDEM = list(TEMA_CONSIDERANDOS.keys())

# ─────────────────────────────────────────────────────────────────────────────
# CRIAR DOCUMENTO EXATAMENTE COMO gerar_word.py
# ─────────────────────────────────────────────────────────────────────────────

def criar_word_com_preamb(path):
    doc = Document()
    
    # ── CONFIGURAÇÃO EXATA (orientação landscape, margens 1.5cm) ─────────────
    configurar_documento(doc)
    
    # ── CAPA / TÍTULO ─────────────────────────────────────────────────────────
    add_page_title(doc)
    
    # ── LEGENDA ───────────────────────────────────────────────────────────────
    add_legenda(doc)
    
    # ── QUEBRA ANTES DOS ARTIGOS ──────────────────────────────────────────────
    doc.add_page_break()
    
    # ── ARTIGOS (UM POR PÁGINA) ───────────────────────────────────────────────
    for i, art in enumerate(ARTIGOS_DATA):
        add_article_section(doc, art)
        if i < len(ARTIGOS_DATA) - 1:
            doc.add_page_break()
    
    # ── QUEBRA ANTES DO PREÂMBULO ─────────────────────────────────────────────
    doc.add_page_break()
    
    # ── CABEÇALHO DO PREÂMBULO (IDÊNTICO AO CABEÇALHO DE ARTIGOS) ──────────────
    t_preamb_title = doc.add_table(rows=1, cols=1)
    set_table_width(t_preamb_title, pct=5000)
    remove_table_borders(t_preamb_title)
    
    cell_preamb = t_preamb_title.cell(0, 0)
    set_cell_bg(cell_preamb, "1A1A2E")
    set_cell_borders(cell_preamb, "1A1A2E")
    p_preamb = cell_preamb.paragraphs[0]
    p_preamb.paragraph_format.space_before = Pt(8)
    p_preamb.paragraph_format.space_after = Pt(8)
    p_preamb.paragraph_format.left_indent = Pt(10)
    add_run_styled(p_preamb, "PREÂMBULO — Considerandos", bold=True, 
                   font_size=11, color_hex="FFFFFF")
    
    doc.add_paragraph()
    
    # ── TEMAS DO PREÂMBULO ────────────────────────────────────────────────────
    for tema in TEMAS_ORDEM:
        considerandos = preamb_por_tema.get(tema, [])
        if not considerandos:
            continue
        
        # Cabeçalho do tema (cor similar a artigos)
        t_tema = doc.add_table(rows=1, cols=2)
        set_table_width(t_tema, pct=5000)
        remove_table_borders(t_tema)
        
        c_tema_id = t_tema.cell(0, 0)
        set_cell_bg(c_tema_id, "1A1A2E")
        set_cell_borders(c_tema_id, "1A1A2E")
        p_tema_id = c_tema_id.paragraphs[0]
        p_tema_id.paragraph_format.space_before = Pt(8)
        p_tema_id.paragraph_format.space_after = Pt(8)
        p_tema_id.paragraph_format.left_indent = Pt(10)
        add_run_styled(p_tema_id, f"TEMA {TEMAS_ORDEM.index(tema)+1}", 
                       bold=True, font_size=11, color_hex="FFFFFF")
        
        c_tema_label = t_tema.cell(0, 1)
        set_cell_bg(c_tema_label, "2C3E6B")
        set_cell_borders(c_tema_label, "2C3E6B")
        p_tema_label = c_tema_label.paragraphs[0]
        p_tema_label.paragraph_format.space_before = Pt(8)
        p_tema_label.paragraph_format.space_after = Pt(8)
        p_tema_label.paragraph_format.left_indent = Pt(10)
        add_run_styled(p_tema_label, tema, bold=True, font_size=11, 
                       color_hex="7EC8E3")
        
        # Considerandos do tema
        for cons in considerandos:
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            
            # Considerando EN
            t_cons_en = doc.add_table(rows=2, cols=1)
            set_table_width(t_cons_en, pct=5000)
            ref_en = f"Considerando {cons['numero']}  ·  {cons['id']}  |  EN"
            cell_header(t_cons_en.cell(0, 0), ref_en, COR["regulamento_header"])
            cell_body(t_cons_en.cell(1, 0),
                     cons['regulamento']['texto'],
                     COR["regulamento_body"],
                     font_size=9.5)
            
            # Considerando PT
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            t_cons_pt = doc.add_table(rows=2, cols=1)
            set_table_width(t_cons_pt, pct=5000)
            ref_pt = f"Considerando {cons['numero']}  ·  PT"
            cell_header(t_cons_pt.cell(0, 0), ref_pt, "2471A3")
            cell_body(t_cons_pt.cell(1, 0),
                     cons['regulamento']['traducao'],
                     COR["regulamento_trad"],
                     font_size=9.5,
                     italic=True)
            
            # Separador
            doc.add_paragraph()
            p_sep = doc.add_paragraph()
            p_sep.paragraph_format.space_before = Pt(0)
            p_sep.paragraph_format.space_after = Pt(8)
            run_sep = p_sep.add_run("─" * 80)
            run_sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run_sep.font.size = Pt(8)
    
    # Salvar
    doc.save(path)
    print(f"✅ {path}")

if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    output = os.path.join(base, "comparativo_reuniao_exemplo_preamb.docx")
    criar_word_com_preamb(output)
    print(f"   - Orientação: Landscape (29.7 x 21.0 cm)")
    print(f"   - Margens: 1.5cm")
    print(f"   - Artigos: {len(ARTIGOS_DATA)}")
    print(f"   - Preâmbulo: {len(TEMAS_ORDEM)} temas")
    print(f"   - Estética: IDÊNTICA a comparativo_reuniao_exemplo.docx")
