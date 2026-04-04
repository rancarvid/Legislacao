from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'start', 'bottom', 'end'):
        tag = f'w:{edge}'
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_header_row(table, col_headers):
    row = table.rows[0]
    for i, hdr in enumerate(col_headers):
        cell = row.cells[i]
        cell.text = hdr
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'

def add_data_row(table, values, italic_col0=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        run.font.name = 'Arial'
        if i == 0 and italic_col0:
            run.italic = True

def add_section_title(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def create_table(doc, col_widths_cm, header_row):
    table = doc.add_table(rows=1, cols=len(col_widths_cm))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, w in enumerate(col_widths_cm):
        table.columns[i].width = Cm(w)
    add_header_row(table, header_row)
    return table

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Title
title = doc.add_paragraph()
run = title.add_run('Quadro Comparativo — Artigos 6.º e 7.º do Regulamento (UE) 2023/0447')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Arial'
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
run2 = sub.add_run(
    'Correspondência norma a norma entre as determinações do Regulamento e as propostas @rgbeac (jun. 2025) e @codigo (DL n.º 214/2013)'
)
run2.font.size = Pt(9)
run2.italic = True
run2.font.name = 'Arial'
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(8)

HDRS = [
    'Critério / Determinação do Regulamento',
    'Regulamento (UE) 2023/0447\n(texto PT-PT)',
    '@rgbeac\n(proposta jun. 2025)',
    '@codigo\n(proposta DL n.º 214/2013)',
]
COL_W = [5.5, 4.5, 4.5, 4.5]

# ============================================================
# ARTIGO 6.º — PRINCÍPIOS GERAIS DE BEM-ESTAR
# ============================================================
add_section_title(doc, 'Artigo 6.º — Princípios gerais de bem-estar')
intro_p = doc.add_paragraph()
run_intro = intro_p.add_run(
    'O art.º 6.º impõe aos operadores a aplicação de princípios gerais de bem-estar aos cães e gatos '
    'criados ou detidos nos seus estabelecimentos. Analisa-se abaixo cada alínea de forma desagregada.'
)
run_intro.font.size = Pt(8.5)
run_intro.font.name = 'Arial'
intro_p.paragraph_format.space_after = Pt(4)

t6 = create_table(doc, COL_W, HDRS)

rows6 = [
    # al. (a) — água e alimentação
    (
        'al. (a) — Água e alimentação\nFornecimento adequado de água e alimentos\n(nutrição e hidratação)',
        'Os cães e gatos recebem água e alimentos de qualidade e numa quantidade que assegura '
        'a sua boa e adequada nutrição e hidratação.\n(al. (a) do art.º 6.º)',
        'Art.º 7.º, n.º 1, al. a) — proibição de fome, sede ou malnutrição.\n'
        'Art.º 10.º, n.º 1, al. a), ii) — alimentos saudáveis, adequados e convenientes; '
        'acesso permanente a água potável.\n'
        'Art.º 51.º, n.º 1 — programa de alimentação bem definido, de valor nutritivo adequado, '
        'em quantidade suficiente para as necessidades das espécies e dos indivíduos.\n'
        'Art.º 51.º, n.º 6 — acesso a água potável sem restrição '
        '(salvo razões médico-veterinárias).\nCoberto.',
        'Art.º 4.º — dever de assegurar as necessidades básicas de bem-estar.\n'
        'Art.º 46.º, n.º 1 — programa de alimentação bem definido, de valor nutritivo adequado, '
        'em quantidade suficiente para as necessidades das espécies e dos indivíduos '
        '(nos locais de criação, manutenção, venda, centros de recolha e hospedagem).\n'
        'Art.º 46.º, n.º 6 — acesso a água potável sem restrição '
        '(salvo razões médico-veterinárias).\n'
        'Coberto no âmbito de estabelecimentos '
        '(art.º 46.º não se aplica à detenção doméstica individual).',
    ),
    # al. (b).1 — ambiente físico: qualidade/limpeza/segurança
    (
        'al. (b) — Ambiente físico (1/2)\nLimpo, seguro e confortável\n(espaço, ar, temperatura, luz, clima)',
        'Os cães e gatos são mantidos num ambiente físico adequado, regularmente limpo, seguro e '
        'confortável, especialmente em termos de espaço, qualidade do ar, temperatura, iluminação '
        'e proteção face a condições climáticas adversas.\n(al. (b) do art.º 6.º)',
        'Art.º 7.º, n.º 1, al. b) — proibição de desconforto físico ou térmico.\n'
        'Art.º 10.º, n.º 1, al. a), v) — abrigo adequado com proteção de condições '
        'atmosféricas adversas (frio, chuva, sol ou calor excessivos); cama seca, limpa '
        'e confortável.\n'
        'Art.º 48.º, n.º 1 — temperatura, ventilação, luminosidade e obscuridade '
        'adequadas à manutenção do bem-estar e conforto das espécies.\n'
        'Art.º 48.º, n.º 3 — luz preferencialmente natural; luz artificial o mais próxima '
        'do espectro solar.\n'
        'Parcialmente coberto (sem referência explícita a qualidade do ar como parâmetro '
        'distinto da ventilação).',
        'Art.º 5.º, n.º 1 — condições de detenção devem salvaguardar bem-estar.\n'
        'Art.º 14.º, n.º 1 — temperatura, ventilação, luminosidade e obscuridade das '
        'instalações adequadas ao bem-estar das espécies.\n'
        'Art.º 14.º, n.º 3 — luz preferencialmente natural; artificial o mais próxima '
        'do espectro solar.\n'
        'Parcialmente coberto (sem referência explícita a qualidade do ar como parâmetro '
        'distinto da ventilação).',
    ),
    # al. (b).2 — espaço suficiente / liberdade de movimento
    (
        'al. (b) — Ambiente físico (2/2)\nEspaço anti-sobrelotação;\nliberdade de movimentos',
        'Os cães e gatos são mantidos num espaço suficientemente grande para prevenir a '
        'sobrelotação e garantir a facilidade de movimentação.\n(al. (b) do art.º 6.º)',
        'Art.º 10.º, n.º 1, al. a), iv) — liberdade de movimento; proibição de contenção '
        'permanente; espaço e enriquecimento ambiental adequados.\n'
        'Art.º 47.º, n.º 1 — espaço adequado às necessidades fisiológicas e etológicas; '
        'prática de exercício físico; fuga de animais agressores.\n'
        'Portaria prevista no art.º 58.º fixará requisitos específicos.\n'
        'Parcialmente coberto (sobrelotação não referenciada de forma autónoma).',
        'Art.º 5.º, n.º 3, al. b) — proibição de restringir a liberdade de movimentos de '
        'forma a impedir levantar-se, deitar-se ou virar-se sobre si próprios.\n'
        'Art.º 13.º, n.º 1, als. a) e b) — espaço adequado às necessidades fisiológicas '
        'e etológicas; prática de exercício físico; fuga de animais agressores.\n'
        'Parcialmente coberto (sobrelotação não referenciada de forma autónoma).',
    ),
    # al. (c).1 — segurança, limpeza e saúde
    (
        'al. (c) — Saúde e higiene (1/2)\nSegurança, limpeza, boa saúde;\nprevenção de doenças, lesões e dor',
        'Os cães e gatos são mantidos seguros, limpos e com boa saúde, prevenindo doenças, '
        'lesões e dor.\n(al. (c) do art.º 6.º)',
        'Art.º 7.º, n.º 1, al. c) — proibição de dor, lesão física ou doença.\n'
        'Art.º 10.º, n.º 1, al. a), iii) — condições higiossanitárias.\n'
        'Art.º 10.º, n.º 1, al. f) — tratamento veterinário preventivo, paliativo ou curativo.\n'
        'Art.º 52.º, n.º 2 — todos os animais alvo de inspeção no início e final do dia '
        'e a cada quatro horas; primeiros cuidados imediatos a animais doentes ou lesionados.\n'
        'Coberto.',
        'Art.º 4.º — dever de salvaguardar a saúde.\n'
        'Art.º 6.º — dever de assegurar cuidados médico-veterinários ao animal ferido '
        'ou doente.\n'
        'Art.º 47.º, n.º 3 — inspeção diária de todos os animais; primeiros cuidados '
        'imediatos a animais doentes, lesionados ou com alterações comportamentais.\n'
        'Coberto em termos gerais (inspeção diária; sem diferenciação de frequência '
        'para animais vulneráveis).',
    ),
    # al. (c).2 — prevenção via maneio e reprodução
    (
        'al. (c) — Saúde e higiene (2/2)\nPrevenção via maneio, manuseamento\ne reprodução adequados',
        'Prevenção de doenças, lesões e dor, nomeadamente através de práticas de maneio, '
        'manuseamento e reprodução adequadas.\n(al. (c) do art.º 6.º)',
        'Art.º 12.º, n.º 1, al. b) — proibição de práticas que causem sofrimento, dano, '
        'stresse ou angústia.\n'
        'Art.º 52.º, n.ºs 3 e 5 — maneio feito de forma a não causar dores, sofrimento '
        'ou distúrbios desnecessários; meios de contenção não devem causar ferimentos, '
        'dores ou angústia.\n'
        'Art.º 70.º (normas reprodutivas) — limitações à reprodução para proteção da saúde.\n'
        'Coberto.',
        'Art.º 5.º, n.º 3, al. c) — proibição de usar equipamentos de maneio, contenção '
        'ou treino que causem sofrimento desnecessário ou lesões.\n'
        'Art.º 47.º, n.ºs 4 e 5 — manuseamento feito de forma a não causar dores, '
        'sofrimento ou distúrbios desnecessários; meios de contenção não devem causar '
        'ferimentos, dores ou angústia.\n'
        'Art.º 8.º, n.º 1 e 2 — reprodução planeada e parâmetros reprodutivos.\n'
        'Coberto.',
    ),
    # al. (d).1 — comportamentos específicos da espécie
    (
        'al. (d) — Comportamento (1/2)\nExpressão de comportamentos da espécie\ne sociais não nocivos',
        'Os cães e gatos são mantidos num ambiente que lhes permite exibir comportamentos '
        'específicos da espécie e comportamentos sociais não nocivos.\n(al. (d) do art.º 6.º)',
        'Art.º 7.º, n.º 1, al. d) — direito de expressar padrões normais de comportamento.\n'
        'Art.º 10.º, n.º 1, al. a), iv) — espaço e enriquecimento ambiental adequados para a '
        'expressão dos comportamentos naturais.\n'
        'Art.º 10.º, n.º 1, al. a), vi) — contato social adequado à espécie, de acordo com '
        'a idade e atividade.\n'
        'Art.º 47.º, n.º 4 — enriquecimento ambiental complexo e estimulante com materiais '
        'e equipamento que estimulem a expressão do repertório de comportamentos naturais.\n'
        'Coberto.',
        'Art.º 5.º, n.º 1 — condições de detenção que salvaguardem parâmetros de bem-estar.\n'
        'Art.º 13.º, n.º 5 — instalações equipadas com materiais e equipamento que '
        'estimulem a expressão dos comportamentos naturais (substrato, cama, ninhos, '
        'ramos, buracos, locais para banhos).\n'
        'Parcialmente coberto (enriquecimento ambiental para comportamentos naturais; '
        'sem referência explícita a comportamentos sociais não nocivos).',
    ),
    # al. (d).2 — relação positiva com humanos
    (
        'al. (d) — Comportamento (2/2)\nRelação positiva com os seres humanos',
        'Os cães e gatos são mantidos num ambiente que lhes permite estabelecer uma relação '
        'positiva com os seres humanos.\n(al. (d) do art.º 6.º)',
        'Art.º 10.º, n.º 1, al. c) — dever de educar o animal com recurso a métodos de '
        'reforço positivo visando a sua vinculação e integração positivas no espaço familiar '
        'e no meio ambiente.\n'
        'Art.º 52.º, n.º 4 — interações do pessoal técnico com os animais devem ser '
        'positivas, regulares, previsíveis e não forçadas.\nCoberto.',
        'Art.º 7.º, n.º 1 — dever de promover o treino dos cães com vista à socialização '
        'e obediência.\n'
        'Art.º 7.º, n.º 2 — treino de acordo com boas práticas.\n'
        'Parcialmente coberto (apenas cães; sem referência explícita a relação positiva '
        'com os seres humanos como fim autónomo).',
    ),
    # al. (e).1 — prevenção/redução de estímulos negativos
    (
        'al. (e) — Estado mental (1/4)\nPrevenir/reduzir estímulos negativos\n(duração e intensidade)',
        'Os cães e gatos são mantidos de forma a otimizar o seu estado mental, prevenindo '
        'ou reduzindo estímulos negativos em duração e intensidade.\n(al. (e) do art.º 6.º)',
        'Art.º 7.º, n.º 1, al. e) — proibição de situações que provoquem stresse, medo '
        'ou ansiedade injustificados.\n'
        'Art.º 12.º, n.º 1, al. b) — proibição de práticas que causem sofrimento, dano, '
        'stresse ou angústia.\nParcialmente coberto (abordagem proibitiva; não preventiva).',
        'Art.º 5.º, n.º 3 — proibição de violência e sofrimento.\n'
        'Sem referência explícita à gestão ativa de estímulos negativos.',
    ),
    # al. (e).2 — maximização de estímulos positivos
    (
        'al. (e) — Estado mental (2/4)\nMaximizar estímulos positivos\n(duração e intensidade)',
        'Os cães e gatos são mantidos de forma a maximizar oportunidades para estímulos '
        'positivos em duração e intensidade.\n(al. (e) do art.º 6.º)',
        'Art.º 10.º, n.º 1, al. a), i) — dever de proporcionar atenção, supervisão, '
        'controlo, exercício físico e estímulo mental.\n'
        'Art.º 47.º, n.ºs 6 e 7 — acesso diário a área de exercício (mínimo 2×30 min); '
        'equipamento de enriquecimento ambiental e agilidade.\nParcialmente coberto.',
        'Art.º 13.º, n.º 5 — materiais e equipamento que estimulem a expressão dos '
        'comportamentos naturais.\n'
        'Sem norma positiva explícita sobre maximização de estímulos positivos.',
    ),
    # al. (e).3 — prevenção de comportamentos repetitivos anormais
    (
        'al. (e) — Estado mental (3/4)\nPrevenir comportamentos repetitivos anormais\nou indicativos de bem-estar negativo',
        'Os cães e gatos são mantidos de forma a prevenir o desenvolvimento de comportamentos '
        'repetitivos anormais ou outros comportamentos indicativos de bem-estar negativo.\n'
        '(al. (e) do art.º 6.º)',
        'Não regulado explicitamente como norma autónoma.\n'
        '(O art.º 7.º, n.º 1, al. d) e e) e o art.º 47.º do @rgbeac indiretamente '
        'impõem condições que previnem o seu aparecimento.)',
        'Não regulado explicitamente.',
    ),
    # al. (e).4 — necessidades individuais do animal
    (
        'al. (e) — Estado mental (4/4)\nNecessidades individuais do animal\n(als. (a) a (d))',
        'Os cães e gatos são mantidos tendo em conta as necessidades individuais do animal nos '
        'domínios referidos nas alíneas (a) a (d).\n(al. (e) do art.º 6.º)',
        'Art.º 10.º, n.º 1, al. a) — bem-estar assegurado "de acordo com a sua espécie, '
        'raça, idade e necessidades físicas e etológicas".\n'
        'Art.º 48.º, n.º 2 — fatores ambientais adequados às necessidades específicas '
        'de animais em fase reprodutiva, recém-nascidos ou doentes.\n'
        'Parcialmente coberto (necessidades da espécie/raça/idade/fase fisiológica; '
        'não explicitamente individuais).',
        'Art.º 4.º — necessidades básicas de bem-estar (sem referência à individualidade '
        'do animal).\n'
        'Art.º 14.º, n.º 2 — fatores ambientais adequados às necessidades específicas de '
        'animais em fase reprodutiva, recém-nascidos ou doentes.\n'
        'Parcialmente coberto.',
    ),
]

for r in rows6:
    add_data_row(t6, r)

# space
doc.add_paragraph()

# ============================================================
# ARTIGO 7.º — OBRIGAÇÕES GERAIS DE BEM-ESTAR
# ============================================================
add_section_title(doc, 'Artigo 7.º — Obrigações gerais de bem-estar')

t7 = create_table(doc, COL_W, HDRS)

rows7 = [
    # n.º 1 — responsabilidade geral
    (
        'n.º 1 — Responsabilidade geral\nBem-estar dos animais;\nminimizar riscos',
        'Os operadores são responsáveis pelo bem-estar dos cães e gatos detidos nos '
        'estabelecimentos sob a sua responsabilidade e controlo, devendo minimizar quaisquer '
        'riscos para o bem-estar desses animais.\n(n.º 1 do art.º 7.º)',
        'Art.º 9.º, n.º 1 — dever geral de todos os cidadãos de notificar infrações.\n'
        'Art.º 10.º, n.º 1 — obrigações especiais dos detentores: assegurar bem-estar.\n'
        'Sem distinção entre detentor doméstico e operador de estabelecimento comercial.\n'
        'Parcialmente coberto.',
        'Art.º 4.º — detenção responsável: cabe aos detentores assegurar as necessidades '
        'básicas de bem-estar e salvaguardar a saúde.\n'
        'Sem atribuição explícita de responsabilidade ao "operador" como sujeito '
        'autónomo de estabelecimento.\nParcialmente coberto.',
    ),
    # n.º 2 — famílias de acolhimento: responsabilidade do operador
    (
        'n.º 2 — Famílias de acolhimento (1/2)\nResponsabilidade do operador;\ninformação sobre bem-estar e necessidades individuais',
        'No caso das famílias de acolhimento, a responsabilidade incumbe ao operador em cujo '
        'nome os cães ou gatos são detidos. O operador presta à família informações adequadas '
        'sobre as obrigações de bem-estar e as necessidades individuais dos animais.\n'
        '(n.º 2 do art.º 7.º)',
        'Art.º 4.º — define "Família de acolhimento temporário" como detentor transitório '
        '(45 dias, extensível).\n'
        'Art.º 101.º, n.ºs 3 e 4 — contrato escrito entre titular e família; registo no SIAC.\n'
        'Art.º 101.º, n.º 5 — destino do animal é da responsabilidade do titular.\n'
        'Parcialmente coberto (responsabilidade no titular; mas não se fixa '
        'o dever de informação ao acolhedor sobre bem-estar e necessidades individuais).',
        'Art.º 9.º — define abandono; não prevê o conceito de família de acolhimento.\n'
        'Não coberto.',
    ),
    # n.º 2 — famílias de acolhimento: limite numérico de animais
    (
        'n.º 2 — Famílias de acolhimento (2/2)\nMáx. 5 animais ou 1 ninhada\n(com/sem mãe) por família',
        'Os operadores não colocam mais do que um total combinado de cinco cães ou gatos '
        'ou uma ninhada, com ou sem a mãe, numa família de acolhimento em qualquer momento.\n'
        '(n.º 2 do art.º 7.º)',
        'Art.º 101.º — acolhimento temporário: sem fixação de limite numérico de animais '
        'por família.\nNão coberto.',
        'Não regulado.',
    ),
    # n.º 3 — derrogação dos EM para maior número em FAT
    (
        'n.º 3 — Derrogação (Estados-Membros)\nMaior número autorizado\nse espaço e cuidadores suficientes',
        'O Estado-Membro pode autorizar maior número de cães, gatos ou ninhadas na família '
        'de acolhimento, desde que as instalações disponham de espaço suficiente (incluindo '
        'exterior) e o número de cuidadores seja suficiente.\n(n.º 3 do art.º 7.º)',
        'Não regulado.',
        'Não regulado.',
    ),
    # n.º 4 — proibição de crueldade, abusos e maus-tratos
    (
        'n.º 4 — Proibição de crueldade e maus-tratos\nInclui participação\nem atividades de crueldade',
        'Os operadores não sujeitam qualquer cão ou gato a crueldade, abusos ou maus-tratos, '
        'incluindo fazendo-os participar em atividades suscetíveis de resultar em crueldade, '
        'abusos ou maus-tratos.\n(n.º 4 do art.º 7.º)',
        'Art.º 12.º, n.º 1, al. a) — proibição de causar a morte em violação das normas.\n'
        'Art.º 12.º, n.º 1, al. b) — proibição de práticas que causem sofrimento, dano, '
        'stresse ou angústia.\n'
        'Art.º 12.º, n.º 1, al. t) — proibição de utilização em eventos que envolvam '
        'crueldade, maus-tratos, sofrimento ou morte.\nCoberto.',
        'Art.º 5.º, n.º 3 — proibição de violência: atos que infligem, sem necessidade, '
        'morte, sofrimento, abuso ou lesões a um animal.\n'
        'Art.º 5.º, n.º 4 — proibição de fins didáticos e lúdicos com dor ou sofrimentos '
        'consideráveis.\nCoberto.',
    ),
    # n.º 5 — proibição de abandono
    (
        'n.º 5 — Proibição de abandono',
        'Os operadores não abandonam os cães ou gatos criados ou detidos por eles.\n'
        '(n.º 5 do art.º 7.º)',
        'Art.º 12.º, n.º 1, al. e) — proibição expressa de abandono.\n'
        'Art.º 9.º, n.º 1 — dever de notificação de infrações (inclui abandono).\n'
        'Coberto.',
        'Art.º 9.º — define abandono; não estabelece proibição autónoma expressa '
        '(proibição implícita pela definição).\n'
        'Parcialmente coberto.',
    ),
    # n.º 6 — realojamento ao cessar atividade
    (
        'n.º 6 — Realojamento ao cessar atividade\nAdoção ou transferência\npara outros operadores/adquirentes',
        'Antes de os operadores cessarem as atividades num estabelecimento, asseguram que '
        'os cães ou gatos aí detidos são realojados, seja assumindo eles próprios a '
        'propriedade de animal de companhia, seja transferindo a responsabilidade para outros '
        'operadores ou adquirentes.\n(n.º 6 do art.º 7.º)',
        'Art.º 101.º, n.º 5 — destino do animal é da responsabilidade do titular; mas esta '
        'norma aplica-se ao acolhimento temporário, não ao encerramento de estabelecimento.\n'
        'Não coberto (para encerramento de estabelecimento comercial).',
        'Não regulado.',
    ),
    # n.º 7 — número suficiente de cuidadores
    (
        'n.º 7 — Cuidadores suficientes e competentes\n(competências do art.º 12.º)',
        'Os operadores asseguram que os cães e gatos são manuseados por um número suficiente '
        'de cuidadores de animais para satisfazer as necessidades de bem-estar, e que esses '
        'cuidadores possuem as competências exigidas no artigo 12.º.\n(n.º 7 do art.º 7.º)',
        'Art.º 52.º, n.º 1 — observação diária dos animais e maneio assegurados por pessoal '
        'técnico com formação teórica e prática específica certificada pela DGAV e em número '
        'adequado à quantidade e espécies animais alojados.\n'
        'Parcialmente coberto (número adequado e formação exigidos; sem remissão para '
        'catálogo de competências específico equivalente ao art.º 12.º do Regulamento).',
        'Art.º 47.º, n.º 1 — observação diária dos animais e maneio assegurados por pessoal '
        'com aptidão para o efeito e em número adequado à quantidade e espécies animais.\n'
        'Art.º 47.º, n.º 2 — maneio por pessoal com experiência ou formação adequada, '
        'sob orientação de médico veterinário.\n'
        'Parcialmente coberto (número adequado e formação exigidos; sem catálogo de '
        'competências específico equivalente ao art.º 12.º do Regulamento).',
    ),
    # n.º 8 — monitorização de indicadores baseados nos animais
    (
        'n.º 8 — Monitorização de indicadores\nComportamento e aparência física;\nações baseadas em resultados',
        'Os operadores asseguram o bem-estar dos cães ou gatos monitorizando indicadores '
        'baseados nos animais relativos ao comportamento e à aparência física, e adotando '
        'ações com base nos resultados desse monitoramento.\n'
        'A Comissão adotará atos delegados que estabelecerão os indicadores e os métodos '
        'de medição.\n(n.ºs 8 e 9 do art.º 7.º)',
        'Art.º 10.º, n.º 1, al. b) — dever de vigiar o animal de modo a evitar que cause '
        'danos (obrigação de vigilância; não equivale a monitorização de indicadores de '
        'bem-estar).\n'
        'Art.º 52.º, n.º 2 — inspeção dos animais no início e final do dia e a cada quatro '
        'horas; primeiros cuidados imediatos.\n'
        'Parcialmente coberto (periodicidade de inspeção regulada; sem sistema formal de '
        'indicadores baseados nos animais nem obrigação de ação estruturada com base nos '
        'resultados).',
        'Art.º 47.º, n.º 3 — inspeção diária de todos os animais; primeiros cuidados '
        'imediatos a doentes, lesionados ou com alterações comportamentais.\n'
        'Parcialmente coberto (inspeção diária; sem sistema formal de indicadores de '
        'bem-estar nem atos delegados equivalentes).',
    ),
]

for r in rows7:
    add_data_row(t7, r)

output_path = '/home/user/Legislacao/tabela_art6_art7_comparativo.docx'
doc.save(output_path)
print(f'Documento gerado: {output_path}')
