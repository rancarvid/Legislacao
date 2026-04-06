#!/usr/bin/env python3
"""
Converte HTML comparativo para DOCX de forma robusta
"""

import re
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('comparativo_reuniao_exemplo_preamb_teste_v2.html', 'r', encoding='utf-8') as f:
    html = f.read()

print("Extraindo dados do HTML...")

# Extrair ARTIGOS de forma mais robusta
try:
    m = re.search(r'const ARTIGOS\s*=\s*(\[[\s\S]*?\n\]\s*;)', html)
    if m:
        json_str = m.group(1).rstrip()[:-1]  # Remove ];
        ARTIGOS = json.loads(json_str)
        print(f"✓ {len(ARTIGOS)} artigos extraídos")
except Exception as e:
    print(f"Erro ao extrair ARTIGOS: {e}")
    ARTIGOS = []

# Extrair PREAMB_POR_TEMA
try:
    m = re.search(r'const PREAMB_POR_TEMA\s*=\s*(\{[\s\S]*?\n\}\s*;)', html)
    if m:
        json_str = m.group(1).rstrip()[:-1]  # Remove };
        PREAMB_POR_TEMA = json.loads(json_str)
        print(f"✓ {len(PREAMB_POR_TEMA)} temas extraídos")
except Exception as e:
    print(f"Erro ao extrair PREAMB_POR_TEMA: {e}")
    PREAMB_POR_TEMA = {}

# Extrair TEMAS_PREAMB
try:
    m = re.search(r'const TEMAS_PREAMB\s*=\s*(\[[\s\S]*?\]\s*;)', html)
    if m:
        json_str = m.group(1).rstrip()[:-1]  # Remove ];
        TEMAS_PREAMB = json.loads(json_str)
        print(f"✓ {len(TEMAS_PREAMB)} temas em ordem extraídos")
except Exception as e:
    print(f"Erro ao extrair TEMAS_PREAMB: {e}")
    TEMAS_PREAMB = []

if not ARTIGOS:
    print("Erro: sem dados para gerar DOCX")
    exit(1)

# ── Criar DOCX ─────────────────────────────────────────────────────────────
doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# Título
title = doc.add_heading('Regulamento (UE) 2023/0447', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Artigos + Preâmbulo (Comparativo)', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── ARTIGOS ────────────────────────────────────────────────────────────────
doc.add_heading('PARTE I: ARTIGOS (39)', level=2)

for art in ARTIGOS:
    # Cabeçalho
    h3 = doc.add_heading(f"{art['id']} — {art.get('tema', '')}", level=3)
    
    # Regulamento EN
    doc.add_heading('Regulamento (EN)', level=4)
    if art.get('regulamento'):
        p = doc.add_paragraph()
        p.add_run('Ref: ').bold = True
        p.add_run(art['regulamento'].get('ref', ''))
        
        p = doc.add_paragraph()
        p.add_run(art['regulamento'].get('texto', ''))
    
    # Regulamento PT (tradução)
    doc.add_heading('Tradução (PT)', level=4)
    if art.get('regulamento'):
        p = doc.add_paragraph()
        p.add_run(art['regulamento'].get('traducao', ''))
    
    # RGBEAC
    if art.get('rgbeac', {}).get('texto'):
        doc.add_heading('@rgbeac', level=4)
        p = doc.add_paragraph()
        p.add_run('Ref: ').bold = True
        p.add_run(art['rgbeac'].get('ref', ''))
        p = doc.add_paragraph()
        p.add_run(art['rgbeac'].get('texto', ''))
    
    # Código
    if art.get('codigo', {}).get('texto'):
        doc.add_heading('@código', level=4)
        p = doc.add_paragraph()
        p.add_run('Ref: ').bold = True
        p.add_run(art['codigo'].get('ref', ''))
        p = doc.add_paragraph()
        p.add_run(art['codigo'].get('texto', ''))
    
    # Legislação
    if art.get('legislacao', {}).get('texto'):
        doc.add_heading('@legislação', level=4)
        p = doc.add_paragraph()
        p.add_run('Ref: ').bold = True
        p.add_run(art['legislacao'].get('ref', ''))
        p = doc.add_paragraph()
        p.add_run(art['legislacao'].get('texto', ''))
    
    # Divergência
    if art.get('divergencia'):
        div = art['divergencia']
        doc.add_heading('Divergências', level=4)
        for key in ['legislacao', 'codigo', 'rgbeac', 'sumario']:
            if div.get(key):
                p = doc.add_paragraph()
                p.add_run(f'{key}: ').bold = True
                p.add_run(div[key])
    
    # Necessidade alteração
    if art.get('necessidade_alteracao'):
        p = doc.add_paragraph()
        p.add_run('Necessidade alteração: ').bold = True
        p.add_run(art['necessidade_alteracao'])
    
    # Notas
    if art.get('notas'):
        p = doc.add_paragraph()
        p.add_run('Notas: ').bold = True
        p.add_run(art['notas'])
    
    # Separador
    doc.add_paragraph('_' * 80)

# ── PREÂMBULO ──────────────────────────────────────────────────────────────
if PREAMB_POR_TEMA and TEMAS_PREAMB:
    doc.add_page_break()
    doc.add_heading('PARTE II: PREÂMBULO (Considerandos)', level=2)
    
    for tema in TEMAS_PREAMB:
        if tema not in PREAMB_POR_TEMA:
            continue
        
        considerandos = PREAMB_POR_TEMA[tema]
        if not considerandos:
            continue
        
        # Título do tema
        doc.add_heading(tema, level=3)
        
        for cons in considerandos:
            # Considerando
            doc.add_heading(f"Considerando {cons.get('numero', '')}", level=4)
            
            # EN
            p = doc.add_paragraph()
            p.add_run('EN: ').bold = True
            p.add_run(cons.get('regulamento', {}).get('texto', ''))
            
            # PT
            p = doc.add_paragraph()
            p.add_run('PT: ').bold = True
            p.add_run(cons.get('regulamento', {}).get('traducao', ''))
            
            doc.add_paragraph()  # espaço

# Salvar
filename = 'comparativo_reuniao_exemplo_preamb_v2.docx'
doc.save(filename)
print(f"\n✅ {filename} criado com sucesso!")
print(f"   - {len(ARTIGOS)} artigos")
if PREAMB_POR_TEMA:
    total_cons = sum(len(c) for c in PREAMB_POR_TEMA.values())
    print(f"   - {len(PREAMB_POR_TEMA)} temas ({total_cons} considerandos)")
