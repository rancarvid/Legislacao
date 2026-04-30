#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera documento Word detalhado com os 4 artigos totalmente novos
do Regulamento 2023/0447
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Define cor de fundo de célula."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level, color=None):
    """Adiciona heading com cor opcional."""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

# Dados dos artigos
ARTIGOS_NOVOS = [
    {
        "numero": 1,
        "titulo_en": "Subject matter",
        "titulo_pt": "Objeto",
        "tema": "Objeto e Escopo",
        "texto_en": """This Regulation lays down:
(a) minimum requirements for the welfare of dogs and cats bred or kept in establishments or placed on the market of the Union; and
(b) rules on the traceability of dogs and cats.""",
        "texto_pt": """O presente regulamento define os requisitos mínimos para:
a) O bem-estar dos cães e gatos criados ou detidos em estabelecimentos ou colocados no mercado da União; e
b) A rastreabilidade dos cães e gatos.""",
        "analise": """Este artigo é completamente novo no ordenamento jurídico português. Define o objeto central do Regulamento Europeu, que tem dois pilares:

1. BEM-ESTAR: estabelece requisitos mínimos para o bem-estar de cães e gatos em qualquer estabelecimento ou quando colocados no mercado europeu.

2. RASTREABILIDADE: institui obrigações de rastreamento de cães e gatos ao longo da cadeia de colocação no mercado.

A legislação portuguesa vigente (DL 276/2001, DL 82/2019, Lei 27/2016) e as propostas (@codigo, @rgbeac) não definem explicitamente estes dois eixos centrais como objeto de regulação integrada.""",
        "impacto": "Obrigação de alinhamento integral do regime nacional com estes dois objetivos.",
    },
    {
        "numero": 2,
        "titulo_en": "Material scope",
        "titulo_pt": "Âmbito de Aplicação Material",
        "tema": "Aplicabilidade Territorial e Material",
        "texto_en": """1. This Regulation applies to the breeding, keeping, tracing, placing on the market and entry into the Union of dogs and cats.
2. This Regulation does not apply to the breeding, keeping or placing on the market or entry into the Union of dogs or cats intended or used for scientific purposes or for clinical trials required for the marketing authorisation of veterinary medicinal products.""",
        "texto_pt": """1. O presente regulamento é aplicável à criação, à detenção, à rastreabilidade, à colocação no mercado e à entrada na União de cães e gatos.
2. O presente regulamento não é aplicável à criação, à detenção, à colocação no mercado ou à entrada na União de cães ou gatos destinados ou usados para fins científicos ou para ensaios clínicos exigidos para efeitos da autorização de introdução no mercado de medicamentos veterinários.""",
        "analise": """Define o âmbito material do Regulamento. Aplica-se a TODA a atividade relativa a cães e gatos: criação, detenção, rastreio e colocação no mercado.

EXCEÇÃO: cães e gatos para fins científicos ou ensaios clínicos de medicamentos veterinários (não se aplica o Regulamento).

A legislação portuguesa não define explicitamente um âmbito material tão abrangente e integrado. O DL 276/2001 e DL 82/2019 tratam aspetos específicos, mas não estabelecem um regime unitário de aplicação a todas as atividades mencionadas.""",
        "impacto": "Necessidade de integração regulatória - o Regulamento tem aplicação DIRETA a todas as atividades com cães/gatos, mesmo que não exista transposição nacional.",
    },
    {
        "numero": 3,
        "titulo_en": "Personal scope",
        "titulo_pt": "Âmbito de Aplicação Pessoal",
        "tema": "Sujeitos Obrigados",
        "texto_en": """1. Chapter II of this Regulation applies to all operators.
2. Chapter III of this Regulation applies to all natural and legal persons owning dogs or cats in the Union.
3. Chapter IV applies to all natural and legal persons who bring dogs or cats into the Union.
4. This Regulation does not apply to farmers offering refuge on their holding to free-roaming stray cats that are useful for pest control, where those farmers are not operators and do not place those cats on the market.""",
        "texto_pt": """1. O capítulo II é aplicável a todos os operadores.
2. O capítulo III é aplicável a todas as pessoas singulares e coletivas proprietárias de cães ou gatos na União.
3. O capítulo IV é aplicável a todas as pessoas singulares e coletivas que introduzam cães ou gatos na União.
4. O presente regulamento não é aplicável a agricultores que, na sua exploração, ofereçam refúgio a gatos errantes que sejam úteis para o controlo de pragas, desde que tais agricultores não sejam operadores e não coloquem esses gatos no mercado.""",
        "analise": """Define QUEM está obrigado pelo Regulamento:

- CAPÍTULO II (Operadores): criadores, vendedores, abrigos — pessoas que colocam cães/gatos no mercado
- CAPÍTULO III (Proprietários): todos os que possuem cães/gatos (responsabilidade do proprietário)
- CAPÍTULO IV (Importadores): quem traz cães/gatos para a UE

EXCEÇÃO: agricultores que acolhem gatos errantes para controlo de pragas (desde que não os comercializem).

A legislação portuguesa não diferencia de forma explícita entre operadores, proprietários e importadores, nem estabelece obrigações segmentadas por tipo de sujeito.""",
        "impacto": "Alargamento da esfera de obrigações para TODOS os proprietários de cães/gatos, não apenas criadores/vendedores.",
    },
    {
        "numero": 4,
        "titulo_en": "Definitions",
        "titulo_pt": "Definições",
        "tema": "Conceitos-Chave",
        "texto_en": """For the purposes of this Regulation, the following definitions apply:
1. 'dog' means an animal of the species Canis lupus familiaris;
2. 'cat' means an animal of the species Felis silvestris catus;
3. 'welfare of dogs and cats' means the physical and mental state of a dog or a cat, which receives appropriate nutrition, is kept in an appropriate environment, is in good health, displays appropriate behaviour, and has an overall a positive mental experience of life;
[... 36 mais definições, total de 39 conceitos]""",
        "texto_pt": """Para efeitos do presente regulamento, entende-se por:
1) «Cão», um animal da espécie Canis lupus familiaris;
2) «Gato», um animal da espécie Felis silvestris catus;
3) «Bem-estar dos cães e gatos», o estado físico e mental de um cão ou de um gato que recebe alimentação adequada, é detido num ambiente adequado, é mantido em boa saúde, exibe um comportamento adequado e, no geral, tem uma experiência mental positiva da vida;
[... 36 mais definições, total de 39 conceitos]""",
        "analise": """Define 39 conceitos fundamentais que estruturam toda a regulação. EXEMPLOS CHAVE:

- "Bem-estar": definição holística (nutrição, ambiente, saúde, comportamento, experiência mental positiva)
- "Operador": pessoa singular/coletiva que coloca cães/gatos no mercado e é responsável por um estabelecimento
- "Estabelecimento": criação, venda, abrigo, lar de acolhimento
- "Lar de acolhimento" (foster home): conceito novo — alojamento privado responsabilizado por operador
- "Amarração" (tethering): ato de prender um cão/gato a um ponto fixo
- "Mutilação": intervenção cirúrgica por razões não terapêuticas (inclui orelhas, cauda, garras)
- "Propriedade responsável": conjunto de comportamentos do proprietário coerentes com bem-estar

NOVIDADE: a legislação portuguesa vigente NÃO possui definições sistemáticas destes 39 conceitos. Cada diploma histórico usa terminologia diferente ou deixa conceitos implícitos.""",
        "impacto": "Necessidade de harmonizar terminologia portuguesa com as 39 definições do Regulamento. Impacto material: conceitos como 'bem-estar', 'operador', 'mutilação' e 'propriedade responsável' mudam de sentido jurídico.",
    },
]

def main():
    print("\n" + "=" * 80)
    print("GERANDO DOCUMENTO WORD - ARTIGOS TOTALMENTE NOVOS")
    print("=" * 80)

    doc = Document()

    # Capa
    title = doc.add_heading('Disposições Totalmente Novas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(102, 126, 234)

    subtitle = doc.add_heading('Regulamento (UE) 2023/0447 sobre bem-estar de cães e gatos', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        'Análise de artigos sem correspondência na legislação portuguesa vigente\n'
        '(@codigo, @rgbeac, DL 276/2001, DL 82/2019, Lei 27/2016)',
        style='Normal'
    )

    doc.add_paragraph()

    # Índice
    doc.add_heading('Índice de Artigos Novos', level=2)
    for art in ARTIGOS_NOVOS:
        doc.add_paragraph(f"Art. {art['numero']}: {art['titulo_pt']}", style='List Bullet')

    doc.add_page_break()

    # Cada artigo
    for art in ARTIGOS_NOVOS:
        # Cabeçalho do artigo
        heading = doc.add_heading(f"Artigo {art['numero']}: {art['titulo_pt']}", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(102, 126, 234)

        # Info do artigo
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_table.rows[0].cells[0].text = "Tema"
        info_table.rows[0].cells[1].text = art['tema']
        set_cell_background(info_table.rows[0].cells[0], "E8E8FF")

        info_table.rows[1].cells[0].text = "Título (EN)"
        info_table.rows[1].cells[1].text = art['titulo_en']
        set_cell_background(info_table.rows[1].cells[0], "E8E8FF")

        info_table.rows[2].cells[0].text = "Classificação"
        info_table.rows[2].cells[1].text = "TOTALMENTE NOVO — Sem correspondência em legislação portuguesa"
        set_cell_background(info_table.rows[2].cells[0], "E8E8FF")

        doc.add_paragraph()

        # Texto EN
        doc.add_heading("Texto Original (Inglês)", level=3)
        p = doc.add_paragraph(art['texto_en'])
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()

        # Texto PT
        doc.add_heading("Tradução (Português)", level=3)
        doc.add_paragraph(art['texto_pt'])

        doc.add_paragraph()

        # Análise
        doc.add_heading("Análise da Disposição", level=3)
        doc.add_paragraph(art['analise'])

        doc.add_paragraph()

        # Impacto
        doc.add_heading("Impacto para Portugal", level=3)
        impact_p = doc.add_paragraph(art['impacto'])
        for run in impact_p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(200, 80, 80)

        doc.add_page_break()

    # Conclusões
    doc.add_heading("Síntese Conclusiva", level=1)
    doc.add_paragraph(
        "Os 4 artigos iniciais do Regulamento 2023/0447 são TOTALMENTE NOVOS no ordenamento português. "
        "Estabelecem a moldura conceptual, material e pessoal do regulamento, e não encontram correspondência "
        "explícita em nenhum documento legislativo ou proposta portuguesa."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Consequências imediatas:",
        style='Heading 3'
    )

    conclusions = [
        "Necessidade de transposição integral e coerente dos 39 conceitos definidos no Art.º 4.º",
        "Alargamento substancial do âmbito pessoal: obrigações para TODOS os proprietários de cães/gatos, não apenas criadores/vendedores",
        "Integração de 5 atividades-chave numa única regulação: criação, detenção, rastreio, colocação no mercado, importação",
        "Harmonização terminológica urgente entre legislação portuguesa e conceitos do Regulamento",
    ]

    for i, conc in enumerate(conclusions, 1):
        doc.add_paragraph(conc, style='List Number')

    # Salva
    output_path = '/home/user/Legislacao/artigos_totalmente_novos_analise_completa.docx'
    doc.save(output_path)

    print(f"\n✓ Documento gravado: {output_path}")
    print(f"  Páginas: {len(doc.paragraphs) + len(doc.tables)}")

if __name__ == "__main__":
    main()
