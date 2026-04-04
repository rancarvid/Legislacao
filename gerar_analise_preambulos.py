#!/usr/bin/env python3
"""
gerar_analise_preambulos.py

Gera análise comparativa dos preâmbulos do @regulamento, @codigo e @rgbeac,
cruzada com os temas recorrentes nas opiniões (@opiniao) e a realidade portuguesa.

Uso:
  python3 gerar_analise_preambulos.py           → gera HTML
  python3 gerar_analise_preambulos.py --word    → gera Word (.docx)
  python3 gerar_analise_preambulos.py --excel   → gera Excel (.xlsx)
  python3 gerar_analise_preambulos.py --all     → gera HTML, Word e Excel
"""

import sys
import argparse
from datetime import date

DATA_PRODUCAO = date.today().strftime("%d de %B de %Y").replace(
    "January", "janeiro").replace("February", "fevereiro").replace(
    "March", "março").replace("April", "abril").replace(
    "May", "maio").replace("June", "junho").replace(
    "July", "julho").replace("August", "agosto").replace(
    "September", "setembro").replace("October", "outubro").replace(
    "November", "novembro").replace("December", "dezembro")

# ─────────────────────────────────────────────────────────────────────────────
# CONTEÚDO ANALÍTICO
# ─────────────────────────────────────────────────────────────────────────────

SEC1_TITULO = "Espírito legislativo do Regulamento Europeu 2023/0447"
SEC1_SUBTITULO = "Bem-estar de cães e gatos e sua rastreabilidade"

SEC1_TEXTO = """
<p>O Regulamento (UE) 2023/0447 é o <strong>primeiro instrumento legislativo da União Europeia
especificamente dedicado a animais de companhia</strong>. A sua base legal dual — artigo 43.º, n.º 2
(política agrícola comum) e artigo 114.º (mercado interno) do TFUE — não é acidental: revela que o
legislador europeu entendeu que a regulação dos animais de companhia é simultaneamente uma questão
de <em>bem-estar animal</em> e de <em>harmonização do mercado interno</em>.</p>

<p>Os 109 considerandos do preâmbulo constroem uma narrativa em três atos. Primeiro, o
<strong>reconhecimento filosófico</strong>: os animais são seres sencientes, capaz de sentir dor e
sofrimento — reconhecimento já inscrito no artigo 13.º do TFUE e no Protocolo n.º 33 do Tratado de
Amesterdão. Segundo, o <strong>diagnóstico dos problemas</strong>: a fragmentação normativa entre
Estados-Membros cria distorções no mercado interno, facilita o comércio ilegal e perpetua condições
de criação que geram sofrimento animal. Terceiro, a <strong>resposta normativa</strong>: um conjunto
articulado de instrumentos destinados a garantir rastreabilidade, elevar os padrões mínimos de bem-estar
e equiparar as condições de entrada de animais provenientes de países terceiros.</p>

<p>O Regulamento diagnostica explicitamente o fenómeno dos <em>puppy mills</em> — explorações de
criação comercial intensiva com recurso a práticas lesivas do bem-estar — como um dos problemas
centrais a resolver. A rastreabilidade (microchip + base de dados interoperável entre todos os EM)
é apresentada como o instrumento transversal que permite identificar a origem de cada animal,
responsabilizar os criadores e detetar redes de tráfico ilegal.</p>

<p>O espírito do Regulamento é, em síntese, <strong>harmonização ascendente</strong>: elevar todos os
Estados-Membros ao nível dos que já possuem legislação mais exigente, eliminando a vantagem
competitiva de quem produz em condições de menor bem-estar. A escolha do instrumento «regulamento»
(de aplicação direta, sem necessidade de transposição) reforça esta intenção: não deixar margem para
implementações fragmentadas ou diluídas a nível nacional.</p>
"""

SEC1_OBJETIVOS = [
    ("Combater o comércio ilegal de animais de companhia",
     "Registo obrigatório antes de qualquer transação; base de dados interoperável entre EM; verificação de plataformas online"),
    ("Garantir bem-estar mínimo na criação comercial",
     "Limiar de ativação (3 cadelas / 2 ninhadas/ano); normas de alojamento; visitas veterinárias; formação de criadores"),
    ("Eliminar práticas lesivas",
     "Proibição de corte de orelhas e de cauda; proibição de coleiras elétricas; proibição de acorrentamento >1 hora"),
    ("Controlar a reprodução",
     "Limite de ninhadas; idade mínima de reprodução; limitação de intervenções de inseminação artificial"),
    ("Equiparar importações de países terceiros",
     "Equivalência obrigatória dos standards de bem-estar para animais importados; registo antes da entrada no EM de destino"),
    ("Regular o mercado digital",
     "Obrigações para plataformas online de verificação do registo do animal e do operador antes de publicar anúncio"),
    ("Gerir populações de gatos errantes",
     "Reconhecimento dos programas de captura-esterilização-devolução (CED) como instrumento de gestão"),
    ("Elevar a consciência pública",
     "Promoção da educação e da adoção em detrimento da compra"),
]

# ─────────────────────────────────────────────────────────────────────────────

SEC2_TITULO = "Espírito legislativo do @codigo — Código do Animal (DL 214/2013)"
SEC2_SUBTITULO = "Proposta de consolidação — perspetiva operacional e de controlo"

SEC2_TEXTO = """
<p>O preâmbulo do <em>Código do Animal</em> (DL n.º 214/2013), produzido em 2013, revela um espírito
<strong>pragmático e administrativo</strong>, centrado no controlo das populações de animais de companhia
e na prevenção do abandono, mais do que num discurso filosófico sobre bem-estar ou senciência.</p>

<p>O problema central identificado é a <strong>criação não planeada e indiscriminada</strong> como
causa do abandono: o excesso de animais gerados sem procura correspondente conduz inevitavelmente à
negligência e ao abandono. A solução preconizada é fundamentalmente <em>instrumental</em>: a identificação
eletrónica (microchip) é apresentada como o mecanismo de responsabilização do detentor, não como
instrumento de bem-estar em si.</p>

<p>O @codigo foi inovador para a sua época em dois aspetos: a <strong>consolidação normativa</strong>
(reunir num diploma único todas as regras sobre reprodução, criação, detenção, maneiro e comércio) e
o <strong>programa CED</strong> (captura, esterilização e devolução) para as colónias de gatos, que
antecipou práticas hoje amplamente reconhecidas.</p>

<p>O que está <strong>ausente</strong> no @codigo é revelador do salto qualitativo ocorrido na década
seguinte: não há referência à senciência animal, ao valor intrínseco dos animais, às cinco liberdades
WOAH, a critérios de bem-estar comportamental, nem a proibições de práticas aversivas. O @codigo é,
na sua essência, um texto de <em>gestão administrativa de populações animais</em>.</p>

<p>Confrontado com o Regulamento europeu, o @codigo surge como um instrumento de primeira geração:
define quem pode ter animais e em que condições, mas não questiona o <em>como</em> se trata os
animais nem o <em>porquê</em> de os proteger para além da utilidade humana.</p>
"""

SEC2_COMPARACAO = [
    ("Fundamento filosófico", "Prevenção do abandono; controlo das populações", "Senciência animal; valor intrínseco; dignidade"),
    ("Instrumento central", "Identificação eletrónica (responsabilização do detentor)", "Rastreabilidade + normas de bem-estar (dupla função)"),
    ("Reprodução", "Evitar criação não planeada", "Limites quantitativos + condições mínimas + registo de criadores"),
    ("Práticas lesivas", "Não abordadas", "Proibições explícitas (coleiras, corte, acorrentamento)"),
    ("Animais errantes", "Programa CED para gatos (inovador à época)", "CED reconhecido + obrigações de registo de errantes"),
    ("Venda", "Incentivo à compra em locais autorizados", "Proibição de venda em pet shops; plataformas online reguladas"),
    ("Importações", "Não abordadas", "Equivalência obrigatória de standards"),
    ("Natureza do instrumento", "Consolidação operacional", "Harmonização ascendente do mercado interno"),
]

# ─────────────────────────────────────────────────────────────────────────────

SEC3_TITULO = "Espírito legislativo do @rgbeac — Regime Geral do Bem-Estar dos Animais de Companhia"
SEC3_SUBTITULO = "Proposta de junho de 2025 — convergência europeia com identidade nacional"

SEC3_TEXTO = """
<p>O preâmbulo do RGBEAC representa um <strong>salto qualitativo significativo</strong> relativamente ao
@codigo. A diferença mais imediata é filosófica: o RGBEAC parte do reconhecimento expresso da
<em>senciência animal</em> — já consagrada na Lei n.º 8/2017, de 3 de março — e do <em>valor
intrínseco</em> dos animais, independente da sua utilidade para o ser humano.</p>

<p>O RGBEAC inscreve-se numa tripla moldura conceptual: as <strong>cinco liberdades</strong> reconhecidas
pela WOAH (Organização Mundial da Saúde Animal), os <strong>doze critérios</strong> de bem-estar do
Projeto Welfare Quality (parceria europeia de investigação), e a <strong>evolução científica</strong>
sobre comportamento animal e efeitos dos métodos aversivos de treino.</p>

<p>A consolidação normativa — objetivo partilhado com o @codigo — é aqui acompanhada de uma <strong>reforma
substantiva</strong>: o RGBEAC não se limita a reunir diplomas existentes, altera-os e melhora-os.
As principais inovações incluem a Lista Nacional de Animais de Companhia (lista positiva de espécies),
a proibição de coleiras elétricas (em linha com ~12 países europeus), a proibição de acorrentamento
permanente (em linha com Alemanha, Países Baixos, Suécia, Áustria e Chipre), e a proibição de venda
de cães, gatos e furões em lojas (em linha com 7 EM).</p>

<p>O RGBEAC dá também especial atenção à <strong>fiscalização</strong>: estabelece um regime bicéfalo
(estabelecimentos abertos ao público vs. propriedade privada), reforça os poderes da DGAV e eleva a
moldura das coimas. A reforma do regime dos animais perigosos abandona a classificação por raça —
demonstrada como cientificamente infundada — em favor de uma abordagem baseada no comportamento
individual e nas condições de detenção.</p>

<p>Em síntese, o RGBEAC é o texto que alinha Portugal com a <em>vanguarda europeia</em> em matéria
de bem-estar de animais de companhia, mantendo instrumentos nacionais consolidados (SIAC, centros de
bem-estar animal, associações zoófilas) e introduzindo o enquadramento filosófico e normativo necessário
para receber e implementar o Regulamento europeu 2023/0447.</p>
"""

SEC3_INOVACOES = [
    ("Lista Nacional de Animais de Companhia", "Lista positiva de espécies permitidas como animais de companhia", "Bélgica, Países Baixos, Finlândia, França, Croácia, Chipre, Luxemburgo, Lituânia, Eslovénia"),
    ("Proibição de coleiras elétricas", "Proibição de uso e comercialização de mecanismos de estímulo negativo", "Alemanha, Dinamarca, Áustria, Finlândia, Países Baixos, Suíça, Suécia, Inglaterra, França, Espanha, Bélgica"),
    ("Proibição de acorrentamento permanente", "Proibição de manter animais permanentemente acorrentados", "Alemanha, Países Baixos, Suécia, Áustria, Chipre"),
    ("Proibição de venda em lojas", "Proibição de venda de cães, gatos e furões em pet shops", "7 Estados-Membros da UE"),
    ("Centros de bem-estar animal", "Substituição dos centros de recolha oficial por centros com função educativa e de promoção do bem-estar", "Boas práticas nacionais já existentes"),
    ("Reforma do regime de animais perigosos", "Abandono da classificação por raça; foco no comportamento e condições de detenção", "Evidência científica consolidada"),
    ("Sistema de Informação de Animais de Companhia (SIAC)", "Extensão do registo obrigatório a todos os animais de companhia (não apenas cães e gatos)", "Generalização do sistema existente"),
    ("Métodos aversivos de treino", "Proibição de uso de métodos aversivos; promoção do reforço positivo", "ESVE, BVA, AFFVAC, AVA, NZVA, The Kennel Club"),
]

# ─────────────────────────────────────────────────────────────────────────────

SEC4_TITULO = "Temas recorrentes nas opiniões (@opiniao)"
SEC4_SUBTITULO = "Análise de 18 pareceres de organizações europeias sobre o Regulamento 2023/0447"

SEC4_INTRO = """
<p>Foram analisados <strong>52 pareceres</strong> produzidos por organizações de 17 países europeus
e internacionais sobre versões anteriores do Regulamento 2023/0447, agrupados em 7 categorias:
organizações de bem-estar animal, veterinárias, kennel clubs e felinologia, caçadores e uso funcional,
indústria e comércio, rastreabilidade e I&R, e autoridades públicas. Embora algumas posições
específicas possam já não se aplicar à versão final do Regulamento, os temas de fundo permanecem
plenamente relevantes para a análise de impacto e implementação.</p>
<p>Use o filtro abaixo para visualizar por grupo de organização.</p>
"""

# Themes: 10 columns
# Position codes: "Favor" / "Contra" / "Reforço" / "Derrogação" / "" (N/A)
OPINIAO_TEMAS = [
    "Limiar\nArt. 4", "Proibição\nvenda lojas", "Castração\ncirúrgica", "Coleiras\nelétricas",
    "Cães\nde caça", "Rastreabi-\nlidade", "Braquice-\nfálicos", "Acorren-\ntamento",
    "Animais\nerrantes", "Venda\nonline"
]

# (Organização, Tipo/descrição, Grupo, [posições por tema])
# Grupos: "Bem-estar animal" | "Veterinária" | "Kennel clubs e felinologia"
#         | "Caçadores e uso funcional" | "Indústria e comércio" | "Rastreabilidade e I&R"
#         | "Autoridades públicas"
OPINIAO_DADOS = [
    # ── GRUPO: Bem-estar animal ────────────────────────────────────────────
    ("Eurogroup for Animals", "ONG pan-europeia (BE)", "Bem-estar animal",
     ["Contra", "Favor", "Favor", "Contra", "", "Favor", "Reforço", "Contra", "Reforço", "Reforço"]),
    ("LAV Onlus", "ONG bem-estar (IT)", "Bem-estar animal",
     ["Contra", "Favor", "", "Contra", "", "Favor", "Reforço", "Contra", "Reforço", ""]),
    ("AWI / Ecuavet", "Clínica veterinária (EU)", "Bem-estar animal",
     ["", "", "", "", "", "Reforço", "", "Contra", "", ""]),
    ("AnimaNaturalis", "ONG bem-estar", "Bem-estar animal",
     ["", "", "", "", "", "Reforço", "", "", "", "Reforço"]),
    ("VETO", "ONG bem-estar EU", "Bem-estar animal",
     ["", "", "", "", "Contra", "Favor", "", "", "Reforço", ""]),
    ("SEY Animal Welfare Finland", "Proteção animal (FI)", "Bem-estar animal",
     ["", "Contra", "", "Contra", "", "Favor", "", "", "Reforço", ""]),
    ("Norwegian Society for Protection", "Proteção animal (NO)", "Bem-estar animal",
     ["Derrogação", "", "", "", "", "Favor", "Reforço", "", "", ""]),
    ("Tierschutz-LASA", "Proteção animal (DE)", "Bem-estar animal",
     ["", "", "", "", "Contra", "", "", "", "Reforço", ""]),
    ("EUCED", "Cooperação económica EU", "Bem-estar animal",
     ["", "", "", "", "", "Reforço", "", "", "Favor", "Reforço"]),
    ("Animal Protection Denmark", "Proteção animal (DK)", "Bem-estar animal",
     ["Contra", "Favor", "Reforço", "", "", "Favor", "Reforço", "Favor", "Reforço", "Favor"]),
    ("Djurskyddet Sverige", "Proteção animal (SE)", "Bem-estar animal",
     ["Contra", "Favor", "Reforço", "Contra", "Reforço", "Favor", "Contra", "Contra", "Reforço", "Favor"]),
    ("Finnish Centre for Animal Welfare", "Centro bem-estar (FI)", "Bem-estar animal",
     ["Reforço", "", "", "", "Reforço", "Favor", "Reforço", "Favor", "", ""]),
    ("Dutch Assoc. Protection Animals", "Proteção animal (NL)", "Bem-estar animal",
     ["Contra", "", "", "", "Derrogação", "Favor", "Reforço", "", "", "Favor"]),
    ("EU Dog & Cat Alliance", "Coligação ONG (EU)", "Bem-estar animal",
     ["Contra", "Favor", "", "", "", "Favor", "Reforço", "", "", "Reforço"]),
    ("Brigitte Bardot Foundation", "ONG bem-estar (FR)", "Bem-estar animal",
     ["Contra", "Favor", "Reforço", "Contra", "Reforço", "Favor", "Contra", "", "Reforço", "Favor"]),
    ("PETA Deutschland", "ONG bem-estar (DE)", "Bem-estar animal",
     ["Derrogação", "Favor", "Favor", "Reforço", "Contra", "Favor", "Reforço", "Reforço", "Favor", "Favor"]),
    ("NOAH – for dyrs rettigheter", "ONG bem-estar (NO)", "Bem-estar animal",
     ["Reforço", "Favor", "Reforço", "Reforço", "Derrogação", "Favor", "Reforço", "Favor", "Favor", "Favor"]),
    ("SOS Galgos", "ONG galgos (ES)", "Bem-estar animal",
     ["Contra", "Favor", "Favor", "", "Contra", "Favor", "Contra", "Contra", "Favor", ""]),
    ("FENPCA", "Fed. proteção cães (ES)", "Bem-estar animal",
     ["Contra", "Favor", "Reforço", "", "Contra", "Reforço", "Reforço", "Reforço", "Reforço", "Reforço"]),
    ("LAU KATU ELKARTEA", "Assoc. proteção animal (ES)", "Bem-estar animal",
     ["Derrogação", "Reforço", "Reforço", "", "Derrogação", "Reforço", "Reforço", "", "Reforço", "Favor"]),
    ("Asociación HAIEKIN", "Assoc. proteção animal (ES)", "Bem-estar animal",
     ["Derrogação", "Reforço", "Reforço", "", "Derrogação", "Reforço", "Reforço", "", "Reforço", "Favor"]),
    ("APARIOJA (La Rioja)", "Assoc. proteção animal (ES)", "Bem-estar animal",
     ["Derrogação", "Reforço", "Reforço", "", "Derrogação", "Reforço", "Reforço", "", "Reforço", "Favor"]),
    ("MIS AMIGAS LAS PALOMAS", "ONG proteção animal (ES)", "Bem-estar animal",
     ["Derrogação", "Favor", "Favor", "", "Contra", "Favor", "Reforço", "Reforço", "Favor", "Favor"]),
    ("INTERcids", "Operadores jurídicos (EU)", "Bem-estar animal",
     ["", "", "", "", "", "", "", "", "", ""]),

    # ── GRUPO: Veterinária ─────────────────────────────────────────────────
    ("FVE (Fed. Veterinários EU)", "Federação veter. europeia", "Veterinária",
     ["Contra", "Favor", "Favor", "", "", "Favor", "Reforço", "", "Favor", "Favor"]),
    ("WSAVA", "Assoc. mundial vet. peq. animais", "Veterinária",
     ["Reforço", "", "Favor", "Contra", "", "Favor", "Reforço", "Contra", "Favor", "Favor"]),
    ("Danish Veterinary Assoc.", "Assoc. veterinária (DK)", "Veterinária",
     ["Contra", "Favor", "Reforço", "", "", "Favor", "Reforço", "", "", "Favor"]),
    ("Federal Chamber Vet. (BTK 🇩🇪)", "Câmara federal veter. (DE)", "Veterinária",
     ["Reforço", "", "", "", "", "Favor", "", "", "", ""]),
    ("Norwegian Vet. Nurses Assoc.", "Enfermeiras veterinárias (NO)", "Veterinária",
     ["", "", "", "", "", "Favor", "", "", "", ""]),
    ("Manteca / UAB", "Académico med. veter.", "Veterinária",
     ["", "", "Contra", "", "", "", "", "", "", ""]),
    ("UAB (Comportamento animal)", "Académico comportamento", "Veterinária",
     ["", "", "Contra", "", "", "", "", "", "", ""]),
    ("VIRBAC", "Empresa farmacêutica veter.", "Veterinária",
     ["", "", "Contra", "", "", "", "", "", "", ""]),

    # ── GRUPO: Kennel clubs e felinologia ──────────────────────────────────
    ("Koninklijke Hondenbescherming", "Proteção cães (NL)", "Kennel clubs e felinologia",
     ["", "Contra", "", "", "", "Favor", "", "", "", ""]),
    ("The Dutch Kennel Club", "Kennel club (NL)", "Kennel clubs e felinologia",
     ["Contra", "", "", "", "", "Favor", "Reforço", "", "", "Favor"]),
    ("Norwegian Kennel Club", "Kennel club (NO)", "Kennel clubs e felinologia",
     ["Favor", "Favor", "Derrogação", "Derrogação", "Derrogação", "Favor", "Reforço", "Favor", "Favor", ""]),
    ("Lithuanian Felinology (Bubaste)", "Assoc. felinológica (LT)", "Kennel clubs e felinologia",
     ["", "", "", "", "", "", "", "", "", ""]),

    # ── GRUPO: Caçadores e uso funcional ───────────────────────────────────
    ("FACE (Fed. Caçadores EU)", "Fed. caça e conservação EU", "Caçadores e uso funcional",
     ["Derrogação", "", "", "Derrogação", "Derrogação", "", "", "", "", ""]),
    ("Féd. Nationale Chasseurs (FR)", "Fed. caçadores (FR)", "Caçadores e uso funcional",
     ["Derrogação", "", "", "Derrogação", "Derrogação", "Derrogação", "", "", "", ""]),
    ("Real Fed. Española de Caza", "Fed. caça (ES)", "Caçadores e uso funcional",
     ["Derrogação", "", "", "", "Favor", "", "", "Contra", "", ""]),

    # ── GRUPO: Indústria e comércio ────────────────────────────────────────
    ("FEDIAF", "Ind. alimentos animais (EU)", "Indústria e comércio",
     ["", "", "", "", "", "", "", "", "", ""]),
    ("MSD Animal Health", "Empresa farmacêutica veter.", "Indústria e comércio",
     ["", "", "", "", "", "Reforço", "", "", "", ""]),
    ("AnimalhealthEurope", "Ind. medicamentos animais", "Indústria e comércio",
     ["Reforço", "Contra", "", "", "", "Favor", "Reforço", "", "", "Favor"]),
    ("Electronic Collar Manufacturers", "Fabr. coleiras eléctricas", "Indústria e comércio",
     ["", "", "", "Contra", "Favor", "", "", "", "", ""]),
    ("Assoc. Responsible Dog Owners", "Assoc. proprietários (UK)", "Indústria e comércio",
     ["", "", "", "Derrogação", "Derrogação", "", "", "", "", ""]),
    ("Classifieds Marketplaces Europe", "Plataformas classificados", "Indústria e comércio",
     ["", "", "", "", "", "", "", "", "", "Derrogação"]),
    ("European Pet Organization (EPO)", "Org. setor animais companhia", "Indústria e comércio",
     ["", "", "", "", "", "", "", "", "", ""]),
    ("World Dog Alliance", "Org. internacional", "Indústria e comércio",
     ["", "", "", "", "", "Favor", "", "", "", "Favor"]),

    # ── GRUPO: Rastreabilidade e I&R ───────────────────────────────────────
    ("EuroPetNet", "Rede bases dados I&R", "Rastreabilidade e I&R",
     ["", "", "", "", "", "Favor", "", "", "", "Favor"]),
    ("K&R Network", "Rede I&R (DE)", "Rastreabilidade e I&R",
     ["Contra", "", "", "", "", "Favor", "", "", "", "Favor"]),
    ("TASSO e.V.", "Maior registo livre (DE)", "Rastreabilidade e I&R",
     ["Contra", "", "", "", "", "Favor", "Reforço", "", "", "Favor"]),

    # ── GRUPO: Autoridades públicas ────────────────────────────────────────
    ("Bundesministerium (🇩🇪)", "Ministério Alimentação e Agric. (DE)", "Autoridades públicas",
     ["Reforço", "", "Favor", "", "", "Favor", "Reforço", "", "", ""]),
    ("Czech National Authority", "Autoridade nacional (CZ)", "Autoridades públicas",
     ["Contra", "Favor", "Reforço", "", "Reforço", "Favor", "Reforço", "Reforço", "", "Favor"]),
    ("Norwegian Food Safety Auth.", "Autoridade alimentar (NO)", "Autoridades públicas",
     ["Derrogação", "Favor", "Favor", "Reforço", "Derrogação", "Favor", "Reforço", "Reforço", "Favor", "Favor"]),
    ("Fundación Vet+i (🇪🇸)", "Plataforma tecn. san. animal (ES)", "Autoridades públicas",
     ["", "", "", "", "", "", "", "", "", ""]),
]

SEC4_ANALISE = """
<p>Da análise cruzada dos <strong>52 pareceres</strong>, organizados em 7 grupos, emergem
<strong>quatro eixos de tensão</strong> fundamentais:</p>
<ul>
  <li><strong>Consenso alargado em rastreabilidade:</strong> A rastreabilidade (identificação
  eletrónica + base de dados interoperável) é o único tema com aprovação quase unânime — todas as
  organizações de bem-estar, veterinárias, kennel clubs e autoridades públicas apoiam. Mesmo a
  indústria e os caçadores não a contestam. O debate limita-se ao <em>âmbito</em>: se deve abranger
  todos os cães e gatos (posição das ONG) ou apenas os transacionados comercialmente.</li>
  <li><strong>Bem-estar vs. uso funcional:</strong> A divisão mais nítida opõe as organizações de
  bem-estar (que querem normas universais, incluindo cães de caça) às federações de caçadores (que
  defendem derrogações). FACE, FNC e Real Fed. Española de Caza rejeitam a aplicação do limiar Art. 4
  à criação não comercial de cães de caça. Em sentido contrário, SOS Galgos e FENPCA documentam que
  os cães de caça representam 70% dos abandonamentos em Espanha — argumento com relevância directa
  para Portugal.</li>
  <li><strong>Proibições vs. regulação:</strong> Coleiras elétricas, venda em pet shops e
  acorrentamento dividem o campo: ONG e autoridades querem proibição total; a indústria
  (Electronic Collar Manufacturers, ARDO) e parte dos kennel clubs propõem exceções e regulação.
  Significativamente, mesmo o Norwegian Kennel Club — que genericamente apoia o Regulamento —
  propõe uma derrogação limitada para uso em treino de cães de caça.</li>
  <li><strong>Debate científico sobre castração:</strong> Manteca/UAB, VIRBAC e UAB Comportamento
  questionam a castração cirúrgica como método-padrão, apontando efeitos adversos na saúde e
  comportamento documentados em 35 raças. A Norwegian Food Safety Authority e a FVE continuam a
  apoiar a castração como instrumento de controlo. Este debate tem implicações directas para as
  políticas de esterilização municipais em Portugal (54,4% dos cães registados já esterilizados).</li>
</ul>
"""

# ─────────────────────────────────────────────────────────────────────────────

SEC5_TITULO = "Questões polémicas e de difícil aplicação na realidade portuguesa"
SEC5_SUBTITULO = "Avaliação do impacto do Regulamento 2023/0447 em Portugal"

SEC5_INTRO = """
<p>A implementação do Regulamento europeu em Portugal levanta questões específicas, decorrentes de
características estruturais da realidade nacional: forte tradição cinegética, elevada taxa de animais
errantes, compliance irregular com as obrigações de registo, e recursos limitados da autoridade
competente (DGAV). A análise que se segue identifica os pontos de maior fricção normativa.</p>
"""

# (Tema, Artigo Reg., Polémica, Impacto, Nível: "alto"/"medio"/"baixo")
SEC5_TEMAS = [
    (
        "Limiar de ativação para criadores",
        "Art. 4.º",
        "O limiar de 3 cadelas / 2 ninhadas por ano captura proprietários de matilhas de caça sem atividade comercial. Em Portugal, a caça com cães (especialmente a caça maior e a caça de perdiz) é uma atividade cultural e económica significativa nas regiões do interior. A Portaria n.º 148/2016 já obriga ao registo das matilhas de caça maior, mas não estabelece limites de reprodução. A aplicação do Art. 4.º sem derrogação específica poderá criminalizar práticas legítimas.",
        "alto",
        "Cultura cinegética consolidada; Portaria 148/2016 já regula matilhas mas sem limites reprodutivos"
    ),
    (
        "Coleiras elétricas (Art. 15.º)",
        "Art. 15.º",
        "O Regulamento proíbe o uso de coleiras elétricas. O RGBEAC também as proíbe, o que representa convergência. No entanto, em Portugal, o uso de coleiras elétricas em contexto de caça e pastoreio (incluindo cães de gado a guardar rebanhos de zonas de lobo) é uma prática estabelecida. As federações de caçadores europeias (FACE, FNC) defendem uma abordagem de regulação em vez de proibição total. A fiscalização desta proibição em zonas rurais remotas é um desafio adicional.",
        "medio",
        "Uso estabelecido em contexto cinegético e pastoril; fiscalização difícil no interior"
    ),
    (
        "Acorrentamento permanente (Art. 14.º, §3)",
        "Art. 14.º, §3",
        "O Regulamento proíbe manter animais presos por mais de 1 hora. O RGBEAC proíbe o acorrentamento permanente. Em Portugal, o acorrentamento de cães — especialmente cães de guarda em propriedades rurais — é uma prática ainda comum, particularmente no interior do país. A OMV já identificou esta prática como fonte de sofrimento. A fiscalização em propriedade privada exige mandado judicial (no âmbito do RGBEAC), o que limita a capacidade de atuação das autoridades.",
        "alto",
        "Prática corrente em zonas rurais; fiscalização em propriedade privada condicionada por mandado judicial"
    ),
    (
        "Identificação e registo universal (Art. 17.º)",
        "Art. 17.º",
        "Em outubro de 2023, estavam registados no SIAC cerca de 1,07 milhão de cães e 629 mil gatos. O total de cães registados em 2023 era de 2,58 milhões, mas o número real de cães em Portugal é provavelmente superior. O compliance com a identificação obrigatória é irregular, especialmente em zonas rurais e para gatos. O Regulamento exige registo antes de qualquer transação, e o RGBEAC estende a obrigação a todos os animais de companhia. O desafio é garantir o cumprimento efetivo, não apenas legislar a obrigação.",
        "alto",
        "SIAC operacional mas com gaps de compliance; alargamento a todos os animais de companhia é desafio operacional significativo"
    ),
    (
        "Gestão de animais errantes",
        "Art. 13.º e ss.",
        "O Censo Nacional de Animais Errantes (2023) identificou 930 mil animais errantes em Portugal Continental (831 mil gatos, 101 mil cães). Portugal abandonou o abate de animais errantes como método de controlo (Lei 27/2016), mas a capacidade dos centros de bem-estar animal é insuficiente para absorver este volume. O Regulamento reconhece os programas CED (captura-esterilização-devolução) para gatos, mas não resolve o problema estrutural de financiamento e capacidade dos centros. O abandono de 42 mil animais/ano (119/dia) agrava permanentemente a situação.",
        "muito-alto",
        "930 mil errantes; 42 mil abandonos/ano; Lei 27/2016 proíbe abate mas recursos insuficientes para alternativa"
    ),
    (
        "Castração cirúrgica como política pública",
        "Art. 9.º (reprodução)",
        "Portugal incentiva ativamente a esterilização cirúrgica — 54,4% dos cães registados no SIAC em 2023 estavam esterilizados. Esta política é defendida como instrumento de controlo do abandono. Contudo, vários pareceres científicos (Manteca/UAB, VIRBAC, UAB comportamental) alertam para efeitos adversos da castração cirúrgica precoce na saúde e comportamento dos animais. A existência de alternativas médicas (implantes de deslorelina, reversíveis) levanta a questão de se a política pública portuguesa está alinhada com a melhor evidência científica disponível.",
        "medio",
        "54,4% dos cães registados já esterilizados; debate científico emergente sobre efeitos adversos da castração precoce"
    ),
    (
        "Raças braquicefálicas (Art. 6.º, §3)",
        "Art. 6.º, §3",
        "O Regulamento exige que a criação de animais com características morfológicas problemáticas (braquicefálicos — bulldogs, pugs, shih tzus, etc.) minimize as consequências negativas para o bem-estar. Organizações como Eurogroup for Animals e LAV Onlus consideram esta formulação insuficiente e defendem a proibição total da criação de raças extremas. Em Portugal, estas raças têm mercado estabelecido e criadores activos. A norma exigirá definição técnica de «características problemáticas» e capacidade de fiscalização especializada.",
        "medio",
        "Mercado de raças braquicefálicas estabelecido; formulação do Art. 6.º, §3 é vaga e de difícil execução"
    ),
    (
        "Venda e publicidade online (Art. 20.º)",
        "Art. 20.º",
        "O Regulamento impõe obrigações de verificação às plataformas online antes de publicar anúncios de venda de animais. Em Portugal, o comércio online e em redes sociais (Facebook Marketplace, OLX) de animais de companhia é um vetor significativo do tráfico ilegal e da criação irresponsável. A fiscalização efetiva de plataformas digitais — muitas delas com sede fora de Portugal — é um desafio regulatório complexo que exige cooperação entre a DGAV, a ANACOM e eventualmente autoridades de outros EM.",
        "alto",
        "Comércio ilegal online relevante; coordenação entre DGAV, ANACOM e plataformas digitais é um desafio novo"
    ),
    (
        "Capacidade institucional da DGAV",
        "Aplicação transversal",
        "A implementação do Regulamento exige um reforço significativo da DGAV em termos de recursos humanos, sistemas de informação e capacidade de fiscalização. A interoperabilidade da base de dados SIAC com os registos dos outros 26 EM é um requisito técnico exigente. O RGBEAC já reforça os poderes da DGAV, mas a capacidade efetiva de fiscalização — especialmente em zonas rurais — é uma preocupação transversal identificada por várias organizações (EUCED, OMV).",
        "muito-alto",
        "DGAV com recursos limitados; interoperabilidade SIAC com sistemas de outros EM é investimento técnico significativo"
    ),
]

# ─────────────────────────────────────────────────────────────────────────────

SEC6_TITULO = "Realidade portuguesa em dados"
SEC6_SUBTITULO = "Contexto estatístico e estrutural — animais de companhia em Portugal"

SEC6_TEXTO = """
<p>A análise do impacto do Regulamento europeu não pode prescindir de uma leitura objetiva da
realidade portuguesa. Os dados disponíveis — provenientes do SIAC (Sistema de Informação de
Animais de Companhia), do Censo Nacional de Animais Errantes (2023) e de publicações
científicas — revelam um panorama com progressos significativos nas últimas décadas, mas com
desafios estruturais que o Regulamento europeu e o RGBEAC procuram endereçar.</p>
"""

SEC6_ESTATISTICAS = [
    {
        "categoria": "Animais registados (SIAC, outubro 2023)",
        "items": [
            ("Cães registados", "~1.075.467", "Fonte: SIAC, outubro 2023"),
            ("Gatos registados", "~629.519", "Fonte: SIAC, outubro 2023"),
            ("Furões registados", "~1.907", "Fonte: SIAC, outubro 2023"),
            ("Total cães registados (2023, todos os registos)", "~2.581.870", "Fonte: estudo ScienceDirect 2025 com base em dados SIAC"),
            ("Cães esterilizados (% do total registado 2023)", "54,4%", "Fonte: estudo ScienceDirect 2025"),
        ]
    },
    {
        "categoria": "Animais errantes (Censo ICNF / Univ. Aveiro, 2023)",
        "items": [
            ("Total de animais errantes em Portugal Continental", "~930.000", "Maior stock de sempre desde a abolição do abate (Lei 27/2016)"),
            ("Gatos errantes", "~830.541", "82% da população de errantes são gatos"),
            ("Cães errantes", "~101.015", "18% da população de errantes são cães"),
            ("Animais recolhidos dos espaços públicos (2023)", "45.148", "Valor mais elevado dos últimos anos"),
            ("Acidentes rodoviários com animais errantes (2019–2022)", "4.640", "4.443 cães + 197 gatos; fonte: dados públicos"),
        ]
    },
    {
        "categoria": "Abandono (dados 2022)",
        "items": [
            ("Animais abandonados/recolhidos por ano", "~42.000", "Corresponde a 119 animais/dia"),
            ("Animais adotados/reencaminhados por ano", "~25.000", "Melhor resultado registado"),
            ("Animais eutanasiados por ano", "~2.000", "Apenas em casos de doença incurável ou agressividade grave"),
            ("Animais que permanecem nos centros", "~15.000", "Causa cumulativa de sobrelotação dos centros"),
            ("Crimes de abandono registados (2015–2019)", "6.711", "Pico em 2019 com 801 casos; fonte: estatísticas judiciais"),
        ]
    },
    {
        "categoria": "Contexto comparativo europeu",
        "items": [
            ("EM que proíbem venda em mercados", "12", "Portugal ainda não se incluía neste grupo antes do RGBEAC"),
            ("EM que proíbem venda em pet shops", "7", "Portugal proíbe com o RGBEAC"),
            ("EM que proíbem coleiras elétricas", "~12", "Portugal proíbe com o RGBEAC"),
            ("EM que proíbem acorrentamento permanente", "5+", "Portugal proíbe com o RGBEAC"),
        ]
    },
]

SEC6_REFLEXAO = """
<p>Os dados revelam uma tensão estrutural: Portugal avançou significativamente na legislação protetora
de animais (abolição do abate, Lei 8/2017, desenvolvimento do SIAC) mas enfrenta um paradoxo — o
crescimento do stock de errantes é em parte consequência das políticas de proteção que proibem o
abate sem que a capacidade alternativa (centros de bem-estar, esterilização, adoção) tenha crescido
proporcionalmente.</p>

<p>A chegada do Regulamento europeu a um país com estas características exige uma reflexão honesta:
as normas de bem-estar e rastreabilidade são acolhidas com o RGBEAC e o SIAC, mas a efetividade da
sua aplicação depende de investimento em capacidade institucional, infraestrutura de registos e
cultura de compliance — especialmente fora das áreas metropolitanas.</p>
"""

# ─────────────────────────────────────────────────────────────────────────────
# GERAÇÃO HTML
# ─────────────────────────────────────────────────────────────────────────────

def gerar_tabela_objetivos(dados):
    linhas = ""
    for obj, instr in dados:
        linhas += f"<tr><td><strong>{obj}</strong></td><td>{instr}</td></tr>\n"
    return f"""<table class="tabela-geral">
<thead><tr><th>Objetivo</th><th>Instrumento / Mecanismo previsto</th></tr></thead>
<tbody>{linhas}</tbody>
</table>"""

def gerar_tabela_comparacao(dados):
    linhas = ""
    for dim, codigo, reg in dados:
        linhas += f"<tr><td><strong>{dim}</strong></td><td class='t-codigo'>{codigo}</td><td class='t-reg'>{reg}</td></tr>\n"
    return f"""<table class="tabela-geral">
<thead><tr><th>Dimensão</th><th>@codigo (DL 214/2013)</th><th>Regulamento 2023/0447</th></tr></thead>
<tbody>{linhas}</tbody>
</table>"""

def gerar_tabela_inovacoes(dados):
    linhas = ""
    for inovacao, desc, refs in dados:
        linhas += f"<tr><td><strong>{inovacao}</strong></td><td>{desc}</td><td class='refs'>{refs}</td></tr>\n"
    return f"""<table class="tabela-geral">
<thead><tr><th>Inovação</th><th>Conteúdo</th><th>Países / Referências</th></tr></thead>
<tbody>{linhas}</tbody>
</table>"""

def gerar_tabela_opiniao(temas, dados):
    # Recolher grupos únicos em ordem de aparecimento
    grupos_vistos = []
    for org, tipo, grupo, posicoes in dados:
        if grupo not in grupos_vistos:
            grupos_vistos.append(grupo)

    cabecalhos = "".join(f"<th>{t}</th>" for t in temas)
    linhas = ""
    grupo_atual = None
    for org, tipo, grupo, posicoes in dados:
        if grupo != grupo_atual:
            grupo_atual = grupo
            ncols = len(temas) + 3  # org + grupo + tipo + temas
            linhas += f"<tr class='grupo-header' data-grupo='{grupo}'><td colspan='{ncols}'><strong>{grupo}</strong></td></tr>\n"
        cels = ""
        for p in posicoes:
            cls = ""
            if p == "Favor": cls = "pos-favor"
            elif p == "Contra": cls = "pos-contra"
            elif p == "Reforço": cls = "pos-reforco"
            elif p == "Derrogação": cls = "pos-derrog"
            cels += f"<td class='{cls}'>{p}</td>"
        linhas += f"<tr data-grupo='{grupo}'><td><strong>{org}</strong></td><td class='tipo-org'>{grupo}</td><td class='tipo-org'>{tipo}</td>{cels}</tr>\n"

    grupos_options = "".join(f'<option value="{g}">{g}</option>' for g in grupos_vistos)

    return f"""<div style="margin-bottom:12px;">
<label style="font-size:12px;font-weight:600;color:#555;">Filtrar por grupo: </label>
<select id="filtro-grupo" onchange="filtrarGrupo(this.value)" style="font-size:12px;padding:4px 8px;border:1px solid #ccc;border-radius:4px;">
<option value="">Todos os grupos</option>
{grupos_options}
</select>
</div>
<div class="tabela-scroll">
<table class="tabela-opiniao" id="tabela-opin">
<thead><tr>
<th>Organização</th><th>Grupo</th><th>Tipo</th>{cabecalhos}
</tr></thead>
<tbody>{linhas}</tbody>
</table>
</div>
<div class="legenda-opiniao">
<span class="pos-favor">Favor / apoia</span>
<span class="pos-contra">Contra / critica / propõe proibição</span>
<span class="pos-reforco">Propõe reforço da norma</span>
<span class="pos-derrog">Propõe derrogação / flexibilização</span>
<span>— Não abordado</span>
</div>
<script>
function filtrarGrupo(grupo) {{
  var rows = document.querySelectorAll('#tabela-opin tbody tr');
  rows.forEach(function(row) {{
    if (!grupo || row.getAttribute('data-grupo') === grupo) {{
      row.style.display = '';
    }} else {{
      row.style.display = 'none';
    }}
  }});
}}
</script>"""

def nivel_classe(nivel):
    return {"alto": "nivel-alto", "medio": "nivel-medio", "baixo": "nivel-baixo", "muito-alto": "nivel-muito-alto"}.get(nivel, "")

def nivel_label(nivel):
    return {"alto": "🔴 Alto", "medio": "🟡 Médio", "baixo": "🟢 Baixo", "muito-alto": "🔴🔴 Muito alto"}.get(nivel, nivel)

def gerar_tabela_polemicas(dados):
    linhas = ""
    for tema, artigo, polemica, nivel, nota in dados:
        cls = nivel_classe(nivel)
        lbl = nivel_label(nivel)
        linhas += f"""<tr>
<td><strong>{tema}</strong><br><small>{artigo}</small></td>
<td>{polemica}</td>
<td class="{cls} nivel-cell">{lbl}</td>
<td class="nota-pt">{nota}</td>
</tr>\n"""
    return f"""<table class="tabela-geral tabela-polemicas">
<thead><tr><th>Tema</th><th>Natureza da polémica</th><th>Impacto PT</th><th>Nota específica para Portugal</th></tr></thead>
<tbody>{linhas}</tbody>
</table>"""

def gerar_blocos_estatisticas(dados):
    html = ""
    for bloco in dados:
        html += f"<h3 class='stat-categoria'>{bloco['categoria']}</h3>"
        html += "<table class='tabela-stats'><thead><tr><th>Indicador</th><th>Valor</th><th>Nota</th></tr></thead><tbody>"
        for ind, val, nota in bloco['items']:
            html += f"<tr><td>{ind}</td><td class='stat-val'><strong>{val}</strong></td><td class='stat-nota'>{nota}</td></tr>"
        html += "</tbody></table>"
    return html

def gerar_html():
    tbl_obj = gerar_tabela_objetivos(SEC1_OBJETIVOS)
    tbl_comp = gerar_tabela_comparacao(SEC2_COMPARACAO)
    tbl_inov = gerar_tabela_inovacoes(SEC3_INOVACOES)
    tbl_opin = gerar_tabela_opiniao(OPINIAO_TEMAS, OPINIAO_DADOS)
    tbl_pol = gerar_tabela_polemicas(SEC5_TEMAS)
    blocos_stat = gerar_blocos_estatisticas(SEC6_ESTATISTICAS)

    html = f"""<!DOCTYPE html>
<html lang="pt">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Análise dos Preâmbulos e Opiniões — Regulamento UE 2023/0447</title>
<style>
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{ font-family: 'Segoe UI', Arial, sans-serif; font-size: 14px; color: #1a1a2e; background: #f4f6fb; }}
a {{ color: inherit; text-decoration: none; }}

/* LAYOUT */
#sidebar {{
  position: fixed; top: 0; left: 0; width: 260px; height: 100vh;
  background: #1a1a2e; color: #e0e0f0; overflow-y: auto;
  padding: 0; z-index: 100;
}}
#main {{ margin-left: 260px; padding: 0 0 60px 0; }}

/* SIDEBAR */
.sidebar-header {{
  background: #16213e; padding: 20px 18px 16px; border-bottom: 1px solid #2a2a4e;
}}
.sidebar-header h2 {{ font-size: 13px; color: #a0b0d0; font-weight: 600; line-height: 1.4; }}
.sidebar-header p {{ font-size: 11px; color: #6070a0; margin-top: 6px; }}
.sidebar-nav {{ padding: 10px 0; }}
.sidebar-nav a {{
  display: block; padding: 10px 18px; font-size: 12.5px; color: #b0c0e0;
  border-left: 3px solid transparent; transition: all 0.2s;
}}
.sidebar-nav a:hover {{ background: #16213e; color: #fff; border-left-color: #4a90d9; }}
.sidebar-nav a.sec1 {{ border-left-color: #0077b6; }}
.sidebar-nav a.sec2 {{ border-left-color: #e06c0a; }}
.sidebar-nav a.sec3 {{ border-left-color: #2a9d5c; }}
.sidebar-nav a.sec4 {{ border-left-color: #9b59b6; }}
.sidebar-nav a.sec5 {{ border-left-color: #e74c3c; }}
.sidebar-nav a.sec6 {{ border-left-color: #2980b9; }}
.nav-num {{ display: inline-block; width: 22px; font-weight: 700; opacity: 0.7; }}

/* SECTIONS */
.section {{ padding: 40px 48px 32px; border-bottom: 2px solid #e0e4ef; background: #fff; margin-bottom: 8px; }}
.section-header {{ margin-bottom: 28px; }}
.section-num {{ font-size: 11px; font-weight: 700; letter-spacing: 2px; text-transform: uppercase; opacity: 0.5; }}
.section-titulo {{ font-size: 22px; font-weight: 700; margin: 6px 0 4px; }}
.section-subtitulo {{ font-size: 14px; color: #555; font-style: italic; }}
.section-divider {{ height: 4px; width: 60px; border-radius: 2px; margin-top: 12px; }}

/* SECTION COLORS */
.s1 .section-titulo, .s1 .section-num {{ color: #0077b6; }}
.s1 .section-divider {{ background: #0077b6; }}
.s2 .section-titulo, .s2 .section-num {{ color: #e06c0a; }}
.s2 .section-divider {{ background: #e06c0a; }}
.s3 .section-titulo, .s3 .section-num {{ color: #2a9d5c; }}
.s3 .section-divider {{ background: #2a9d5c; }}
.s4 .section-titulo, .s4 .section-num {{ color: #9b59b6; }}
.s4 .section-divider {{ background: #9b59b6; }}
.s5 .section-titulo, .s5 .section-num {{ color: #e74c3c; }}
.s5 .section-divider {{ background: #e74c3c; }}
.s6 .section-titulo, .s6 .section-num {{ color: #2980b9; }}
.s6 .section-divider {{ background: #2980b9; }}

/* HERO */
#hero {{
  background: linear-gradient(135deg, #1a1a2e 0%, #16213e 60%, #0f3460 100%);
  color: #fff; padding: 52px 48px 40px; margin-left: 0;
}}
#hero h1 {{ font-size: 28px; font-weight: 700; line-height: 1.3; margin-bottom: 12px; }}
#hero p {{ color: #a0b8d8; font-size: 14px; margin-bottom: 6px; }}
.hero-tags {{ display: flex; gap: 10px; flex-wrap: wrap; margin-top: 20px; }}
.hero-tag {{
  background: rgba(255,255,255,0.12); padding: 4px 12px; border-radius: 20px;
  font-size: 12px; color: #c0d8f0;
}}

/* TEXTO */
.texto-analise {{ line-height: 1.75; color: #2a2a3e; }}
.texto-analise p {{ margin-bottom: 14px; }}
.texto-analise ul {{ margin: 12px 0 14px 20px; }}
.texto-analise li {{ margin-bottom: 8px; line-height: 1.6; }}

/* TABELAS */
.tabela-geral {{
  width: 100%; border-collapse: collapse; margin: 20px 0;
  font-size: 13px; background: #fff;
  box-shadow: 0 1px 4px rgba(0,0,0,0.08);
  border-radius: 6px; overflow: hidden;
}}
.tabela-geral thead tr {{ background: #1a1a2e; color: #e0e8f8; }}
.tabela-geral th {{ padding: 10px 14px; text-align: left; font-size: 12px; font-weight: 600; }}
.tabela-geral td {{ padding: 10px 14px; border-bottom: 1px solid #e8eaf6; vertical-align: top; }}
.tabela-geral tbody tr:last-child td {{ border-bottom: none; }}
.tabela-geral tbody tr:hover {{ background: #f0f4ff; }}
.t-codigo {{ background: #fff3e8; }}
.t-reg {{ background: #e8f4ff; }}
.refs {{ font-size: 11px; color: #666; font-style: italic; }}

/* TABELA OPINIÕES */
.tabela-scroll {{ overflow-x: auto; margin: 20px 0; }}
.tabela-opiniao {{
  width: 100%; border-collapse: collapse; font-size: 11.5px;
  min-width: 900px;
}}
.tabela-opiniao thead tr {{ background: #1a1a2e; color: #e0e8f8; }}
.tabela-opiniao th {{ padding: 8px 10px; text-align: center; font-size: 11px; white-space: pre-line; }}
.tabela-opiniao th:first-child, .tabela-opiniao th:nth-child(2) {{ text-align: left; }}
.tabela-opiniao td {{ padding: 7px 10px; border-bottom: 1px solid #e8eaf6; text-align: center; vertical-align: middle; }}
.tabela-opiniao td:first-child {{ text-align: left; white-space: nowrap; }}
.tabela-opiniao td:nth-child(2) {{ text-align: left; }}
.tabela-opiniao tbody tr:nth-child(even) {{ background: #f8f9fc; }}
.tipo-org {{ font-size: 11px; color: #666; font-style: italic; }}
.pos-favor {{ background: #e8f8e8 !important; color: #1a7a1a; font-weight: 600; font-size: 11px; }}
.pos-contra {{ background: #fde8e8 !important; color: #b31a1a; font-weight: 600; font-size: 11px; }}
.pos-reforco {{ background: #fff3cd !important; color: #7a5a00; font-size: 11px; }}
.pos-derrog {{ background: #ede8ff !important; color: #5a00aa; font-size: 11px; }}
.legenda-opiniao {{
  display: flex; gap: 18px; flex-wrap: wrap; margin-top: 12px;
  font-size: 11px; padding: 10px 0;
}}
.legenda-opiniao span {{ padding: 3px 10px; border-radius: 4px; }}
.legenda-opiniao .pos-favor {{ background: #e8f8e8; color: #1a7a1a; }}
.legenda-opiniao .pos-contra {{ background: #fde8e8; color: #b31a1a; }}
.legenda-opiniao .pos-reforco {{ background: #fff3cd; color: #7a5a00; }}
.legenda-opiniao .pos-derrog {{ background: #ede8ff; color: #5a00aa; }}

/* TABELA POLÉMICAS */
.tabela-polemicas .nivel-cell {{ text-align: center; font-weight: 700; font-size: 13px; white-space: nowrap; }}
.nivel-alto {{ background: #fde8e8; color: #b31a1a; }}
.nivel-muito-alto {{ background: #f8b4b4; color: #8b0000; }}
.nivel-medio {{ background: #fff3cd; color: #7a5a00; }}
.nivel-baixo {{ background: #e8f8e8; color: #1a7a1a; }}
.nota-pt {{ font-size: 12px; color: #555; font-style: italic; }}

/* ESTATÍSTICAS */
.stat-categoria {{ font-size: 15px; color: #2980b9; margin: 24px 0 10px; font-weight: 600; }}
.tabela-stats {{ width: 100%; border-collapse: collapse; font-size: 13px; margin-bottom: 20px; }}
.tabela-stats th {{ background: #2980b9; color: #fff; padding: 8px 14px; text-align: left; }}
.tabela-stats td {{ padding: 8px 14px; border-bottom: 1px solid #e0e8f0; }}
.tabela-stats tbody tr:hover {{ background: #eef5ff; }}
.stat-val {{ color: #1a1a2e; font-size: 15px; white-space: nowrap; }}
.stat-nota {{ font-size: 11px; color: #888; font-style: italic; }}

/* NOTA METODOLÓGICA */
.nota-met {{
  background: #f0f4ff; border-left: 4px solid #4a90d9;
  padding: 14px 18px; margin: 20px 0; font-size: 12.5px; color: #444;
  border-radius: 0 4px 4px 0;
}}
.nota-met strong {{ color: #1a1a2e; }}

/* FOOTER */
#footer {{
  text-align: center; padding: 30px; font-size: 11px; color: #999;
  background: #f4f6fb; margin-left: 0;
}}

@media (max-width: 900px) {{
  #sidebar {{ display: none; }}
  #main {{ margin-left: 0; }}
  .section {{ padding: 24px 20px; }}
}}
</style>
</head>
<body>

<!-- SIDEBAR -->
<div id="sidebar">
  <div class="sidebar-header">
    <h2>Análise dos Preâmbulos e Opiniões</h2>
    <p>Regulamento UE 2023/0447 · @codigo · @rgbeac</p>
    <p style="margin-top:8px;color:#4a5a80;">{DATA_PRODUCAO}</p>
  </div>
  <nav class="sidebar-nav">
    <a href="#sec1" class="sec1"><span class="nav-num">1</span> Espírito do Regulamento Europeu</a>
    <a href="#sec2" class="sec2"><span class="nav-num">2</span> Espírito do @codigo</a>
    <a href="#sec3" class="sec3"><span class="nav-num">3</span> Espírito do @rgbeac</a>
    <a href="#sec4" class="sec4"><span class="nav-num">4</span> Temas recorrentes nas opiniões</a>
    <a href="#sec5" class="sec5"><span class="nav-num">5</span> Questões polémicas em Portugal</a>
    <a href="#sec6" class="sec6"><span class="nav-num">6</span> Realidade portuguesa em dados</a>
  </nav>
</div>

<!-- MAIN -->
<div id="main">

<!-- HERO -->
<div id="hero">
  <h1>Análise dos Preâmbulos e Opiniões<br>sobre o Regulamento UE 2023/0447</h1>
  <p>Reflexão sobre o espírito legislativo, os temas recorrentes nas opiniões e a realidade portuguesa</p>
  <p>Produzido em {DATA_PRODUCAO}</p>
  <div class="hero-tags">
    <span class="hero-tag">@regulamento</span>
    <span class="hero-tag">@codigo</span>
    <span class="hero-tag">@rgbeac</span>
    <span class="hero-tag">@opiniao (52 pareceres, 7 grupos)</span>
    <span class="hero-tag">Realidade PT</span>
    <span class="hero-tag">Síntese executiva</span>
  </div>
</div>

<!-- SEC 1 -->
<section id="sec1" class="section s1">
  <div class="section-header">
    <div class="section-num">Secção 1</div>
    <h2 class="section-titulo">{SEC1_TITULO}</h2>
    <div class="section-subtitulo">{SEC1_SUBTITULO}</div>
    <div class="section-divider"></div>
  </div>
  <div class="texto-analise">{SEC1_TEXTO}</div>
  <h3 style="margin: 24px 0 12px; font-size:15px; color:#0077b6;">Objetivos e instrumentos do Regulamento</h3>
  {tbl_obj}
</section>

<!-- SEC 2 -->
<section id="sec2" class="section s2">
  <div class="section-header">
    <div class="section-num">Secção 2</div>
    <h2 class="section-titulo">{SEC2_TITULO}</h2>
    <div class="section-subtitulo">{SEC2_SUBTITULO}</div>
    <div class="section-divider"></div>
  </div>
  <div class="texto-analise">{SEC2_TEXTO}</div>
  <h3 style="margin: 24px 0 12px; font-size:15px; color:#e06c0a;">Comparação @codigo vs. Regulamento Europeu</h3>
  {tbl_comp}
</section>

<!-- SEC 3 -->
<section id="sec3" class="section s3">
  <div class="section-header">
    <div class="section-num">Secção 3</div>
    <h2 class="section-titulo">{SEC3_TITULO}</h2>
    <div class="section-subtitulo">{SEC3_SUBTITULO}</div>
    <div class="section-divider"></div>
  </div>
  <div class="texto-analise">{SEC3_TEXTO}</div>
  <h3 style="margin: 24px 0 12px; font-size:15px; color:#2a9d5c;">Principais inovações do RGBEAC</h3>
  {tbl_inov}
</section>

<!-- SEC 4 -->
<section id="sec4" class="section s4">
  <div class="section-header">
    <div class="section-num">Secção 4</div>
    <h2 class="section-titulo">{SEC4_TITULO}</h2>
    <div class="section-subtitulo">{SEC4_SUBTITULO}</div>
    <div class="section-divider"></div>
  </div>
  <div class="texto-analise">{SEC4_INTRO}</div>
  <div class="nota-met"><strong>Nota metodológica:</strong> Os pareceres foram produzidos sobre versões anteriores do Regulamento. Algumas posições específicas podem já ter sido acolhidas ou ultrapassadas pela versão final. Os temas de fundo, no entanto, permanecem relevantes para a análise de implementação.</div>
  {tbl_opin}
  <div class="texto-analise" style="margin-top:24px;">{SEC4_ANALISE}</div>
</section>

<!-- SEC 5 -->
<section id="sec5" class="section s5">
  <div class="section-header">
    <div class="section-num">Secção 5</div>
    <h2 class="section-titulo">{SEC5_TITULO}</h2>
    <div class="section-subtitulo">{SEC5_SUBTITULO}</div>
    <div class="section-divider"></div>
  </div>
  <div class="texto-analise">{SEC5_INTRO}</div>
  {tbl_pol}
</section>

<!-- SEC 6 -->
<section id="sec6" class="section s6">
  <div class="section-header">
    <div class="section-num">Secção 6</div>
    <h2 class="section-titulo">{SEC6_TITULO}</h2>
    <div class="section-subtitulo">{SEC6_SUBTITULO}</div>
    <div class="section-divider"></div>
  </div>
  <div class="texto-analise">{SEC6_TEXTO}</div>
  {blocos_stat}
  <div class="texto-analise" style="margin-top:20px;">{SEC6_REFLEXAO}</div>
  <div class="nota-met" style="margin-top:20px;">
    <strong>Fontes:</strong> SIAC (Sistema de Informação de Animais de Companhia); Censo Nacional de Animais Errantes 2023 (ICNF / Universidade de Aveiro); estudo ScienceDirect 2025 sobre dados SIAC 2023; The Portugal News (oct. 2023); Euroweekly News (oct. 2023); OMV — Ordem dos Médicos Veterinários.
  </div>
</section>

</div><!-- /main -->

<div id="footer">
  Análise produzida no âmbito do projeto de análise comparativa da legislação portuguesa e europeia sobre animais de companhia · {DATA_PRODUCAO}
</div>

</body>
</html>"""
    return html

# ─────────────────────────────────────────────────────────────────────────────
# GERAÇÃO WORD
# ─────────────────────────────────────────────────────────────────────────────

def limpar_html(texto):
    """Remove tags HTML simples para usar em texto Word."""
    import re
    texto = re.sub(r'<strong>(.*?)</strong>', r'\1', texto, flags=re.DOTALL)
    texto = re.sub(r'<em>(.*?)</em>', r'\1', texto, flags=re.DOTALL)
    texto = re.sub(r'<p>(.*?)</p>', r'\1\n', texto, flags=re.DOTALL)
    texto = re.sub(r'<ul>(.*?)</ul>', r'\1', texto, flags=re.DOTALL)
    texto = re.sub(r'<li>(.*?)</li>', r'  • \1\n', texto, flags=re.DOTALL)
    texto = re.sub(r'<[^>]+>', '', texto)
    return texto.strip()

def gerar_word():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("ERRO: python-docx não está instalado. Execute: pip install python-docx")
        sys.exit(1)

    doc = Document()

    # Margens
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    def add_titulo_doc(doc, texto):
        p = doc.add_heading(texto, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(26, 26, 46)
            run.font.size = Pt(18)

    def add_subtitulo_doc(doc, texto):
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 130)
            run.font.size = Pt(11)

    def add_sec_titulo(doc, num, titulo, subtitulo, cor):
        doc.add_page_break()
        p = doc.add_heading(f"Secção {num} — {titulo}", level=1)
        for run in p.runs:
            run.font.color.rgb = cor
        if subtitulo:
            ps = doc.add_paragraph(subtitulo)
            for run in ps.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 120)
                run.font.size = Pt(10.5)

    def add_texto(doc, texto_html):
        import re
        # Extrair parágrafos
        paragrafos = re.findall(r'<p>(.*?)</p>', texto_html, re.DOTALL)
        for par in paragrafos:
            texto = limpar_html(f'<p>{par}</p>')
            p = doc.add_paragraph(texto)
            p.paragraph_format.space_after = Pt(8)
        # Extrair bullets
        bullets = re.findall(r'<li>(.*?)</li>', texto_html, re.DOTALL)
        for b in bullets:
            texto = limpar_html(f'<li>{b}</li>').replace('  • ', '')
            doc.add_paragraph(texto, style='List Bullet')

    def add_tabela_2col(doc, cabecalhos, linhas, cor_header=None):
        if cor_header is None:
            cor_header = RGBColor(26, 26, 46)
        table = doc.add_table(rows=1, cols=len(cabecalhos))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, cab in enumerate(cabecalhos):
            hdr[i].text = cab
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
            tc = hdr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '%02x%02x%02x' % cor_header)
            tcPr.append(shd)
        for linha in linhas:
            row = table.add_row().cells
            for i, val in enumerate(linha):
                row[i].text = val
                for run in row[i].paragraphs[0].runs:
                    run.font.size = Pt(9)
        doc.add_paragraph()

    # Capa
    add_titulo_doc(doc, "Análise dos Preâmbulos e Opiniões")
    add_titulo_doc(doc, "Regulamento UE 2023/0447")
    doc.add_paragraph()
    add_subtitulo_doc(doc, "Reflexão sobre o espírito legislativo, os temas recorrentes nas opiniões e a realidade portuguesa")
    add_subtitulo_doc(doc, f"Produzido em {DATA_PRODUCAO}")
    doc.add_paragraph()
    add_subtitulo_doc(doc, "@regulamento · @codigo · @rgbeac · @opiniao (52 pareceres, 7 grupos) · Realidade PT")

    # SEC 1
    add_sec_titulo(doc, 1, SEC1_TITULO, SEC1_SUBTITULO, RGBColor(0, 119, 182))
    add_texto(doc, SEC1_TEXTO)
    doc.add_heading("Objetivos e instrumentos do Regulamento", level=2)
    add_tabela_2col(doc, ["Objetivo", "Instrumento / Mecanismo previsto"],
                    [(o, i) for o, i in SEC1_OBJETIVOS],
                    cor_header=RGBColor(0, 119, 182))

    # SEC 2
    add_sec_titulo(doc, 2, SEC2_TITULO, SEC2_SUBTITULO, RGBColor(224, 108, 10))
    add_texto(doc, SEC2_TEXTO)
    doc.add_heading("Comparação @codigo vs. Regulamento Europeu", level=2)
    add_tabela_2col(doc, ["Dimensão", "@codigo (DL 214/2013)", "Regulamento 2023/0447"],
                    [(d, c, r) for d, c, r in SEC2_COMPARACAO],
                    cor_header=RGBColor(224, 108, 10))

    # SEC 3
    add_sec_titulo(doc, 3, SEC3_TITULO, SEC3_SUBTITULO, RGBColor(42, 157, 92))
    add_texto(doc, SEC3_TEXTO)
    doc.add_heading("Principais inovações do RGBEAC", level=2)
    add_tabela_2col(doc, ["Inovação", "Conteúdo", "Países / Referências"],
                    [(i, d, r) for i, d, r in SEC3_INOVACOES],
                    cor_header=RGBColor(42, 157, 92))

    # SEC 4
    add_sec_titulo(doc, 4, SEC4_TITULO, SEC4_SUBTITULO, RGBColor(155, 89, 182))
    add_texto(doc, SEC4_INTRO)
    nota = doc.add_paragraph("Nota metodológica: Os pareceres foram produzidos sobre versões anteriores do Regulamento. Algumas posições específicas podem já ter sido acolhidas ou ultrapassadas pela versão final. Os temas de fundo permanecem relevantes para a análise de implementação.")
    nota.paragraph_format.left_indent = Cm(1)
    for run in nota.runs:
        run.font.italic = True
        run.font.size = Pt(9)

    # Tabela opiniões — formato condensado (3 colunas para caber em Word)
    doc.add_heading("Posições por organização e tema (síntese)", level=2)
    temas_abrev = ["L1", "L2", "L3", "L4", "L5", "L6", "L7", "L8", "L9", "L10"]
    linhas_opin = []
    grupo_ant = None
    for org, tipo, grupo, posicoes in OPINIAO_DADOS:
        if grupo != grupo_ant:
            grupo_ant = grupo
            linhas_opin.append([f"── {grupo} ──", "", ""])
        pos_sintet = "; ".join(
            f"{temas_abrev[i]}: {p}" for i, p in enumerate(posicoes) if p
        ) or "—"
        linhas_opin.append([org, grupo, pos_sintet])
    add_tabela_2col(doc, ["Organização", "Grupo", "Principais posições"],
                    linhas_opin, cor_header=RGBColor(155, 89, 182))
    add_texto(doc, SEC4_ANALISE)

    # SEC 5
    add_sec_titulo(doc, 5, SEC5_TITULO, SEC5_SUBTITULO, RGBColor(231, 76, 60))
    add_texto(doc, SEC5_INTRO)
    doc.add_heading("Análise por tema", level=2)
    for tema, artigo, polemica, nivel, nota in SEC5_TEMAS:
        p = doc.add_heading(f"{tema} ({artigo})", level=3)
        for run in p.runs:
            run.font.color.rgb = RGBColor(180, 50, 30)
        doc.add_paragraph(f"Impacto em Portugal: {nivel_label(nivel)}")
        doc.add_paragraph(polemica)
        nota_p = doc.add_paragraph(f"Nota PT: {nota}")
        for run in nota_p.runs:
            run.font.italic = True
            run.font.size = Pt(10)
        doc.add_paragraph()

    # SEC 6
    add_sec_titulo(doc, 6, SEC6_TITULO, SEC6_SUBTITULO, RGBColor(41, 128, 185))
    add_texto(doc, SEC6_TEXTO)
    for bloco in SEC6_ESTATISTICAS:
        doc.add_heading(bloco['categoria'], level=2)
        add_tabela_2col(doc, ["Indicador", "Valor", "Nota"],
                        [(i, v, n) for i, v, n in bloco['items']],
                        cor_header=RGBColor(41, 128, 185))
    add_texto(doc, SEC6_REFLEXAO)
    nota_fontes = doc.add_paragraph("Fontes: SIAC; Censo Nacional de Animais Errantes 2023 (ICNF / Universidade de Aveiro); estudo ScienceDirect 2025; The Portugal News (out. 2023); Euroweekly News (out. 2023); OMV.")
    for run in nota_fontes.runs:
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

    return doc

# ─────────────────────────────────────────────────────────────────────────────
# GERAÇÃO EXCEL
# ─────────────────────────────────────────────────────────────────────────────

def gerar_excel():
    try:
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("ERRO: openpyxl não está instalado. Execute: pip install openpyxl")
        sys.exit(1)

    wb = Workbook()
    wb.remove(wb.active)  # Remove default sheet

    # Paleta de cores (RGB hex)
    cor_header = "CBD5E1"  # cinzento-azul suave
    cor_favor = "DCEEDE"   # verde pastel
    cor_contra = "EAD6D6"  # rosa suave
    cor_reforco = "F5E7D1" # bege suave
    cor_derrog = "E8E0F0"  # roxo pastel
    cor_impacto_alto = "D9D4D4"  # cinzento suave
    cor_fundo = "FAFBFC"   # cinzento muito claro
    cor_linha = "E5E7EB"   # cinzento claro

    # Borders
    thin_border = Border(
        left=Side(style='thin', color='E5E7EB'),
        right=Side(style='thin', color='E5E7EB'),
        top=Side(style='thin', color='E5E7EB'),
        bottom=Side(style='thin', color='E5E7EB')
    )

    def adicionar_sec(titulo, linhas_conteudo, cor_header_sec=cor_header):
        """Adiciona uma nova folha com conteúdo estruturado."""
        ws = wb.create_sheet(titulo)

        # Larguras de coluna
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 45
        ws.column_dimensions['C'].width = 35

        row = 1
        for conteudo in linhas_conteudo:
            if isinstance(conteudo, dict):
                if conteudo['tipo'] == 'titulo':
                    ws.merge_cells(f'A{row}:C{row}')
                    cell = ws[f'A{row}']
                    cell.value = conteudo['valor']
                    cell.font = Font(name='Calibri', size=14, bold=True, color='5B7A8F')
                    cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                    row += 1
                    row += 1  # Espacamento
                elif conteudo['tipo'] == 'tabela':
                    # Cabeçalho
                    for col_idx, cab in enumerate(conteudo['cabecalhos'], 1):
                        cell = ws.cell(row=row, column=col_idx)
                        cell.value = cab
                        cell.fill = PatternFill(start_color=cor_header_sec, end_color=cor_header_sec, fill_type='solid')
                        cell.font = Font(name='Calibri', size=10, bold=True, color='1A1A2E')
                        cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                        cell.border = thin_border
                    row += 1

                    # Dados
                    for linha in conteudo['dados']:
                        for col_idx, val in enumerate(linha, 1):
                            cell = ws.cell(row=row, column=col_idx)
                            cell.value = val
                            cell.font = Font(name='Calibri', size=9)
                            cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
                            cell.border = thin_border
                            # Alternância de cores
                            if row % 2 == 0:
                                cell.fill = PatternFill(start_color='F8F9FC', end_color='F8F9FC', fill_type='solid')
                        row += 1
                    row += 2  # Espacamento entre tabelas

        return ws

    # SEC 1 — Regulamento Europeu
    adicionar_sec("Sec1", [
        {'tipo': 'titulo', 'valor': 'Espírito do Regulamento Europeu 2023/0447'},
        {'tipo': 'tabela', 'cabecalhos': ['Objetivo', 'Instrumento / Mecanismo'],
         'dados': [(o, i) for o, i in SEC1_OBJETIVOS]},
    ], cor_favor)

    # SEC 2 — @codigo
    adicionar_sec("Sec2", [
        {'tipo': 'titulo', 'valor': 'Espírito do @codigo (DL 214/2013) vs. Regulamento'},
        {'tipo': 'tabela', 'cabecalhos': ['Dimensão', '@codigo', 'Regulamento 2023/0447'],
         'dados': [(d, c, r) for d, c, r in SEC2_COMPARACAO]},
    ], "DCC9C6")

    # SEC 3 — RGBEAC
    adicionar_sec("Sec3", [
        {'tipo': 'titulo', 'valor': 'Principais inovações do RGBEAC'},
        {'tipo': 'tabela', 'cabecalhos': ['Inovação', 'Conteúdo', 'Países / Referências'],
         'dados': [(i, d, r) for i, d, r in SEC3_INOVACOES]},
    ], "C8DCC6")

    # SEC 4 — Opiniões (matriz com coluna Grupo)
    ws_sec4 = wb.create_sheet("Sec4")
    ws_sec4.column_dimensions['A'].width = 22  # Organização
    ws_sec4.column_dimensions['B'].width = 22  # Grupo
    for idx in range(len(OPINIAO_TEMAS)):
        ws_sec4.column_dimensions[get_column_letter(idx+3)].width = 14

    # Cores pastel por grupo
    cores_grupo = {
        "Bem-estar animal":         "EAF4EA",
        "Veterinária":              "EAF0FA",
        "Kennel clubs e felinologia": "FFF8EA",
        "Caçadores e uso funcional": "FFF0EA",
        "Indústria e comércio":     "F5EAFA",
        "Rastreabilidade e I&R":    "EAF7FA",
        "Autoridades públicas":     "F0EAF5",
    }

    row = 1
    # Título
    n_cols = len(OPINIAO_TEMAS) + 2
    ws_sec4.merge_cells(f'A1:{get_column_letter(n_cols)}{row}')
    cell_title = ws_sec4['A1']
    cell_title.value = 'Matriz de Opiniões: 52 Organizações × 10 Temas (7 grupos)'
    cell_title.font = Font(name='Calibri', size=14, bold=True, color='5B7A8F')
    cell_title.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    row += 2

    # Cabeçalho
    for col_idx, cab in enumerate(['Organização', 'Grupo'], 1):
        cell = ws_sec4.cell(row=row, column=col_idx)
        cell.value = cab
        cell.fill = PatternFill(start_color=cor_header, end_color=cor_header, fill_type='solid')
        cell.font = Font(name='Calibri', size=9, bold=True, color='1A1A2E')
        cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        cell.border = thin_border

    for col_idx, tema in enumerate(OPINIAO_TEMAS, 3):
        cell = ws_sec4.cell(row=row, column=col_idx)
        cell.value = tema.replace('\n', ' ')
        cell.fill = PatternFill(start_color=cor_header, end_color=cor_header, fill_type='solid')
        cell.font = Font(name='Calibri', size=8, bold=True, color='1A1A2E')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
    row += 1

    # Dados opiniões — separador por grupo
    grupo_atual = None
    for org, tipo, grupo, posicoes in OPINIAO_DADOS:
        # Linha separadora de grupo
        if grupo != grupo_atual:
            grupo_atual = grupo
            ws_sec4.merge_cells(f'A{row}:{get_column_letter(n_cols)}{row}')
            cell_grp = ws_sec4[f'A{row}']
            cell_grp.value = grupo
            cor_grp = cores_grupo.get(grupo, "F0F0F0")
            cell_grp.fill = PatternFill(start_color=cor_grp, end_color=cor_grp, fill_type='solid')
            cell_grp.font = Font(name='Calibri', size=9, bold=True, color='2D4059')
            cell_grp.alignment = Alignment(horizontal='left', vertical='center')
            ws_sec4.row_dimensions[row].height = 16
            row += 1

        cor_grp = cores_grupo.get(grupo, "FFFFFF")

        # Organização
        cell_org = ws_sec4.cell(row=row, column=1)
        cell_org.value = org
        cell_org.font = Font(name='Calibri', size=9, bold=True)
        cell_org.border = thin_border
        cell_org.fill = PatternFill(start_color=cor_grp, end_color=cor_grp, fill_type='solid')

        # Grupo
        cell_grp2 = ws_sec4.cell(row=row, column=2)
        cell_grp2.value = grupo
        cell_grp2.font = Font(name='Calibri', size=8, color='666666')
        cell_grp2.border = thin_border
        cell_grp2.fill = PatternFill(start_color=cor_grp, end_color=cor_grp, fill_type='solid')

        for col_idx, pos in enumerate(posicoes, 3):
            cell = ws_sec4.cell(row=row, column=col_idx)
            cell.value = pos if pos else ""
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.font = Font(name='Calibri', size=8)

            if pos == "Favor":
                cell.fill = PatternFill(start_color=cor_favor, end_color=cor_favor, fill_type='solid')
            elif pos == "Contra":
                cell.fill = PatternFill(start_color=cor_contra, end_color=cor_contra, fill_type='solid')
            elif pos == "Reforço":
                cell.fill = PatternFill(start_color=cor_reforco, end_color=cor_reforco, fill_type='solid')
            elif pos == "Derrogação":
                cell.fill = PatternFill(start_color=cor_derrog, end_color=cor_derrog, fill_type='solid')
        row += 1

    # SEC 5 — Questões polémicas
    sec5_dados = []
    for tema, artigo, polemica, nivel, nota in SEC5_TEMAS:
        sec5_dados.append([f"{tema}\n({artigo})", polemica, nivel_label(nivel), nota])

    adicionar_sec("Sec5", [
        {'tipo': 'titulo', 'valor': 'Questões polémicas / Difícil aplicação em Portugal'},
        {'tipo': 'tabela', 'cabecalhos': ['Tema', 'Polémica', 'Impacto PT', 'Nota específica'],
         'dados': sec5_dados},
    ], cor_contra)

    # SEC 6 — Dados estatísticos
    ws_sec6 = wb.create_sheet("Sec6")
    ws_sec6.column_dimensions['A'].width = 30
    ws_sec6.column_dimensions['B'].width = 18
    ws_sec6.column_dimensions['C'].width = 35

    row = 1
    for bloco in SEC6_ESTATISTICAS:
        # Título do bloco
        ws_sec6.merge_cells(f'A{row}:C{row}')
        cell = ws_sec6[f'A{row}']
        cell.value = bloco['categoria']
        cell.font = Font(name='Calibri', size=12, bold=True, color='2980B9')
        cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        row += 1

        # Cabeçalho da tabela
        for col_idx, cab in enumerate(['Indicador', 'Valor', 'Nota'], 1):
            cell = ws_sec6.cell(row=row, column=col_idx)
            cell.value = cab
            cell.fill = PatternFill(start_color=cor_header, end_color=cor_header, fill_type='solid')
            cell.font = Font(name='Calibri', size=9, bold=True, color='1A1A2E')
            cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
            cell.border = thin_border
        row += 1

        # Dados
        for ind, val, n in bloco['items']:
            ws_sec6.cell(row=row, column=1).value = ind
            ws_sec6.cell(row=row, column=2).value = val
            ws_sec6.cell(row=row, column=3).value = n

            for col in range(1, 4):
                cell = ws_sec6.cell(row=row, column=col)
                cell.border = thin_border
                cell.font = Font(name='Calibri', size=9)
                cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
                if row % 2 == 0:
                    cell.fill = PatternFill(start_color='F8F9FC', end_color='F8F9FC', fill_type='solid')
            row += 1
        row += 2  # Espacamento

    # Folha de Legenda
    ws_leg = wb.create_sheet("Legenda")
    ws_leg.column_dimensions['A'].width = 20
    ws_leg.column_dimensions['B'].width = 40

    leg_row = 1
    ws_leg.cell(row=leg_row, column=1).value = "Paleta de cores — Legenda"
    ws_leg.cell(row=leg_row, column=1).font = Font(name='Calibri', size=14, bold=True, color='5B7A8F')
    leg_row += 2

    legendas = [
        ("Verde pastel", cor_favor, "Posição Favor / Apoio nas opiniões"),
        ("Rosa suave", cor_contra, "Posição Contra / Crítica nas opiniões"),
        ("Bege suave", cor_reforco, "Propõe reforço / Crítica moderada"),
        ("Roxo pastel", cor_derrog, "Propõe derrogação / Flexibilização"),
        ("Cinzento suave", cor_impacto_alto, "Impacto Muito Alto em Portugal"),
    ]

    for cor_nome, cor_hex, descricao in legendas:
        ws_leg.cell(row=leg_row, column=1).value = cor_nome
        ws_leg.cell(row=leg_row, column=2).value = descricao
        ws_leg.cell(row=leg_row, column=1).fill = PatternFill(start_color=cor_hex, end_color=cor_hex, fill_type='solid')
        ws_leg.cell(row=leg_row, column=1).font = Font(name='Calibri', size=10)
        ws_leg.cell(row=leg_row, column=2).font = Font(name='Calibri', size=10)
        ws_leg.cell(row=leg_row, column=1).border = thin_border
        ws_leg.cell(row=leg_row, column=2).border = thin_border
        ws_leg.cell(row=leg_row, column=1).alignment = Alignment(horizontal='center', vertical='center')
        ws_leg.cell(row=leg_row, column=2).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        leg_row += 1

    return wb

# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Gera análise dos preâmbulos e opiniões.")
    parser.add_argument('--word', action='store_true', help='Gera documento Word (.docx)')
    parser.add_argument('--excel', action='store_true', help='Gera documento Excel (.xlsx)')
    parser.add_argument('--all', action='store_true', help='Gera HTML, Word e Excel')
    args = parser.parse_args()

    gerar_html_flag = not args.word and not args.excel or args.all
    gerar_word_flag = args.word or args.all
    gerar_excel_flag = args.excel or args.all

    if gerar_html_flag:
        html = gerar_html()
        ficheiro_html = "analise_preambulos_opinoes.html"
        with open(ficheiro_html, 'w', encoding='utf-8') as f:
            f.write(html)
        print(f"✓ HTML gerado: {ficheiro_html}")

    if gerar_word_flag:
        doc = gerar_word()
        ficheiro_docx = "analise_preambulos_opinoes.docx"
        doc.save(ficheiro_docx)
        print(f"✓ Word gerado: {ficheiro_docx}")

    if gerar_excel_flag:
        wb = gerar_excel()
        ficheiro_xlsx = "analise_preambulos_opinoes.xlsx"
        wb.save(ficheiro_xlsx)
        print(f"✓ Excel gerado: {ficheiro_xlsx}")

    if not gerar_html_flag and not gerar_word_flag and not gerar_excel_flag:
        print("Uso: python3 gerar_analise_preambulos.py [--word] [--excel] [--all]")

if __name__ == '__main__':
    main()
