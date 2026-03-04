"""
Gera documento Word (.docx) comparativo artigo a artigo — visualmente fiel ao HTML.
Lê os dados de gerar_comparativo_reuniao.py e produz comparativo_reuniao_exemplo.docx
"""

import os
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Importar dados do script existente
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from gerar_comparativo_reuniao import ARTIGOS, COR

# ---------------------------------------------------------------------------
# UTILITÁRIOS DE COR E ESTILO
# ---------------------------------------------------------------------------

def hex_to_rgb(hex_str):
    """Converte string hex (sem #) para tuplo (R, G, B)."""
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color):
    """Define cor de fundo de uma célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    """Define bordas finas em todas as faces da célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_vertical_align(cell, align="top"):
    """Define alinhamento vertical da célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def add_run_styled(paragraph, text, bold=False, italic=False,
                   font_name="Calibri", font_size=10,
                   color_hex=None):
    """Adiciona run com estilo completo ao parágrafo."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color_hex:
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    return run


def remove_table_borders(table):
    """Remove bordas externas da tabela (deixa apenas bordas de células)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def set_table_width(table, width_cm=None, pct=None):
    """Define largura da tabela (cm ou percentagem 0-5000)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    if pct is not None:
        tblW.set(qn("w:w"), str(pct))
        tblW.set(qn("w:type"), "pct")
    elif width_cm is not None:
        tblW.set(qn("w:w"), str(int(width_cm * 567)))
        tblW.set(qn("w:type"), "dxa")
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)


def set_col_width(col_cells, width_cm):
    """Define largura de coluna via primeira célula."""
    for cell in col_cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width_cm * 567)))
        tcW.set(qn("w:type"), "dxa")
        existing = tcPr.find(qn("w:tcW"))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(tcW)


def cell_header(cell, text, bg_hex, text_hex="FFFFFF",
                font_size=9, bold=True, align="center"):
    """Preenche célula como cabeçalho colorido."""
    set_cell_bg(cell, bg_hex)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run_styled(p, text, bold=bold, font_size=font_size, color_hex=text_hex)


_RE_ALINEA  = re.compile(r"^[a-z]\)|^\([a-z-]+\)")
_RE_SUB     = re.compile(r"^[—–]\s|^[—–]$")
_RE_ART_HDR = re.compile(r"^Artigo\s+\d+", re.IGNORECASE)


def _classify_line(line):
    """Devolve 'art-header', 'sub', 'alinea' ou 'normal'."""
    t = line.strip()
    if _RE_ART_HDR.match(t):
        return "art-header"
    if _RE_SUB.match(t):
        return "sub"
    if _RE_ALINEA.match(t):
        return "alinea"
    return "normal"


def _add_para(cell, first_flag):
    """Cria novo parágrafo na célula; na primeira chamada reutiliza o existente."""
    from docx.text.paragraph import Paragraph as _Para
    if first_flag[0]:
        first_flag[0] = False
        return cell.paragraphs[0]
    new_p = OxmlElement("w:p")
    cell._tc.append(new_p)
    return _Para(new_p, cell)


def cell_body(cell, text, bg_hex, ref_text=None,
              font_size=9.5, bold=False, italic=False,
              text_hex="222222"):
    """Preenche célula detectando automaticamente alíneas e subalíneas."""
    set_cell_bg(cell, bg_hex)
    set_cell_borders(cell)
    set_cell_vertical_align(cell, "top")

    first = [True]  # mutable flag para _add_para

    # ── Referência (opcional) ────────────────────────────────────────────
    if ref_text:
        p = _add_para(cell, first)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(4)
        p.paragraph_format.right_indent = Pt(4)
        add_run_styled(p, ref_text, bold=True, font_size=8,
                       color_hex="555555", italic=True)

    # ── Corpo do texto: blocos separados por \n\n, linhas por \n ─────────
    blocos = [b for b in text.split("\n\n") if b.strip()]
    for bi, bloco in enumerate(blocos):
        # Detectar bloco [dim] (texto secundário a cinza)
        is_dim = bloco.startswith("[dim]")
        bloco_txt = bloco[5:].strip() if is_dim else bloco
        dim_hex = "AAAAAA" if is_dim else text_hex

        linhas = [l.strip() for l in bloco_txt.split("\n") if l.strip()]
        for li, linha in enumerate(linhas):
            kind = _classify_line(linha)
            p = _add_para(cell, first)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.right_indent = Pt(4)

            # Espaçamento: maior entre blocos (\n\n), menor entre linhas (\n)
            if bi == 0 and li == 0 and not ref_text:
                p.paragraph_format.space_before = Pt(4)
            elif li == 0:
                p.paragraph_format.space_before = Pt(7) if kind != "alinea" else Pt(5)
            else:
                p.paragraph_format.space_before = Pt(1)

            # Indentação por tipo
            if kind == "art-header":
                p.paragraph_format.left_indent = Pt(4)
                p.paragraph_format.first_line_indent = Pt(0)
                # Linha separadora via espaço antes aumentado (exceto primeiro)
                if bi > 0:
                    p.paragraph_format.space_before = Pt(10)
                add_run_styled(p, linha, bold=True, italic=False,
                               font_size=font_size - 0.5, color_hex="333333")
            elif kind == "sub":
                p.paragraph_format.left_indent = Pt(34)
                p.paragraph_format.first_line_indent = Pt(-14)
                add_run_styled(p, linha, bold=bold, italic=italic,
                               font_size=font_size, color_hex=dim_hex)
            elif kind == "alinea":
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.first_line_indent = Pt(-14)
                add_run_styled(p, linha, bold=bold, italic=italic,
                               font_size=font_size, color_hex=dim_hex)
            else:
                p.paragraph_format.left_indent = Pt(4)
                p.paragraph_format.first_line_indent = Pt(0)
                add_run_styled(p, linha, bold=bold, italic=italic,
                               font_size=font_size, color_hex=dim_hex)


# ---------------------------------------------------------------------------
# SECÇÕES DO DOCUMENTO
# ---------------------------------------------------------------------------

def add_page_title(doc):
    """Adiciona página de título / cabeçalho do documento."""
    # Bloco de título com fundo escuro (simulado via tabela 1x1)
    t = doc.add_table(rows=1, cols=1)
    set_table_width(t, pct=5000)
    cell = t.cell(0, 0)
    set_cell_bg(cell, "1A1A2E")
    set_cell_borders(cell, "1A1A2E")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    add_run_styled(p, "Comparativo Artigo a Artigo", bold=True,
                   font_size=16, color_hex="FFFFFF")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(14)
    add_run_styled(p2, "Regulamento 2023/0447 — Cães e Gatos",
                   bold=False, font_size=11, color_hex="7EC8E3")
    doc.add_paragraph()  # espaço


def add_legenda(doc):
    """Adiciona tabela de legenda das cores."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run_styled(p, "LEGENDA DO SISTEMA DE CORES", bold=True,
                   font_size=9, color_hex="4A4A4A")

    legenda = [
        ("@regulamento (texto EN)", COR["regulamento_body"]),
        ("@regulamento (tradução PT)", COR["regulamento_trad"]),
        ("@rgbeac (proposta jun. 2025)", COR["rgbeac_body"]),
        ("@codigo (DL n.º 214/2013)", COR["codigo_body"]),
        ("@legislacao (legislação vigente)", COR["legislacao_body"]),
        ("Divergência face ao Regulamento", COR["divergencia_body"]),
        ("Notas de Reunião", COR["notas_body"]),
    ]
    t = doc.add_table(rows=len(legenda), cols=2)
    set_table_width(t, pct=2500)
    for i, (label, cor) in enumerate(legenda):
        c0 = t.cell(i, 0)
        c1 = t.cell(i, 1)
        set_cell_bg(c0, cor)
        set_cell_bg(c1, cor)
        set_cell_borders(c0)
        set_cell_borders(c1)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        p0.paragraph_format.left_indent = Pt(6)
        add_run_styled(p0, label, bold=True, font_size=9, color_hex="222222")
    doc.add_paragraph()


def add_article_section(doc, art):
    """Adiciona secção completa de um artigo."""

    # ── Cabeçalho do Artigo ──────────────────────────────────────────────
    t_badge = doc.add_table(rows=1, cols=2)
    set_table_width(t_badge, pct=5000)
    remove_table_borders(t_badge)

    c_id = t_badge.cell(0, 0)
    set_cell_bg(c_id, "1A1A2E")
    set_cell_borders(c_id, "1A1A2E")
    p_id = c_id.paragraphs[0]
    p_id.paragraph_format.space_before = Pt(8)
    p_id.paragraph_format.space_after = Pt(8)
    p_id.paragraph_format.left_indent = Pt(10)
    add_run_styled(p_id, art["id"], bold=True, font_size=11, color_hex="FFFFFF")

    c_tema = t_badge.cell(0, 1)
    set_cell_bg(c_tema, "2C3E6B")
    set_cell_borders(c_tema, "2C3E6B")
    p_tema = c_tema.paragraphs[0]
    p_tema.paragraph_format.space_before = Pt(8)
    p_tema.paragraph_format.space_after = Pt(8)
    p_tema.paragraph_format.left_indent = Pt(10)
    add_run_styled(p_tema, art["tema"], bold=True, font_size=11, color_hex="7EC8E3")

    # ── Texto original EN (leitura primária) ─────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    t_en = doc.add_table(rows=2, cols=1)
    set_table_width(t_en, pct=5000)
    titulo_en = art["regulamento"].get("titulo", "")
    hdr_en = f"@regulamento — {titulo_en}  ·  {art['regulamento']['ref']}  |  EN"
    cell_header(t_en.cell(0, 0), hdr_en, COR["regulamento_header"])
    cell_body(t_en.cell(1, 0),
              art["regulamento"]["texto"],
              COR["regulamento_body"])

    # ── Tradução PT-PT ────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    t_trad = doc.add_table(rows=2, cols=1)
    set_table_width(t_trad, pct=5000)
    cell_header(t_trad.cell(0, 0),
                f"@regulamento — Tradução PT-PT  ·  {art['regulamento']['ref']}",
                "2471A3")
    cell_body(t_trad.cell(1, 0),
              art["regulamento"]["traducao"],
              COR["regulamento_trad"],
              italic=True)

    # ── Tabela 3 colunas: RGBEAC | Código | Legislação vigente ───────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    t3 = doc.add_table(rows=2, cols=3)
    set_table_width(t3, pct=5000)

    # Cabeçalhos
    cell_header(t3.cell(0, 0), "@rgbeac (proposta jun. 2025)",
                COR["rgbeac_header"])
    cell_header(t3.cell(0, 1), "@codigo (DL n.º 214/2013)",
                COR["codigo_header"])
    cell_header(t3.cell(0, 2), "@legislacao (legislação vigente)",
                COR["legislacao_header"])

    # Corpos
    cell_body(t3.cell(1, 0),
              art["rgbeac"]["texto"],
              COR["rgbeac_body"],
              ref_text=art["rgbeac"]["ref"])
    cell_body(t3.cell(1, 1),
              art["codigo"]["texto"],
              COR["codigo_body"],
              ref_text=art["codigo"]["ref"])
    cell_body(t3.cell(1, 2),
              art["legislacao"]["texto"],
              COR["legislacao_body"],
              ref_text=art["legislacao"]["ref"])

    # ── Divergência estruturada (3 colunas + sumário merged) ───────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    t_div = doc.add_table(rows=3, cols=3)
    set_table_width(t_div, pct=5000)

    div = art.get("divergencia", {})

    # Linha 0: Headers (3 colunas) — Ordem: RGBEAC | CÓDIGO | LEGISLAÇÃO
    cell_header(t_div.cell(0, 0), "@rgbeac", COR["rgbeac_header"])
    cell_header(t_div.cell(0, 1), "@codigo", COR["codigo_header"])
    cell_header(t_div.cell(0, 2), "@legislacao", COR["legislacao_header"])

    # Linha 1: Textos (3 colunas) — Ordem: RGBEAC | CÓDIGO | LEGISLAÇÃO
    cell_body(t_div.cell(1, 0),
              div.get("rgbeac", ""),
              COR["rgbeac_body"])
    cell_body(t_div.cell(1, 1),
              div.get("codigo", ""),
              COR["codigo_body"])
    cell_body(t_div.cell(1, 2),
              div.get("legislacao", ""),
              COR["legislacao_body"])

    # Linha 2: Sumário (merged 3 colunas com amplitude total)
    cell_sumario = t_div.cell(2, 0)
    set_cell_bg(cell_sumario, COR["divergencia_body"])
    set_cell_borders(cell_sumario)
    set_cell_vertical_align(cell_sumario, "top")

    # Merge cells (2, 1) e (2, 2) com (2, 0)
    for col_idx in [1, 2]:
        cell_to_merge = t_div.cell(2, col_idx)
        tc = cell_to_merge._tc
        tcPr = tc.get_or_add_tcPr()
        tcMerge = OxmlElement("w:tcMerge")
        tcPr.append(tcMerge)

    # Header do sumário (primeira linha da célula)
    p_header = cell_sumario.paragraphs[0]
    p_header.paragraph_format.space_before = Pt(4)
    p_header.paragraph_format.space_after = Pt(6)
    p_header.paragraph_format.left_indent = Pt(6)
    p_header.paragraph_format.right_indent = Pt(4)
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_styled(p_header,
                   f"Sumário / Proposta  ·  Necessidade de alteração: {art['necessidade_alteracao']}",
                   bold=True, font_size=9, color_hex="222222")

    # Corpo do sumário com parsing de alíneas
    sumario_texto = div.get("sumario", "")
    first = [True]
    first[0] = False  # Já usamos o primeiro parágrafo para header

    blocos = [b for b in sumario_texto.split("\n\n") if b.strip()]
    for bi, bloco in enumerate(blocos):
        is_dim = bloco.startswith("[dim]")
        bloco_txt = bloco[5:].strip() if is_dim else bloco
        dim_hex = "AAAAAA" if is_dim else "222222"

        linhas = [l.strip() for l in bloco_txt.split("\n") if l.strip()]
        for li, linha in enumerate(linhas):
            kind = _classify_line(linha)
            p = _add_para(cell_sumario, first)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.right_indent = Pt(4)

            if li == 0 and bi == 0:
                p.paragraph_format.space_before = Pt(2)
            elif li == 0:
                p.paragraph_format.space_before = Pt(7) if kind != "alinea" else Pt(5)
            else:
                p.paragraph_format.space_before = Pt(1)

            if kind == "art-header":
                p.paragraph_format.left_indent = Pt(4)
                p.paragraph_format.first_line_indent = Pt(0)
                if bi > 0:
                    p.paragraph_format.space_before = Pt(10)
                add_run_styled(p, linha, bold=True, italic=False,
                               font_size=8.5, color_hex="333333")
            elif kind == "sub":
                p.paragraph_format.left_indent = Pt(34)
                p.paragraph_format.first_line_indent = Pt(-14)
                add_run_styled(p, linha, bold=False, italic=False,
                               font_size=9, color_hex=dim_hex)
            elif kind == "alinea":
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.first_line_indent = Pt(-14)
                add_run_styled(p, linha, bold=False, italic=False,
                               font_size=9, color_hex=dim_hex)
            else:
                p.paragraph_format.left_indent = Pt(4)
                p.paragraph_format.first_line_indent = Pt(0)
                add_run_styled(p, linha, bold=False, italic=False,
                               font_size=9, color_hex=dim_hex)

    # ── Notas de Reunião ──────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    t_notas = doc.add_table(rows=2, cols=1)
    set_table_width(t_notas, pct=5000)
    cell_header(t_notas.cell(0, 0),
                "NOTAS DE REUNIÃO",
                COR["notas_header"])
    # Célula com espaço em branco para preenchimento manual
    c_notas = t_notas.cell(1, 0)
    set_cell_bg(c_notas, COR["notas_body"])
    set_cell_borders(c_notas)
    notas_text = art["notas"] if art["notas"] else "(espaço para anotações da reunião)"
    p_notas = c_notas.paragraphs[0]
    p_notas.paragraph_format.space_before = Pt(6)
    p_notas.paragraph_format.space_after = Pt(40)  # altura mínima para escrita
    p_notas.paragraph_format.left_indent = Pt(4)
    add_run_styled(p_notas, notas_text, italic=True,
                   font_size=9.5, color_hex="888888" if not art["notas"] else "222222")

    # Separador entre artigos
    doc.add_paragraph()
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after = Pt(16)
    run_sep = p_sep.add_run("─" * 80)
    run_sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run_sep.font.size = Pt(8)


# ---------------------------------------------------------------------------
# CONFIGURAÇÃO DO DOCUMENTO
# ---------------------------------------------------------------------------

def configurar_documento(doc):
    """Define margens e estilos base do documento."""
    from docx.oxml.ns import nsmap
    section = doc.sections[0]
    section.page_width = Cm(29.7)   # A4 horizontal
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Estilo normal base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def criar_word(path):
    doc = Document()
    configurar_documento(doc)

    # Capa / título
    add_page_title(doc)

    # Legenda de cores
    add_legenda(doc)

    # Quebra de página antes dos artigos
    doc.add_page_break()

    # Um artigo por página (exceto o último)
    for i, art in enumerate(ARTIGOS):
        add_article_section(doc, art)
        if i < len(ARTIGOS) - 1:
            doc.add_page_break()

    doc.save(path)
    print(f"Word guardado: {path}")


if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    output = os.path.join(base, "comparativo_reuniao_exemplo.docx")
    criar_word(output)
    print("Concluído.")
