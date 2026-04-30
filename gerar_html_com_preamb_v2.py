#!/usr/bin/env python3
"""
gerar_html_com_preamb_v2.py

Gera HTML com preâmbulo integrado na versão completa original.

Estratégia:
1. Gera HTML completo com criar_html() original
2. Extrai preâmbulo (considerandos EN+PT)
3. Remove ▌ de todo o documento
4. Injeta temas do preâmbulo na sidebar (após "Artigos")
5. Adiciona secção de preâmbulo ao final do HTML (após últimos artigos)

Output: comparativo_reuniao_exemplo_preamb_teste_v2.html
"""

import os
import sys
import json
import re
from docx import Document

# Importar dados e funções do script principal
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from gerar_comparativo_reuniao import (
    ARTIGOS,
    extrair_glossario_pt,
    extrair_glossario_en,
    criar_html,
)

# ─────────────────────────────────────────────────────────────────────────────
# EXTRAÇÃO DO PREÂMBULO (IDÊNTICO À V1)
# ─────────────────────────────────────────────────────────────────────────────

TEMA_CONSIDERANDOS = {
    # Posição 1: Âmbito e Exclusões (PRIMEIRO na sidebar)
    "Âmbito e Exclusões": [19, 20, 21, 22, 23],
    # Posição 2: Motivos e Objetivos (logo após Âmbito e Exclusões)
    "Motivos e Objetivos": [1, 2, 7, 16, 17, 3, 4, 5],
    # Resto em ordem alfabética
    "Abrigos": [27, 28, 84],
    "Alojamento": [48],
    "Amarração": [51],
    "Autoridades Competentes": [31],
    "Cães de guarda de gado/pastoreio": [58],
    "Cães militares/polícia/aduaneiros": [57],
    "Competências de Execução": [79, 85],
    "Conformações Extremas e Genótipos": [41, 42, 77],
    "Consanguinidade": [43],
    "Contentores": [47, 52],
    "Formação": [36, 37, 38],
    "Híbridos": [44],
    "Lares de acolhimento temporário": [29, 30],
    "Lojas de Venda": [24, 25, 26],
    "Luz": [52],
    "Mutilações": [56],
    "Obrigação de informação sobre detenção responsável": [28, 60],
    "Países Terceiros": [73, 74, 75],
    "Práticas dolorosas": [34],
    "Princípios gerais de bem-estar animal": [13],
    "Proteção de Dados": [67, 68, 69, 70, 71, 72],
    "Publicidade": [63],
    "Rastreabilidade": [8, 9, 10, 11, 12, 14, 15, 18, 61, 62, 65],
    "Registo/Aprovação de Estabelecimentos": [35, 59],
    "Regras específicas de bem-estar animal": [24, 25, 45, 46],
    "Regras mais restritivas": [80, 81],
    "Relatórios Anuais": [32, 66, 82],
    "Reprodução": [49, 50, 53, 54],
    "Sanções": [83, 6],
    "Saúde": [33, 40, 73],
    "Sociabilização": [55, 76],
    "Treino": [64],
    "Visitas Médico-Veterinárias de aconselhamento de bem-estar": [33, 39, 78],
}


def extrair_preamb():
    """Extrai preâmbulo (considerandos) de ambos os documentos, removendo ▌."""
    doc_en = Document('11.12.2025 Regulamento cães e gatos-ocr - sem rasuras.docx')
    doc_pt = Document('pe00002.pt26.PB.aftermeeting 2.docx')

    considerandos_en = {}
    considerandos_pt = {}

    # Extrair EN
    for p in doc_en.paragraphs:
        if p.style.name == 'Considérant' and p.text.strip():
            txt = p.text.strip()
            # Ignorar linhas só com ▌
            if not re.match(r'^[▌\s]+$', txt):
                m = re.match(r'^\((\d+)\)', txt)
                if m:
                    num = int(m.group(1))
                    # REMOVER ▌ do texto
                    txt_limpo = txt.replace('▌', '').strip()
                    considerandos_en[num] = txt_limpo

    # Extrair PT
    for p in doc_pt.paragraphs:
        if p.style.name == 'Considérant' and p.text.strip():
            txt = p.text.strip()
            # Ignorar linhas só com ▌
            if not re.match(r'^[▌\s]+$', txt):
                m = re.match(r'^\((\d+)\)', txt)
                if m:
                    num = int(m.group(1))
                    # REMOVER ▌ do texto
                    txt_limpo = txt.replace('▌', '').strip()
                    considerandos_pt[num] = txt_limpo

    # Criar estrutura PREAMB (com duplicação por tema)
    preamb = []
    for tema, nums in TEMA_CONSIDERANDOS.items():
        for num in nums:
            preamb.append({
                'id': f'PREAMB-{num:02d}',
                'numero': num,
                'tema': tema,
                'regulamento': {
                    'texto': considerandos_en.get(num, ''),
                    'traducao': considerandos_pt.get(num, '')
                }
            })

    # Retornar organizado por tema para sidebar (preservando ordem de TEMA_CONSIDERANDOS)
    preamb_por_tema = {}
    # Inicializar dicionário na ordem correta dos temas
    for tema in TEMA_CONSIDERANDOS.keys():
        preamb_por_tema[tema] = []

    # Preencher considerandos em cada tema
    for item in preamb:
        tema = item['tema']
        preamb_por_tema[tema].append(item)

    return preamb, preamb_por_tema


def modificar_html_para_adicionar_preamb(html_original, preamb_por_tema, artigos):
    """
    Modifica HTML gerado pela criar_html() original para:
    1. Remover ▌ restantes
    2. Adicionar temas do preâmbulo à sidebar (após "Artigos")
    3. Injetar dados do preâmbulo no JavaScript
    """

    # 1. REMOVER ▌ (cleanup extra, como precaução)
    html_limpo = html_original.replace('▌', '')

    # 2. ADICIONAR TEMAS DO PREÂMBULO À SIDEBAR
    # Padrão: após </nav> há um comentário ou fim
    # Vamos injetar JavaScript que adiciona temas na sidebar

    temas_preamb = list(preamb_por_tema.keys())  # Manter ordem do dicionário (Âmbito e Exclusões primeiro)
    dados_preamb_json = json.dumps(preamb_por_tema, ensure_ascii=False)

    # Injetar dados do preâmbulo no script
    script_inject = f"""
// ══════════════════════════════════════════════════════════════════════════════
// PREÂMBULO — Dados e Funcionalidade Adicional
// ══════════════════════════════════════════════════════════════════════════════

const PREAMB_POR_TEMA = {dados_preamb_json};
const TEMAS_PREAMB = {json.dumps(temas_preamb, ensure_ascii=False)};
let currentPreambuloSearchResults = null; // Guardar resultados da pesquisa

function highlightPreambuloText(text) {{
  // Aplicar highlight igual aos artigos, usando a variável searchTerm global
  if (!searchTerm || !text) return text;
  const re = new RegExp(searchTerm.replace(/[.*+?^${{}}()|[\\]\\\\]/g, '\\\\$&'), 'gi');
  return text.replace(re, m => `<mark>${{m}}</mark>`);
}}

function exibirTemaPreambulo(tema) {{
  const considerandos = PREAMB_POR_TEMA[tema];
  const container = document.getElementById('main-content');
  container.innerHTML = '';

  // Título do tema
  const titulo = document.createElement('div');
  titulo.className = 'preamb-titulo';
  titulo.innerHTML = '<h2>' + tema + '</h2><small>' + considerandos.length + ' considerando(s)</small>';
  container.appendChild(titulo);

  // Exibir cada considerando
  for (const cons of considerandos) {{
    const artBadge = document.createElement('div');
    artBadge.className = 'art-badge preamb';
    artBadge.textContent = 'PREAMB-' + String(cons.numero).padStart(2, '0');
    container.appendChild(artBadge);

    // Card EN
    const cardEn = document.createElement('div');
    cardEn.className = 'card reg';
    cardEn.style.marginBottom = '14px';
    cardEn.innerHTML = `
      <div class="card-header">
        @regulamento — Considerando ${{cons.numero}} (EN)
        <span class="card-header-ref">Texto Original EN</span>
      </div>
      <div class="card-body">${{highlightPreambuloText(cons.regulamento.texto).replace(/\\n/g, '<br>')}}</div>
    `;
    container.appendChild(cardEn);

    // Card PT
    const cardPt = document.createElement('div');
    cardPt.className = 'card reg-tr';
    cardPt.style.marginBottom = '20px';
    cardPt.innerHTML = `
      <div class="card-header">
        @regulamento — Considerando ${{cons.numero}} (PT)
        <span class="card-header-ref">Tradução PT-PT</span>
      </div>
      <div class="card-body">${{highlightPreambuloText(cons.regulamento.traducao).replace(/\\n/g, '<br>')}}</div>
    `;
    container.appendChild(cardPt);
  }}
}}

// Função para procurar em considerandos
function considerandoMatch(cons, searchTerm) {{
  const q = searchTerm.toLowerCase();
  return cons.numero.toString().includes(q) ||
         cons.tema.toLowerCase().includes(q) ||
         cons.regulamento.texto.toLowerCase().includes(q) ||
         cons.regulamento.traducao.toLowerCase().includes(q);
}}

// Função para restaurar resultados da pesquisa do preâmbulo na sidebar
function restorePreambuloSearchResults() {{
  if (!currentPreambuloSearchResults || currentPreambuloSearchResults.length === 0) return;

  const nav = document.getElementById('sidebar');
  if (!nav) return;

  // Remover seção anterior se existir
  const existing = nav.querySelector('[data-preamb-search-results]');
  if (existing) existing.remove();

  // Recriar seção de resultados do preâmbulo
  const prembSection = document.createElement('div');
  prembSection.setAttribute('data-preamb-search-results', 'true');
  prembSection.style.marginTop = '1rem';
  prembSection.style.borderTop = '1px solid rgba(255,255,255,0.2)';
  prembSection.style.paddingTop = '1rem';

  const prembLabel = document.createElement('p');
  prembLabel.textContent = 'Preâmbulo (' + currentPreambuloSearchResults.length + ')';
  prembLabel.style.color = '#9B8B9E';
  prembLabel.style.fontSize = '0.9rem';
  prembLabel.style.fontWeight = 'bold';
  prembLabel.style.marginBottom = '0.5rem';
  prembSection.appendChild(prembLabel);

  // Recriar botões
  currentPreambuloSearchResults.forEach(function(item) {{
    const btn = document.createElement('button');
    btn.className = 'preamb-search-btn';
    btn.innerHTML = 'PREAMB-' + String(item.cons.numero).padStart(2, '0') + ' — ' + item.tema +
                   '<small>' + item.cons.regulamento.traducao.substring(0, 50) + '…</small>';
    btn.style.width = '100%';
    btn.style.background = 'rgba(155, 139, 158, 0.15)';
    btn.style.border = 'none';
    btn.style.borderLeft = '3px solid #9B8B9E';
    btn.style.color = 'white';
    btn.style.padding = '0.75rem 1rem';
    btn.style.textAlign = 'left';
    btn.style.cursor = 'pointer';
    btn.style.fontSize = '0.85rem';
    btn.style.marginBottom = '0.5rem';
    btn.style.transition = 'all 0.2s';
    btn.onmouseover = function() {{ this.style.background = 'rgba(155, 139, 158, 0.25)'; }};
    btn.onmouseout = function() {{ this.style.background = 'rgba(155, 139, 158, 0.15)'; }};
    btn.onclick = function() {{
      exibirTemaPreambulo(item.tema);
    }};

    prembSection.appendChild(btn);
  }});

  nav.appendChild(prembSection);
}}

// Hook na pesquisa original para incluir preâmbulo na sidebar
function setupPreambuloSearch() {{
  if (typeof window.pesquisar !== 'function') {{
    console.log('pesquisar not ready, retrying...');
    setTimeout(setupPreambuloSearch, 100);
    return;
  }}

  const pesquisarOriginal = window.pesquisar;
  const renderOriginal = typeof window.render !== 'undefined' ? window.render : null;
  console.log('Setting up preamble search hook');

  window.pesquisar = function(q) {{
    console.log('Pesquisando:', q);

    // Chamar pesquisa original PRIMEIRO
    pesquisarOriginal.call(this, q);

    const searchTerm = q.trim().toLowerCase();
    if (!searchTerm) {{
      currentPreambuloSearchResults = null;
      return;
    }}

    // Agora procurar no preâmbulo
    console.log('Procurando no preâmbulo por:', searchTerm);
    const matchingConsiderandos = [];

    for (const tema of TEMAS_PREAMB) {{
      const considerandos = PREAMB_POR_TEMA[tema];
      for (const cons of considerandos) {{
        if (considerandoMatch(cons, searchTerm)) {{
          console.log('Match encontrado:', cons.numero, tema);
          matchingConsiderandos.push({{ tema: tema, cons: cons }});
        }}
      }}
    }}

    console.log('Total de matches no preâmbulo:', matchingConsiderandos.length);

    // Guardar resultados para restaurar depois
    if (matchingConsiderandos.length > 0) {{
      currentPreambuloSearchResults = matchingConsiderandos;
      // Restaurar imediatamente
      restorePreambuloSearchResults();
    }} else {{
      currentPreambuloSearchResults = null;
    }}
  }};

  // Hook na função render() original para restaurar resultados após clicar em artigo
  if (renderOriginal && typeof window.render === 'function') {{
    window.render = function() {{
      renderOriginal.call(this);
      // Restaurar resultados do preâmbulo APÓS render
      setTimeout(restorePreambuloSearchResults, 50);
    }};
  }}
}}

setupPreambuloSearch();

// Hook na função limparPesquisa() para também limpar preâmbulo
setTimeout(function() {{
  if (typeof window.limparPesquisa === 'function') {{
    const limparOriginal = window.limparPesquisa;
    window.limparPesquisa = function() {{
      currentPreambuloSearchResults = null;
      limparOriginal.call(this);
    }};
  }}
}}, 100);

// Adicionar botões de preâmbulo à sidebar DEPOIS de renderSidebar() ser chamada
setTimeout(function() {{
  const nav = document.getElementById('sidebar');
  if (!nav) return;

  // Adicionar separador de preâmbulo
  const sep = document.createElement('h2');
  sep.style.marginTop = '1rem';
  sep.style.borderTop = '1px solid rgba(255,255,255,0.2)';
  sep.style.paddingTop = '1rem';
  sep.textContent = 'Preâmbulo';
  nav.appendChild(sep);

  // Adicionar botões de temas
  for (const tema of TEMAS_PREAMB) {{
    const btn = document.createElement('button');
    btn.className = 'preamb-theme-btn';
    btn.innerHTML = tema + '<small>(' + PREAMB_POR_TEMA[tema].length + ' considerando)</small>';
    btn.onclick = () => exibirTemaPreambulo(tema);
    nav.appendChild(btn);
  }}
}}, 100);

// Estilos para preâmbulo
const stylePreambulo = document.createElement('style');
stylePreambulo.textContent = `
  .preamb-theme-btn {{
    width: 100%;
    background: rgba(155, 139, 158, 0.1) !important;
    border: none;
    border-left: 3px solid #9B8B9E;
    color: white;
    padding: 0.75rem 1rem;
    text-align: left;
    cursor: pointer;
    font-size: 0.9rem;
    transition: all 0.2s;
  }}
  .preamb-theme-btn:hover {{
    background: rgba(155, 139, 158, 0.2) !important;
  }}
  .preamb-theme-btn small {{
    display: block;
    font-size: 0.75rem;
    opacity: 0.7;
    margin-top: 2px;
  }}
  .preamb-titulo {{
    margin-bottom: 1.5rem;
    padding-bottom: 1rem;
    border-bottom: 2px solid #9B8B9E;
  }}
  .preamb-titulo h2 {{
    color: #9B8B9E;
    font-size: 1.3rem;
    margin-bottom: 0.25rem;
  }}
  .preamb-titulo small {{
    color: #999;
    font-size: 0.9rem;
  }}
  .art-badge.preamb {{
    background: #9B8B9E !important;
  }}
  .card.preamb {{
    border-left-color: #9B8B9E !important;
    background: #F8F5FB !important;
  }}
  .card.preamb .card-header {{
    color: #9B8B9E !important;
  }}
`;
document.head.appendChild(stylePreambulo);
"""

    # Localizar onde injetar (antes de fechar </body>)
    if '</body>' in html_limpo:
        # Injetar antes de </body>
        html_limpo = html_limpo.replace('</body>', f'<script>\n{script_inject}\n</script>\n</body>')

    return html_limpo


def gerar_html_completo_com_preamb(output_path):
    """Gera HTML completo com preâmbulo integrado."""

    print("1. Gerando HTML base (artigos completos)...")
    # Gerar HTML original completo
    temp_path = '/tmp/comparativo_temp.html'
    criar_html(temp_path, ARTIGOS)

    # Ler HTML gerado
    with open(temp_path, 'r', encoding='utf-8') as f:
        html_original = f.read()

    print("2. Extraindo preâmbulo...")
    preamb, preamb_por_tema = extrair_preamb()
    print(f"   ✓ {len(preamb)} considerandos extraídos ({len(preamb_por_tema)} temas)")

    print("3. Modificando HTML: removendo ▌ e adicionando preâmbulo...")
    html_final = modificar_html_para_adicionar_preamb(html_original, preamb_por_tema, ARTIGOS)

    print(f"4. Gravando {output_path}...")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_final)

    print(f"✓ Concluído: {output_path}")
    import os
    size_mb = os.path.getsize(output_path) / (1024 * 1024)
    print(f"  Tamanho: {size_mb:.2f} MB")


if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    output = os.path.join(base, "comparativo_reuniao_exemplo_preamb_teste_v2.html")
    gerar_html_completo_com_preamb(output)
