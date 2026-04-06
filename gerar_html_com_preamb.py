#!/usr/bin/env python3
"""
Script para gerar HTML com preâmbulo adicional.
Baseado em gerar_comparativo_reuniao.py mas com suporte para preâmbulo.

NOTA: Este é um ficheiro de teste. O objetivo é validar a integração do
preâmbulo antes de mesclar com a versão principal.
"""

import os
import sys
import json
from docx import Document
import re

# Importar dados do script principal
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from gerar_comparativo_reuniao import (
    ARTIGOS,
    extrair_glossario_pt,
    extrair_glossario_en,
    COR,
    criar_excel,
)

# ─────────────────────────────────────────────────────────────────────────────
# EXTRAÇÃO DO PREÂMBULO
# ─────────────────────────────────────────────────────────────────────────────

TEMA_CONSIDERANDOS = {
    "Abrigos": [27, 28, 84],
    "Alojamento": [48],
    "Amarração": [51],
    "Âmbito e Exclusões": [19, 20, 21, 22, 23],
    "Autoridades Competentes": [31],
    "Cães de guarda de gado/pastoreio": [58],
    "Cães militares/polícia/aduaneiros": [57],
    "Comunicação dos incumprimentos": [6],
    "Competências de Execução": [79, 85],
    "Conformações Extremas e Genótipos": [41, 42, 77],
    "Consanguinidade": [43],
    "Contentores": [47, 52],
    "Formação": [36, 37, 38],
    "Híbridos": [44],
    "Lares de acolhimento temporário": [29, 30],
    "Lojas de Venda": [24, 25, 26],
    "Luz": [52],
    "Motivos": [1, 2, 7],
    "Mutilações": [56],
    "Objetivos": [16, 17],
    "Objetivos (proteção ao mercado, ao consumidor e ao bem-estar animal)": [3, 4, 5],
    "Obrigação de informação sobre detenção responsável": [28, 60],
    "Países Terceiros": [73, 74, 75],
    "Práticas dolorosas": [34],
    "Princípios gerais de bem-estar animal": [13],
    "Proteção de Dados": [67, 68, 69, 70, 71, 72],
    "Rastreabilidade": [8, 9, 10, 11, 12, 14, 15, 18, 61, 62, 65],
    "Registo/Aprovação de Estabelecimentos": [35, 59],
    "Regras específicas de bem-estar animal": [24, 25, 45, 46],
    "Regras mais restritivas": [80, 81],
    "Relatórios Anuais": [32, 66, 82],
    "Reprodução": [49, 50, 53, 54],
    "Requisitos relativos à publicidade": [63],
    "Sanções": [83],
    "Saúde": [33, 40, 73],
    "Sociabilização": [55, 76],
    "Treino": [64],
    "Visitas Médico-Veterinárias de aconselhamento de bem-estar": [33, 39, 78],
}


def extrair_preamb():
    """Extrai preâmbulo (considerandos) de ambos os documentos."""
    doc_en = Document('11.12.2025 Regulamento cães e gatos-ocr - sem rasuras.docx')
    doc_pt = Document('pe00002.pt26.PB.aftermeeting 2.docx')

    considerandos_en = {}
    considerandos_pt = {}

    # Extrair EN
    for p in doc_en.paragraphs:
        if p.style.name == 'Considérant' and p.text.strip():
            txt = p.text.strip()
            if not re.match(r'^[▌\s]+$', txt):
                m = re.match(r'^\((\d+)\)', txt)
                if m:
                    num = int(m.group(1))
                    considerandos_en[num] = txt

    # Extrair PT
    for p in doc_pt.paragraphs:
        if p.style.name == 'Considérant' and p.text.strip():
            txt = p.text.strip()
            if not re.match(r'^[▌\s]+$', txt):
                m = re.match(r'^\((\d+)\)', txt)
                if m:
                    num = int(m.group(1))
                    considerandos_pt[num] = txt

    # Criar estrutura PREAMB (com duplicação por tema)
    preamb = []
    for tema, nums in TEMA_CONSIDERANDOS.items():
        for num in nums:
            preamb.append({
                'id': f'PREAMB-{num:02d}',
                'numero': num,
                'tema': tema,
                'regulamento': {
                    'texto': considerandos_en.get(num, ''),
                    'traducao': considerandos_pt.get(num, '')
                }
            })

    # Retornar em ordem numérica para apresentação
    return sorted(preamb, key=lambda x: x['numero'])


def criar_html_com_preamb(path, artigos, preamb):
    """Gera HTML com artigos + preâmbulo ao final."""
    print(f"Gerando HTML com {len(artigos)} artigos + {len(preamb)} considerandos...")

    # Extrair glossários
    glossario_pt = extrair_glossario_pt(artigos)
    glossario_en = extrair_glossario_en(artigos)
    glossario_pt_json = json.dumps(glossario_pt, ensure_ascii=False)
    glossario_en_json = json.dumps(glossario_en, ensure_ascii=False)

    # Dados combinados: ARTIGOS + PREÂMBULO
    # Estrutura: { artigos: [...], preamb: {...temas: [...considerandos...]...} }
    dados_preamb_por_tema = {}
    for item in preamb:
        tema = item['tema']
        if tema not in dados_preamb_por_tema:
            dados_preamb_por_tema[tema] = []
        dados_preamb_por_tema[tema].append(item)

    dados_combinados = {
        'artigos': artigos,
        'preamb': dados_preamb_por_tema,
        'temas_preamb': sorted(dados_preamb_por_tema.keys())
    }
    dados_json = json.dumps(dados_combinados, ensure_ascii=False, indent=2)

    # JavaScript para glossário (igual ao original)
    js_glossario = """
// ── GLOSSÁRIO INTERATIVO: Tooltips e Marcação de Termos (APENAS REGULAMENTO) ──
let tooltipTimeoutId = null;

function marcarGlossario(htmlStr, glossario) {
  if (!glossario || Object.keys(glossario).length === 0) return htmlStr;

  // ★ CRÍTICO: Normalizar hífens no HTML ANTES de qualquer matching
  let htmlNormalizado = htmlStr
    .replace(/\\u2011/g, '-')
    .replace(/\\u2010/g, '-');

  const termos = Object.keys(glossario).sort((a, b) => b.length - a.length);
  let result = htmlNormalizado;
  let debugCount = 0;

  for (const termo of termos) {
    let isSubstring = false;
    for (const other of termos) {
      if (termo !== other && other.includes(termo)) {
        if (result.includes('data-termo="' + other + '"')) {
          isSubstring = true;
          break;
        }
      }
    }
    if (isSubstring) continue;

    const regexEscaped = termo.replace(/[.*+?^${}()|[\\\\]\\\\\\\\]/g, '\\\\\\\\$&');
    const regex = new RegExp('\\\\\\\\b' + regexEscaped + '\\\\\\\\b(?![^<]*>)', 'gi');

    result = result.replace(regex, m => {
      debugCount++;
      return '<span class="glossario-termo" data-termo="' + termo + '">' + m + '</span>';
    });
  }
  return result;
}

function mostrarTooltip(termo, elemento, glossario) {
  if (!elemento || !elemento.classList) return;
  const tooltip = document.getElementById('tooltip-glossario');
  const definicao = glossario[termo.toLowerCase()];
  if (!definicao) return;
  tooltip.textContent = definicao;
  tooltip.classList.add('visible');
  const rect = elemento.getBoundingClientRect();
  let top = rect.bottom + 8;
  let left = rect.left;
  tooltip.style.left = Math.max(8, Math.min(left, window.innerWidth - 328)) + 'px';
  tooltip.style.top = top + 'px';
  clearTimeout(tooltipTimeoutId);
  tooltipTimeoutId = setTimeout(() => esconderTooltip(), 6000);
}

function esconderTooltip() {
  const tooltip = document.getElementById('tooltip-glossario');
  tooltip.classList.remove('visible');
}

// Event delegation para glossário
document.addEventListener('mouseenter', (e) => {
  if (e.target && e.target.classList && e.target.classList.contains('glossario-termo')) {
    const glossario = GLOSSARIO_EN;
    if (glossario) mostrarTooltip(e.target.dataset.termo, e.target, glossario);
  }
}, true);
"""

    html = f"""<!DOCTYPE html>
<html lang="pt">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Comparativo Artigo a Artigo + Preâmbulo — Regulamento 2023/0447</title>
<style>
  :root {{
    --reg: #8AAFCF;
    --reg-bg: #F0F7FC;
    --reg-tr: #F7FAFD;
    --rgb: #7FAA8C;
    --rgb-bg: #F0F7F4;
    --cod: #B8956A;
    --cod-bg: #FFF9F2;
    --leg: #75A6AE;
    --leg-bg: #F5FCFD;
    --div: #A689C6;
    --div-bg: #F9F6FD;
    --nota: #909090;
    --nota-bg: #FAFAFA;
    --dark: #2C3E50;
    --preamb: #9B8B9E;
    --preamb-bg: #F8F5FB;
  }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Segoe UI', Calibri, sans-serif; background: #F8F9FB; color: #333; }}

  /* HEADER */
  header {{
    background: var(--dark);
    color: white;
    padding: 1.5rem;
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
  }}
  header h1 {{ font-size: 1.3rem; margin-bottom: 0.5rem; }}
  header p {{ opacity: 0.9; font-size: 0.95rem; }}

  /* MAIN LAYOUT */
  .container {{
    display: flex;
    min-height: calc(100vh - 80px);
  }}

  /* SIDEBAR */
  nav {{
    width: 280px;
    background: var(--dark);
    color: white;
    padding: 1.5rem 0;
    overflow-y: auto;
    border-right: 1px solid rgba(255,255,255,0.1);
  }}
  nav h2 {{
    color: white;
    font-size: 0.75rem;
    text-transform: uppercase;
    padding: 0.75rem 1rem;
    margin-bottom: 0.5rem;
    border-bottom: 1px solid rgba(255,255,255,0.1);
  }}
  nav button {{
    width: 100%;
    background: rgba(255,255,255,0.05);
    border: none;
    color: white;
    padding: 0.75rem 1rem;
    text-align: left;
    cursor: pointer;
    font-size: 0.9rem;
    transition: all 0.2s;
  }}
  nav button:hover {{ background: rgba(255,255,255,0.1); }}
  nav button.active {{
    background: rgba(255,255,255,0.2);
    border-left: 3px solid var(--reg);
    padding-left: calc(1rem - 3px);
  }}
  nav button small {{
    display: block;
    font-size: 0.75rem;
    opacity: 0.7;
    margin-top: 2px;
  }}
  nav button.theme-button {{
    background: rgba(155,139,158,0.1);
    border-left: 3px solid var(--preamb);
  }}
  nav button.theme-button:hover {{
    background: rgba(155,139,158,0.2);
  }}

  /* MAIN CONTENT */
  #main-content {{
    flex: 1;
    padding: 2rem;
    overflow-y: auto;
  }}

  /* CARDS */
  .card {{
    background: white;
    border-left: 4px solid #ccc;
    padding: 1.5rem;
    margin-bottom: 1.5rem;
    border-radius: 4px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.08);
  }}
  .card.reg {{ border-left-color: var(--reg); }}
  .card.reg-tr {{ border-left-color: var(--reg); background: var(--reg-tr); }}
  .card.preamb {{ border-left-color: var(--preamb); background: var(--preamb-bg); }}

  .card-header {{
    font-weight: bold;
    margin-bottom: 1rem;
    font-size: 1rem;
    padding-bottom: 0.5rem;
    border-bottom: 1px solid rgba(0,0,0,0.08);
  }}
  .card.reg .card-header {{ color: var(--reg); }}
  .card.preamb .card-header {{ color: var(--preamb); }}

  .card-ref {{
    font-size: 0.85rem;
    opacity: 0.7;
    margin-top: 0.5rem;
  }}

  .card-body {{
    line-height: 1.6;
    font-size: 0.95rem;
    color: #444;
  }}

  .tema-title {{
    font-size: 1.2rem;
    font-weight: bold;
    margin-bottom: 1.5rem;
    color: var(--preamb);
  }}

  /* BADGE */
  .art-badge {{
    display: inline-block;
    background: var(--dark);
    color: white;
    padding: 0.4rem 0.8rem;
    border-radius: 3px;
    font-size: 0.85rem;
    font-weight: bold;
    margin-bottom: 1rem;
  }}
  .art-badge.preamb {{
    background: var(--preamb);
  }}

  /* GLOSSÁRIO */
  #tooltip-glossario {{
    position: fixed;
    background: #222;
    color: white;
    padding: 0.75rem;
    border-radius: 4px;
    font-size: 0.9rem;
    max-width: 300px;
    z-index: 1000;
    opacity: 0;
    pointer-events: none;
    transition: opacity 0.2s;
  }}
  #tooltip-glossario.visible {{ opacity: 0.95; pointer-events: auto; }}

  .glossario-termo {{
    text-decoration: underline;
    text-decoration-style: dashed;
    cursor: help;
    color: var(--reg);
  }}

  /* FOOTER NAVIGATION */
  .nav-btns {{
    position: fixed;
    bottom: 0;
    right: 0;
    padding: 1rem;
    display: flex;
    gap: 0.5rem;
  }}
  .nav-btns button {{
    background: var(--dark);
    color: white;
    border: none;
    padding: 0.5rem 1rem;
    cursor: pointer;
    border-radius: 4px;
    font-size: 0.9rem;
  }}
  .nav-btns button:hover {{ opacity: 0.8; }}
  .nav-btns button:disabled {{ opacity: 0.5; cursor: not-allowed; }}
</style>
</head>
<body>

<header>
  <h1>Regulamento (UE) 2023/0447 — Bem-estar de Cães e Gatos</h1>
  <p>Visualizador Interativo: Artigos + Preâmbulo | Comparação EN ↔ PT</p>
</header>

<div class="container">
  <nav id="sidebar">
    <h2>Artigos</h2>
  </nav>
  <div id="main-content"></div>
</div>

<div id="tooltip-glossario"></div>
<div class="nav-btns">
  <button id="prev-btn">← Anterior</button>
  <button id="next-btn">Próximo →</button>
</div>

<script>
// Dados injetados
const DADOS = {dados_json};
const GLOSSARIO_EN = {glossario_en_json};
const GLOSSARIO_PT = {glossario_pt_json};

{js_glossario}

// NAVEGAÇÃO E RENDERIZAÇÃO
let currentIndex = 0;
let currentType = 'artigo'; // 'artigo' ou 'preamb'
let currentTema = null;
let allItems = [];

// Construir lista de itens navegáveis
function construirListaItens() {{
  allItems = [];

  // Adicionar artigos
  for (const art of DADOS.artigos) {{
    allItems.push({{ type: 'artigo', data: art }});
  }}

  // Adicionar preâmbulo (por tema + considerandos)
  for (const tema of DADOS.temas_preamb) {{
    const considerandos = DADOS.preamb[tema];
    for (const cons of considerandos) {{
      allItems.push({{ type: 'preamb', tema: tema, data: cons }});
    }}
  }}
}}

function renderSidebar() {{
  const nav = document.getElementById('sidebar');
  nav.innerHTML = '<h2>Artigos</h2>';

  // Botões de artigos
  for (const art of DADOS.artigos) {{
    const btn = document.createElement('button');
    btn.textContent = art.id + ' — ' + art.tema;
    btn.onclick = () => navegar('artigo', art.id);
    if (currentType === 'artigo' && currentIndex === DADOS.artigos.indexOf(art)) {{
      btn.classList.add('active');
    }}
    nav.appendChild(btn);
  }}

  // Separador
  const sep = document.createElement('h2');
  sep.style.marginTop = '1rem';
  sep.textContent = 'Preâmbulo';
  nav.appendChild(sep);

  // Botões de temas do preâmbulo
  for (const tema of DADOS.temas_preamb) {{
    const btn = document.createElement('button');
    btn.classList.add('theme-button');
    btn.innerHTML = tema + '<small>' + DADOS.preamb[tema].length + ' considerando(s)</small>';
    btn.onclick = () => exibirTema(tema);
    if (currentType === 'preamb' && currentTema === tema) {{
      btn.classList.add('active');
    }}
    nav.appendChild(btn);
  }}
}}

function exibirTema(tema) {{
  currentType = 'preamb';
  currentTema = tema;
  renderSidebar();

  const considerandos = DADOS.preamb[tema];
  const content = document.getElementById('main-content');

  content.innerHTML = '<div class="tema-title">' + tema + '</div>';

  for (const cons of considerandos) {{
    const enHtml = marcarGlossario(cons.regulamento.texto.replace(/\\n/g, '<br>'), GLOSSARIO_EN);
    const ptHtml = marcarGlossario(cons.regulamento.traducao.replace(/\\n/g, '<br>'), GLOSSARIO_PT);

    content.innerHTML += `
      <div class="art-badge preamb">PREAMB-${{cons.numero.toString().padStart(2, '0')}}</div>
      <div class="card preamb" style="margin-bottom: 14px;">
        <div class="card-header">
          Regulamento — Texto Original EN
          <span class="card-ref">Considerando ${{cons.numero}}</span>
        </div>
        <div class="card-body">${{enHtml}}</div>
      </div>
      <div class="card preamb" style="margin-bottom: 20px;">
        <div class="card-header">
          Regulamento — Tradução PT-PT
          <span class="card-ref">Considerando ${{cons.numero}}</span>
        </div>
        <div class="card-body">${{ptHtml}}</div>
      </div>
    `;
  }}
}}

function navegar(tipo, id) {{
  if (tipo === 'artigo') {{
    const idx = DADOS.artigos.findIndex(a => a.id === id);
    if (idx >= 0) {{
      currentType = 'artigo';
      currentIndex = idx;
      renderArtigo(DADOS.artigos[idx]);
      renderSidebar();
    }}
  }}
}}

function renderArtigo(art) {{
  const content = document.getElementById('main-content');

  const enHtml = marcarGlossario(art.regulamento.texto.replace(/\\n/g, '<br>'), GLOSSARIO_EN);
  const ptHtml = marcarGlossario(art.regulamento.traducao.replace(/\\n/g, '<br>'), GLOSSARIO_PT);

  content.innerHTML = `
    <div class="art-badge">${{art.id}}</div>
    <div class="tema-title">${{art.tema}}</div>

    <div class="card reg" style="margin-bottom: 14px;">
      <div class="card-header">
        Regulamento — ${{art.regulamento.titulo}}
        <span class="card-ref">${{art.regulamento.ref}} · Texto Original EN</span>
      </div>
      <div class="card-body">${{enHtml}}</div>
    </div>

    <div class="card reg-tr" style="margin-bottom: 20px;">
      <div class="card-header">
        Regulamento — Tradução PT-PT
        <span class="card-ref">${{art.regulamento.ref}}</span>
      </div>
      <div class="card-body">${{ptHtml}}</div>
    </div>
  `;
}}

// Inicializar
construirListaItens();
renderSidebar();
if (DADOS.artigos.length > 0) {{
  renderArtigo(DADOS.artigos[0]);
}}

// Event listeners para navegação
document.getElementById('prev-btn').onclick = () => {{
  if (currentIndex > 0 && currentType === 'artigo') {{
    currentIndex--;
    renderArtigo(DADOS.artigos[currentIndex]);
    renderSidebar();
  }}
}};

document.getElementById('next-btn').onclick = () => {{
  if (currentIndex < DADOS.artigos.length - 1 && currentType === 'artigo') {{
    currentIndex++;
    renderArtigo(DADOS.artigos[currentIndex]);
    renderSidebar();
  }}
}};

// Atualizar botões de navegação
setInterval(() => {{
  document.getElementById('prev-btn').disabled = (currentIndex === 0 || currentType !== 'artigo');
  document.getElementById('next-btn').disabled = (currentIndex === DADOS.artigos.length - 1 || currentType !== 'artigo');
}}, 100);
</script>

</body>
</html>
"""

    with open(path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"✓ HTML gerado: {path}")


if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))

    # Extrair preâmbulo
    print("Extraindo preâmbulo...")
    preamb = extrair_preamb()
    print(f"✓ Preâmbulo extraído: {len(preamb)} considerandos")

    # Gerar HTML
    output_path = os.path.join(base, "comparativo_reuniao_exemplo_preamb_teste.html")
    criar_html_com_preamb(output_path, ARTIGOS, preamb)

    print("✓ Concluído com sucesso!")
