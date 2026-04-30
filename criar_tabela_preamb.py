#!/usr/bin/env python3
"""
Cria documento Word com tabela de temas do preâmbulo
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Adiciona cor de fundo a célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

TEMA_CONSIDERANDOS = {
    "Âmbito e Exclusões": [19, 20, 21, 22, 23],
    "Motivos e Objetivos": [1, 2, 7, 16, 17, 3, 4, 5],
    "Abrigos": [27, 28, 84],
    "Alojamento": [48],
    "Amarração": [51],
    "Autoridades Competentes": [31],
    "Cães de guarda de gado/pastoreio": [58],
    "Cães militares/polícia/aduaneiros": [57],
    "Competências de Execução": [79, 85],
    "Conformações Extremas e Genótipos": [41, 42, 77],
    "Consanguinidade": [43],
    "Contentores": [47, 52],
    "Formação": [36, 37, 38],
    "Híbridos": [44],
    "Lares de acolhimento temporário": [29, 30],
    "Lojas de Venda": [24, 25, 26],
    "Luz": [52],
    "Mutilações": [56],
    "Obrigação de informação sobre detenção responsável": [28, 60],
    "Países Terceiros": [73, 74, 75],
    "Práticas dolorosas": [34],
    "Princípios gerais de bem-estar animal": [13],
    "Proteção de Dados": [67, 68, 69, 70, 71, 72],
    "Publicidade": [63],
    "Rastreabilidade": [8, 9, 10, 11, 12, 14, 15, 18, 61, 62, 65],
    "Registo/Aprovação de Estabelecimentos": [35, 59],
    "Regras específicas de bem-estar animal": [24, 25, 45, 46],
    "Regras mais restritivas": [80, 81],
    "Relatórios Anuais": [32, 66, 82],
    "Reprodução": [49, 50, 53, 54],
    "Sanções": [83, 6],
    "Saúde": [33, 40, 73],
    "Sociabilização": [55, 76],
    "Treino": [64],
    "Visitas Médico-Veterinárias de aconselhamento de bem-estar": [33, 39, 78],
}

# Criar documento
doc = Document()

# Adicionar título
title = doc.add_heading('Preâmbulo — Divisão por Temas', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adicionar subtítulo
subtitle = doc.add_paragraph('Regulamento (UE) 2023/0447 — Considerandos organizados por tema')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(11)
subtitle_format.font.italic = True
subtitle_format.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # Espaço

# Criar tabela com 3 colunas
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'

# Cabeçalho
header_cells = table.rows[0].cells
header_cells[0].text = ''
header_cells[1].text = 'Tema'
header_cells[2].text = 'Considerandos'

# Formatar cabeçalho
for i, cell in enumerate(header_cells):
    shade_cell(cell, '4472C4')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adicionar linhas
for tema, considerandos in TEMA_CONSIDERANDOS.items():
    row_cells = table.add_row().cells
    row_cells[0].text = ''
    row_cells[1].text = tema
    row_cells[2].text = ', '.join(str(c) for c in considerandos)
    
    # Alinhar coluna 3 ao centro
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Ajustar largura das colunas
for row in table.rows:
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(3.2)
    row.cells[2].width = Inches(1.8)

# Salvar
filename = 'Preamb_Temas_Considerandos.docx'
doc.save(filename)
print(f"✅ Documento criado: {filename}")
print(f"   - {len(TEMA_CONSIDERANDOS)} temas")
print(f"   - Coluna 1: vazia")
print(f"   - Coluna 2: Tema")
print(f"   - Coluna 3: números dos considerandos")
