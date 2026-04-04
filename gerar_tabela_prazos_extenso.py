"""
Gera tabela Word EXTENSA de prazos de implementação do Regulamento 2023/0447 (cães e gatos).

Cada linha inclui dois blocos verbatim ipsis verbis:
  [NORMA]  — texto substantivo do artigo (o que a norma exige)
  [PRAZO]  — cláusula que fixa o prazo (de Art. 33 ou do próprio artigo)

Output: tabela_prazos_regulamento_extenso.docx

Correções face à versão resumida:
  - Removidas 2 linhas incorrectas com ref. a "Art. 7(3) atos delegados"
  - Art. 23(7) → Art. 23(4), 2.º parágrafo (Art. 23 tem apenas 4 §§)
  - Art. 26(5) → Art. 26(4), 5.º § (Art. 26 tem apenas 4 §§)
  - Art. 33(2) → Art. 33 (sem numeração de parágrafo)
"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Cores
# ---------------------------------------------------------------------------
C_HEADER_BG   = "1F4E79"
C_HEADER_TEXT = "FFFFFF"
C_GROUP_BG    = "D9E8F5"
C_GROUP_TEXT  = "1F4E79"
C_ROW_EVEN    = "EBF3FB"
C_ROW_ODD     = "FFFFFF"
C_CHECK       = "1F4E79"
C_REF         = "1F4E79"
C_NORMA       = "303030"   # cinzento quase preto para texto substantivo
C_PRAZO       = "1F4E79"   # azul para cláusula de prazo
C_VIA         = "808080"   # cinzento para "via Art. X"
C_BORDER      = "B8CCE4"

# ---------------------------------------------------------------------------
# Dimensões (landscape A4, margens 1.5 cm cada lado → 26.7 cm utilizáveis)
# ---------------------------------------------------------------------------
PAGE_W_CM    = 29.7
MARGIN_CM    = 1.5
USABLE_CM    = PAGE_W_CM - 2 * MARGIN_CM   # 26.7 cm
COL1_CM      = 14.0                         # mais larga para texto extenso
N_PRAZO_COLS = 12
PRAZO_CM     = (USABLE_CM - COL1_CM) / N_PRAZO_COLS   # ≈ 1.058 cm

# ---------------------------------------------------------------------------
# Colunas de prazo (cronológicas)
# ---------------------------------------------------------------------------
COLUNAS = [
    {"key": "2a",       "label": "2\nanos"},
    {"key": "3a",       "label": "3\nanos"},
    {"key": "4a",       "label": "4\nanos"},
    {"key": "5a",       "label": "5\nanos"},
    {"key": "6a",       "label": "6\nanos"},
    {"key": "7a",       "label": "7\nanos"},
    {"key": "8a",       "label": "8\nanos"},
    {"key": "10a",      "label": "10\nanos"},
    {"key": "14a",      "label": "14\nanos"},
    {"key": "15a",      "label": "15\nanos"},
    {"key": "1Jul2030", "label": "1 Jul.\n2030"},
    {"key": "1Jul2036", "label": "1 Jul.\n2036"},
]

# ---------------------------------------------------------------------------
# Dados: cada norma tem:
#   ref       — referência articulada (str; \n para sub-linhas "via ...")
#   norma     — texto verbatim ipsis verbis do artigo substantivo
#   prazo     — cláusula verbatim que fixa o prazo (None = embutido na norma)
#   prazo_src — fonte do prazo para label (ex.: "Art. 33, ponto (i)")
#   prazos    — lista de chaves das colunas a marcar com ✓
# ---------------------------------------------------------------------------

GRUPOS = [
    # =========================================================================
    {
        "titulo": "Grupo 1 — Bem-estar e Reprodução",
        "normas": [
            {
                "ref": "Art. 8(1)\n[via Art. 33, ponto (iii)]",
                "norma": (
                    "1. Operators of breeding establishments shall ensure that their breeding "
                    "strategies minimise the risk of producing dogs or cats with genotypes "
                    "associated with effects detrimental to the health and welfare of those animals."
                ),
                "prazo": (
                    "Article 8(1) shall apply from 1 July 2036."
                ),
                "prazo_src": "Art. 33, ponto (iii)",
                "prazos": ["1Jul2036"],
            },
            {
                "ref": "Art. 8(2)\n[via Art. 33, ponto (iii)]",
                "norma": (
                    "2. Operators of breeding establishments shall not use for reproduction dogs "
                    "or cats that have excessive conformational traits leading to a high risk of "
                    "detrimental effects on the welfare of such dogs or cats, or of their offspring. "
                    "Before selecting a dog or cat that might have an excessive conformational trait "
                    "for breeding, the operator shall consult a veterinarian or an independent "
                    "qualified person acting under the responsibility of a veterinarian. That "
                    "veterinarian or independent qualified person shall assess whether the dog or "
                    "cat has an excessive conformational trait."
                ),
                "prazo": (
                    "Article 8(2) shall apply from 1 July 2030."
                ),
                "prazo_src": "Art. 33, ponto (iii)",
                "prazos": ["1Jul2030"],
            },
            {
                "ref": "Art. 12(2)(3)\n[via Art. 33, ponto (v)]",
                "norma": (
                    "2. The competences referred to in paragraph 1 may be acquired through "
                    "education, training or professional experience. Only documented education, "
                    "training or professional experience shall be taken into account when "
                    "determining whether an animal carer has the competences referred to in "
                    "paragraph 1.\n"
                    "3. Operators shall ensure that at least one animal carer, other than a "
                    "volunteer or intern, at the establishment has completed the training courses "
                    "referred to in Article 22. Operators shall ensure that that animal carer "
                    "transfers his or her knowledge to the other animal carers of the establishment."
                ),
                "prazo": (
                    "Article 12(2) and (3) shall apply from … [seven years from the date of entry "
                    "into force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (v)",
                "prazos": ["7a"],
            },
            {
                "ref": "Art. 12(4)",
                "norma": (
                    "4. The Commission shall adopt implementing acts, laying down minimum "
                    "requirements concerning the formal education, training or professional "
                    "experience referred to in paragraph 2 of this Article necessary to determine "
                    "whether an animal carer has the competences referred to in paragraph 1 and "
                    "for the training courses referred to in paragraph 3. The implementing act "
                    "concerning the training courses referred to in paragraph 3 shall be adopted "
                    "by ... [3 years from the date of entry into force of this Regulation]. "
                    "Those implementing acts shall be adopted in accordance with the examination "
                    "procedure referred to in Article 29."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["3a"],
            },
            {
                "ref": "Art. 13(1)(a)",
                "norma": (
                    "1. Operators shall:\n"
                    "(a) ensure that the establishments for which they are responsible receive "
                    "a visit by a veterinarian for the purpose of identifying and assessing any "
                    "risk factor for the welfare of the dogs or cats and advising the operator "
                    "on measures to address those risks initially by ... [date three years after "
                    "the date of entry into force of this Regulation] or one year following the "
                    "notification of the new establishment, and thereafter when appropriate, "
                    "based on a risk analysis by the competent authorities, or on an annual "
                    "basis if Member States so provide in their national law;"
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["3a"],
            },
            {
                "ref": "Art. 13(2)",
                "norma": (
                    "2. By ... [date 24 months from the date of entry into force of this "
                    "Regulation], the Commission shall adopt delegated acts in accordance with "
                    "Article 28 supplementing this Article by laying down the minimum criteria "
                    "to be assessed by the veterinarian during the advisory welfare visit."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 16\n[via Art. 33, ponto (i)]",
                "norma": (
                    "1. Operators shall ensure that:\n"
                    "(a) dogs or cats for which they are responsible are inspected by animal "
                    "carers at least once a day and that vulnerable dogs and cats, such as "
                    "newborns, ill or injured dogs and cats, and peri-partum bitches and queens, "
                    "are inspected more frequently;\n"
                    "(b) dogs or cats with compromised welfare are, where necessary, transferred "
                    "without undue delay to a separate area and, where necessary, receive "
                    "appropriate treatment;\n"
                    "(c) where the recovery of a dog or a cat whose welfare is compromised is "
                    "not achievable and the dog or cat experiences severe pain or suffering, a "
                    "veterinarian is consulted without undue delay to decide whether the dog or "
                    "cat is to be euthanised to end its suffering, and, if that is the case, to "
                    "perform the euthanasia using anaesthesia and analgesia;\n"
                    "(d) measures are taken to prevent and control external and internal "
                    "parasites, and vaccinations are carried out to prevent common diseases to "
                    "which dogs or cats are likely to be exposed;\n"
                    "(e) enrichments that are used do not present a significant risk to dogs and "
                    "cats of injury or biological or chemical contamination or any other health "
                    "risk.\n"
                    "2. Operators of breeding establishments shall ensure that:\n"
                    "(a) measures are taken to safeguard the health of dogs or cats in "
                    "accordance with point 3 of Annex I;\n"
                    "(b) bitches or queens are bred only if they have reached a minimum age and "
                    "skeletal maturity in accordance with point 3 of Annex I, and only if they "
                    "have no diagnosed disease, clinical sign of diseases or physical conditions "
                    "which could negatively impact their pregnancy and welfare;\n"
                    "(c) the litter-giving pregnancies of bitches or queens follows a maximum "
                    "frequency in accordance with point 3 of Annex I;\n"
                    "(d) lactating queens are not mated or inseminated;\n"
                    "(e) dogs and cats which are no longer used for reproduction, including as "
                    "a result of the provisions of this Regulation, are either kept or sold, "
                    "donated or rehomed, and not killed or abandoned."
                ),
                "prazo": (
                    "Article 16 shall apply from … [three years from the date of entry into "
                    "force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (i)",
                "prazos": ["3a"],
            },
        ],
    },
    # =========================================================================
    {
        "titulo": "Grupo 2 — Identificação e Rastreabilidade",
        "normas": [
            {
                "ref": "Art. 20(2)",
                "norma": (
                    "2. Operators shall ensure that dogs and cats born in their establishments "
                    "are individually identified within three months after their birth, and in "
                    "any event before the date that they are placed on the market.\n"
                    "Operators of selling establishments and shelters, and operators who place "
                    "and are responsible for dogs and cats in foster homes shall ensure that "
                    "dogs and cats that enter their establishments or come under their "
                    "responsibility are individually identified within 30 days of arrival and "
                    "in any event before the date of their placing on the market.\n"
                    "Pet owners and any other natural or legal persons other than operators "
                    "who own dogs or cats, shall ensure that every dog or cat is individually "
                    "identified at the latest when it reaches the age of three months or, if "
                    "the dog or cat is placed on the market, before the date of that placing "
                    "on the market.\n"
                    "The implantation of the transponder shall be performed by a veterinarian. "
                    "Member States may allow the implantation of transponders by persons other "
                    "than veterinarians provided that they adopt national rules laying down the "
                    "minimum qualifications that such persons are required to have.\n"
                    "Where dogs and cats have been individually identified by means of an "
                    "injectable transponder containing a microchip in accordance with Union or "
                    "national law before … [date two years after the date of entry into force "
                    "of this Regulation] they shall be considered to be compliant with the "
                    "requirements in paragraph 1 and the first, second, third and fourth "
                    "subparagraphs of this paragraph, provided that the microchip is still readable."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 20(7)(a)",
                "norma": (
                    "7. The identification and registration requirements of this Article shall "
                    "apply as follows:\n"
                    "(a) for operators and natural or legal persons placing dogs and cats on "
                    "the market from ... [4 years from the entry into force of this Regulation];"
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["4a"],
            },
            {
                "ref": "Art. 20(7)(b)",
                "norma": (
                    "(b) for pet owners and other natural or legal persons other than operators, "
                    "who do not place dogs on the market: from … [10 years from entry into force "
                    "of this Regulation];"
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["10a"],
            },
            {
                "ref": "Art. 20(7)(c)",
                "norma": (
                    "(c) for pet owners and other natural or legal persons other than operators, "
                    "who do not place cats on the market: from … [15 years from entry into force "
                    "of this Regulation]."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["15a"],
            },
        ],
    },
    # =========================================================================
    {
        "titulo": "Grupo 3 — Bases de Dados e Publicidade Online",
        "normas": [
            {
                "ref": "Art. 21(3)\n[via Art. 33, ponto (ii)]",
                "norma": (
                    "3. When placing a dog or a cat on the market in the Union, the natural or "
                    "legal person placing the dog or cat on the market shall provide the acquirer "
                    "with:\n"
                    "(a) proof of the identification and registration of the dog or cat in "
                    "compliance with Article 20;\n"
                    "(b) the following information on the dog or cat:\n"
                    "(i) its species;\n"
                    "(ii) its sex;\n"
                    "(iii) its date and country of birth; and\n"
                    "(iv) where relevant, its breed.\n"
                    "Where a natural or legal person advertises a dog or cat online with a view "
                    "to placing it on the Union market, that person shall use the system referred "
                    "to in paragraph 5 to generate a unique verification token. That person shall "
                    "include that token in the advertisement, along with a weblink to the system "
                    "referred to in paragraph 5.\n"
                    "The system referred to in paragraph 5 shall enable acquirers to verify the "
                    "authenticity of the identification, registration and ownership of dogs or "
                    "cats advertised online."
                ),
                "prazo": (
                    "Article 21(3) and Article 23(1) shall apply from … [four years from entry "
                    "into force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (ii)",
                "prazos": ["4a"],
            },
            {
                "ref": "Art. 21(6)(a)",
                "norma": (
                    "6. The Commission shall adopt implementing acts laying down:\n"
                    "(a) the information to be provided by natural and legal persons placing "
                    "dogs or cats on the market as proof of identification and registration of "
                    "the dogs and cats in accordance with point (a) of paragraph 3;\n"
                    "The implementing acts referred to in point (a) shall be adopted by ... "
                    "[two years after the date of entry into force of this Regulation] and the "
                    "implementing act referred to in points (b) and (c) shall be adopted by ... "
                    "[three years from date of entry into force of this Regulation]. Those "
                    "implementing acts shall be adopted in accordance with the examination "
                    "procedure referred to in Article 29."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 21(6)(b)(c)",
                "norma": (
                    "(b) the information to be provided by natural and legal persons advertising "
                    "dogs or cats to the verification system referred to in paragraph 5 for the "
                    "purpose of demonstrating the authenticity of the identification, registration "
                    "and ownership of the dog or cat advertised.\n"
                    "(c) the following characteristics of the system referred to in paragraph 5:\n"
                    "— the key functions of the system;\n"
                    "— the technical, electronic and cryptographic requirements for the system;\n"
                    "— the procedural steps to be followed, and the information to be provided, "
                    "by the natural or legal person placing the dog or cat on the market, and "
                    "the steps and information required of the acquirer, in order for the online "
                    "verification system to work.\n"
                    "The implementing act referred to in points (b) and (c) shall be adopted by "
                    "... [three years from date of entry into force of this Regulation]."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["3a"],
            },
            {
                "ref": "Art. 23(1)\n[via Art. 33, ponto (ii)]",
                "norma": (
                    "1. Member States shall be responsible for establishing and maintaining "
                    "databases for the registration of identified dogs and cats in accordance "
                    "with Article 20(1) and (2) and Article 26(3) and Article 26(4), second "
                    "subparagraph."
                ),
                "prazo": (
                    "Article 21(3) and Article 23(1) shall apply from … [four years from entry "
                    "into force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (ii)",
                "prazos": ["4a"],
            },
            {
                "ref": "Art. 23(4), second subparagraph,\npontos (a)(c)",
                "norma": (
                    "4. The Commission shall establish and maintain an index database containing "
                    "the minimum set of fields laid down in the implementing acts referred to in "
                    "subparagraph 2, point (b). […] The Commission shall adopt implementing acts "
                    "laying down detailed arrangements concerning:\n"
                    "(a) the minimum content of the databases referred to in paragraph 1;\n"
                    "(b) the interoperability between Member States' databases and the index "
                    "database, including the minimum set of fields to be transmitted to the "
                    "index database and the intervals of the transmission;\n"
                    "(c) the functionality for providing proof of the identification and "
                    "registration of a dog or a cat, as referred to in Article 21(3) point (a);\n"
                    "(d) the registry where Member States will declare their databases, and the "
                    "necessary parameters for connecting those databases with one another in "
                    "accordance with the provisions established pursuant to point (b);\n"
                    "(e) the interconnection between the Member States' databases referred to "
                    "in paragraph 1, the pet travellers' database referred to in Article 26, "
                    "paragraph 4, and the Information Management System for Official Controls "
                    "(IMSOC), where relevant.\n"
                    "The Commission shall adopt the implementing acts referred to in the second "
                    "subparagraph, points (a) and (c) by ... [date two years after the date of "
                    "entry into force of this Regulation]."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 23(4), second subparagraph,\npontos (b)(d)(e)",
                "norma": (
                    "It shall adopt the implementing acts referred to in the second subparagraph, "
                    "points (b), (d) and (e) by ... [three years from the date of entry into "
                    "force of this Regulation]. Those implementing acts shall be adopted in "
                    "accordance with the examination procedure referred to in Article 29."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["3a"],
            },
        ],
    },
    # =========================================================================
    {
        "titulo": "Grupo 4 — Detenção Responsável e Formação",
        "normas": [
            {
                "ref": "Art. 15\n[via Art. 33, ponto (iv)]",
                "norma": (
                    "1. The operators of breeding and selling establishments shall ensure that "
                    "dogs and cats are housed in accordance with point 2 of Annex I. The "
                    "operators of shelters shall ensure that dogs and cats are housed in "
                    "accordance with point 2.2 of Annex I.\n"
                    "2. Operators shall ensure that:\n"
                    "(a) the establishments where dogs or cats are kept and the equipment used "
                    "therein are suitable for the types and the number of dogs or cats, and make "
                    "possible the necessary access to, and the thorough inspection of, all dogs "
                    "or cats;\n"
                    "(b) all building components of the establishment, including the flooring "
                    "and roof, and space divisions, as well as the equipment used for dogs or "
                    "cats, are constructed and maintained properly, to ensure that they do not "
                    "pose any risks to the welfare of the dogs or cats;\n"
                    "(c) all building components of the establishment, including the flooring, "
                    "and space divisions, as well as the equipment used for dogs or cats, are "
                    "kept clean to ensure that they do not pose any risks to the welfare of the "
                    "dogs or cats;\n"
                    "(d) in breeding and selling establishments where dogs or cats are kept "
                    "indoors, the dust levels, the temperature, and the relative air humidity "
                    "and gas concentrations are not harmful to dogs or cats and that ventilation "
                    "is sufficient to avoid overheating;\n"
                    "(e) dogs or cats have enough space to be able to move around freely and to "
                    "express species-specific behaviour according to their needs with the "
                    "possibility to withdraw and rest;\n"
                    "(f) dogs or cats have clean, comfortable and dry resting places that are "
                    "sufficiently large and numerous to ensure that all of them can lie down and "
                    "rest in a natural position at the same time;\n"
                    "(g) appropriate structures and measures are in place for dogs or cats that "
                    "are kept outdoors in order to protect them from adverse weather conditions, "
                    "including to prevent thermal stress, sunburn and frostbite.\n"
                    "3. Operators shall not keep dogs or cats in containers. However, containers "
                    "may be used for transportation, for the short term isolation of individual "
                    "dogs or cats and for participation in shows, exhibitions and competitions, "
                    "for puppies or kittens with reduced thermoregulation capacity or for puppies "
                    "or kittens together with their mothers, provided that, for the dogs or cats "
                    "concerned, stress is minimised and suffering is avoided, and they are able "
                    "to stand, turn around and lie down in a natural position.\n"
                    "4. Operators shall not keep dogs older than 8 weeks exclusively indoors. "
                    "Such dogs shall have daily access to an outdoor area, or be walked daily, "
                    "to allow exercise, exploration and socialisation. The minimum combined "
                    "duration of such daily access or walk shall be one hour in total. The "
                    "operator may only deviate from these requirements based on the written "
                    "advice of a veterinarian.\n"
                    "5. When cats are kept in catteries, operators shall design and construct "
                    "individual enclosures to allow cats to move around freely and to exhibit "
                    "their natural behaviour.\n"
                    "6. Operators of breeding and selling establishments shall ensure that in "
                    "indoor areas where dogs and cats are kept, an appropriate thermoneutral "
                    "zone is maintained that takes into account their coat type, age, size, "
                    "breed, and health.\n"
                    "7. Operators of breeding and selling establishments shall, where necessary, "
                    "use heating or cooling systems in the indoor enclosures at their "
                    "establishments to maintain good air quality and an appropriate temperature "
                    "and to remove excessive moisture.\n"
                    "8. Operators shall ensure that dogs or cats are exposed to light, and are "
                    "able to stay in the dark for sufficient and uninterrupted periods in order "
                    "to maintain a normal circadian rhythm."
                ),
                "prazo": (
                    "Article 15 … shall apply from … [five years from the date of entry into "
                    "force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (iv)",
                "prazos": ["5a"],
            },
            {
                "ref": "Art. 21(3), 2.º parágrafo +\nArt. 21(4)(5)\n[via Art. 33, ponto (iv)]",
                "norma": (
                    "Art. 21(3), 2.º §: Where a natural or legal person advertises a dog or "
                    "cat online with a view to placing it on the Union market, that person shall "
                    "use the system referred to in paragraph 5 to generate a unique verification "
                    "token. That person shall include that token in the advertisement, along "
                    "with a weblink to the system referred to in paragraph 5. The system "
                    "referred to in paragraph 5 shall enable acquirers to verify the "
                    "authenticity of the identification, registration and ownership of dogs or "
                    "cats advertised online.\n"
                    "Art. 21(4): Providers of online platforms shall ensure that their online "
                    "interface is designed and organised in such a way that makes it easier for "
                    "operators or other natural or legal persons who are placing dogs or cats on "
                    "the market to comply with their obligations under paragraphs 1 to 3 of this "
                    "Article, and in line with Article 31 of Regulation (EU) 2022/2065, and "
                    "shall inform acquirers, in a visible manner, of the possibility to verify "
                    "the authenticity of the identification, registration and ownership of the "
                    "dog or cat on the online verification system referred to in paragraph 5 "
                    "accessed via a weblink.\n"
                    "Art. 21(5): The Commission shall ensure that a verification system for "
                    "performing automated checks of the authenticity of the identification, "
                    "registration and ownership of dogs or cats advertised online, using the "
                    "database referred to in Article 23, is publicly available online, free of "
                    "charge and generates the unique verification token referred to in paragraph "
                    "3(2) of this Article."
                ),
                "prazo": (
                    "Article 21(3), second subparagraph, (4) and (5) … shall apply from … "
                    "[five years from the date of entry into force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (iv)",
                "prazos": ["5a"],
            },
            {
                "ref": "Art. 22(1)(a)(b)(c)\n[via Art. 33, ponto (iv)]",
                "norma": (
                    "1. For the purposes of Article 12, the competent authorities shall be "
                    "responsible for:\n"
                    "(a) ensuring that training courses are available for animal carers;\n"
                    "(b) approving the content of the training courses referred to in point (a), "
                    "in accordance with the minimum requirements laid down by the implementing "
                    "acts referred to in Article 12(4);\n"
                    "(c) certifying animal carers who have successfully completed the training "
                    "courses referred to in point (a).\n"
                    "The competent authorities may delegate the task referred to in point (c) "
                    "to providers of training courses."
                ),
                "prazo": (
                    "Article 22(1), points (a), (b) and (c) … shall apply from … [five years "
                    "from the date of entry into force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (iv)",
                "prazos": ["5a"],
            },
            {
                "ref": "Art. 23(3)(4)\n[via Art. 33, ponto (iv)]",
                "norma": (
                    "3. Member States shall ensure that their databases, as referred to in "
                    "paragraph 1, comply with the requirements laid down by the implementing "
                    "act referred to in paragraph 4, second subparagraph, point (b), to ensure "
                    "their interoperability.\n"
                    "4. The Commission shall establish and maintain an index database containing "
                    "the minimum set of fields laid down in the implementing acts referred to "
                    "in subparagraph 2, point (b). The Commission may entrust the development, "
                    "maintenance and operation of that index database to an independent entity, "
                    "following a public selection process pursuant to the relevant provisions "
                    "of Title VII of the Regulation (EU, Euratom) 2024/2509."
                ),
                "prazo": (
                    "Article 23(3) and (4) … shall apply from … [five years from the date of "
                    "entry into force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (iv)",
                "prazos": ["5a"],
            },
        ],
    },
    # =========================================================================
    {
        "titulo": "Grupo 5 — Importação e Movimento",
        "normas": [
            {
                "ref": "Art. 10\n[via Art. 33, ponto (vi)]",
                "norma": (
                    "1. Operators of breeding establishments that either produce or intend to "
                    "produce more than five litters per calendar year or that keep more than a "
                    "combined total of five bitches or queens at any given time shall place "
                    "dogs or cats on the market only after their establishment has been approved "
                    "by the competent authority.\n"
                    "2. The competent authority shall perform on-site inspections to verify that "
                    "the establishment meets the requirements of this Regulation. Member States "
                    "may allow such inspections to be carried out remotely, provided that the "
                    "means of distance communication used provides sufficient evidence for the "
                    "competent authority to perform reliable inspections. The competent authority "
                    "shall grant certificates of approval only to breeding establishments that "
                    "meet the requirements of this Regulation.\n"
                    "3. The competent authority shall maintain a publicly available list "
                    "including the following information for each approved establishment:\n"
                    "(a) the name, contact details and, where available, the URL of the website "
                    "of the establishment;\n"
                    "(b) the address of the establishment;\n"
                    "(c) the name of the operator;\n"
                    "(d) the species and, if relevant, the breeds related to the establishment "
                    "activities approved;\n"
                    "(e) the unique approval number assigned to the establishment by the "
                    "competent authority and the date of the approval and cessation of activities."
                ),
                "prazo": (
                    "Article 10 shall apply from … [eight years from the date of entry into "
                    "force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (vi)",
                "prazos": ["8a"],
            },
            {
                "ref": "Art. 26(1)(2)(3)\n[via Art. 33, ponto (iv)]",
                "norma": (
                    "1. Dogs and cats may be brought into the Union for the purpose of being "
                    "placed on the Union market only if the following conditions are met:\n"
                    "(a) they have been bred and kept in compliance with any of the following "
                    "requirements:\n"
                    "(i) the requirements contained in Chapter II of this Regulation;\n"
                    "(ii) requirements recognised by the Union, in accordance with Article 129 "
                    "of Regulation (EU) 2017/625, as being equivalent to those set out by "
                    "Chapter II of this Regulation; or\n"
                    "(iii) where applicable, requirements contained in a specific agreement "
                    "between the Union and the exporting country.\n"
                    "(b) they come from a third country or territory and from an establishment "
                    "listed in accordance with Articles 126 and 127 of Regulation (EU) 2017/625.\n"
                    "2. The official certificate referred to in Article 126(2), point (c), of "
                    "Regulation (EU) 2017/625 accompanying dogs and cats brought into the Union "
                    "from third countries and territories for the purpose of being placed on the "
                    "Union market, shall contain an attestation certifying compliance with "
                    "paragraph 1 of this Article.\n"
                    "3. Dogs and cats brought into the Union for the purpose of being placed on "
                    "the Union market shall be identified before their entry into the Union by a "
                    "veterinarian by means of an injectable transponder containing a readable "
                    "microchip that complies with the requirements set out in Annex II.\n"
                    "The operator responsible for the import of the dogs or cats into the Union "
                    "shall ensure that they are registered in a national database referred to in "
                    "Article 23(1), by a veterinarian, within five working days after they were "
                    "brought into the Union."
                ),
                "prazo": (
                    "Article 26(1), (2) and (3) shall apply from … [five years from the date "
                    "of entry into force of this Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (iv)",
                "prazos": ["5a"],
            },
            {
                "ref": "Art. 26(4), §§ 1-4\n[via Art. 33, ponto (vii)]",
                "norma": (
                    "4. The non-commercial movement of a dog or cat from a third country or "
                    "territory to the Union shall be prenotified by its owner to an online "
                    "Union pet travellers' database at least five working days before the dog "
                    "or cat crosses the Union border, except in the following cases:\n"
                    "(a) where the dog or cat is brought into the Union directly from third "
                    "countries or from territories fulfilling the conditions set out in Article "
                    "17(1), point (a), of Commission Delegated Regulation (EU) …/…, and\n"
                    "(b) where the dog or cat is registered in a Member State database referred "
                    "to in Article 23(1) of this Regulation.\n"
                    "Where the dog or cat stays more than six months in the Union, the owner "
                    "shall ensure that it is registered in the database of the Member State of "
                    "residence referred to in Article 23(1), by a veterinarian, within five "
                    "working days after the expiry of the sixth month since it entered the Union.\n"
                    "The Commission shall establish and maintain the Union pet travellers' "
                    "database referred to in the first subparagraph. The Commission may entrust "
                    "the development, maintenance and operation of that database to an independent "
                    "entity, following a public selection process pursuant to the relevant "
                    "provisions of Title VII of the Regulation (EU, Euratom) 2024/2509. Access "
                    "to that database shall be restricted to Member States' competent authorities "
                    "and to the Commission.\n"
                    "The Commission shall ensure that the database triggers iRASFF notifications "
                    "for pre-notified movements that present a risk of fraud. The Member State "
                    "receiving the notification shall take appropriate measures to follow it up "
                    "in accordance with Article 105(2) of Regulation (EU) 2017/625."
                ),
                "prazo": (
                    "Article 26(4) shall apply from … [10 years from the entry into force of "
                    "the Regulation]."
                ),
                "prazo_src": "Art. 33, ponto (vii)",
                "prazos": ["10a"],
            },
            {
                "ref": "Art. 26(4),\natos de execução\n(5.º parágrafo)",
                "norma": (
                    "The Commission shall by ... [8 years after the date of entry into force "
                    "of this Regulation] adopt implementing acts laying down detailed "
                    "arrangements for the following:\n"
                    "(i) the information to be pre-notified by owners in accordance with "
                    "paragraph (4) of this Article in the Union pet travellers' database, "
                    "taking into account the personal data protection requirements of Regulation "
                    "(EU) 2018/1725 and Regulation (EU) 2016/679;\n"
                    "(ii) the procedure by which the risk of fraud is to be established, which "
                    "is to take into account the activities carried out by the AAC network.\n"
                    "Those implementing acts shall be adopted in accordance with the examination "
                    "procedure referred to in Article 29."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["8a"],
            },
        ],
    },
    # =========================================================================
    {
        "titulo": "Grupo 6 — Relatórios e Avaliação",
        "normas": [
            {
                "ref": "Art. 24(2)",
                "norma": (
                    "2. The competent authorities shall draw up and transmit to the Commission "
                    "a report in electronic form, on the data on animal welfare, set out in "
                    "Annex III, by 31 August at three-yearly intervals. The first such report "
                    "shall be drawn up and transmitted to the Commission by ... [date 6 years "
                    "from the date of entry into force of this Regulation]. Each report shall "
                    "contain a summary of the data gathered during the previous three years."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["6a"],
            },
            {
                "ref": "Art. 31(1)",
                "norma": (
                    "1. On the basis of the reports received in accordance with Article 24 and "
                    "any additional relevant information, the Commission shall publish, by ... "
                    "[7 years from the date of entry into force of this Regulation] and "
                    "thereafter at three-yearly intervals, a monitoring report on the welfare "
                    "of dogs and cats placed on the market in the Union."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["7a"],
            },
            {
                "ref": "Art. 31(2)",
                "norma": (
                    "2. By … [14 years from the date of entry into force of this Regulation], "
                    "the Commission shall carry out an evaluation of this Regulation and present "
                    "a report on the main findings to the European Parliament, the Council, the "
                    "European Economic and Social Committee, and the Committee of the Regions. "
                    "In that evaluation and report the Commission shall assess, in particular:\n"
                    "(a) the extent to which this Regulation has contributed to ensuring a high "
                    "level of welfare for dogs and cats, improving their traceability, and "
                    "reducing the illegal trade in them;\n"
                    "(b) the impact that this Regulation has had on operators of breeding and "
                    "selling establishments and shelters, and of operators who place dogs and "
                    "cats in foster homes, taking into account inter alia the administrative "
                    "burden and compliance costs."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["14a"],
            },
        ],
    },
    # =========================================================================
    {
        "titulo": "Grupo 7 — Disposições Institucionais",
        "normas": [
            {
                "ref": "Art. 30(2)",
                "norma": (
                    "2. Member States shall, by ... [two years after the date of entry into "
                    "force of this Regulation], inform the Commission about any existing stricter "
                    "national rules that they intend to maintain in accordance with paragraph 1. "
                    "Thereafter, Member States shall inform the Commission about stricter national "
                    "rules before their adoption, unless the Member States have already notified "
                    "the draft national rules as a draft technical regulation under Article 5 of "
                    "Directive (EU) 2015/1535 of the European Parliament and of the Council. "
                    "The Commission shall bring them to the attention of the other Member States."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["2a"],
            },
            {
                "ref": "Art. 33\n(norma geral de aplicação)",
                "norma": (
                    "This Regulation shall enter into force on the twentieth day following that "
                    "of its publication in the Official Journal of the European Union.\n"
                    "It shall apply from ... [two years from the date of entry into force of "
                    "this Regulation]. However,\n"
                    "(i) Article 16 shall apply from … [three years from the date of entry into "
                    "force of this Regulation];\n"
                    "(ii) Article 21(3) and Article 23(1) shall apply from … [four years from "
                    "entry into force of this Regulation];\n"
                    "(iii) Article 8(1) shall apply from 1 July 2036 and Article 8(2) shall "
                    "apply from 1 July 2030;\n"
                    "(iv) Article 15, Article 21(3), second subparagraph, (4) and (5), Article "
                    "22(1), points (a),(b) and (c), Article 23(3) and (4), and Article 26(1), "
                    "(2) and (3) shall apply from … [five years from the date of entry into "
                    "force of this Regulation];\n"
                    "(v) Article 12(2) and (3) shall apply from … [seven years from the date "
                    "of entry into force of this Regulation];\n"
                    "(vi) Article 10 shall apply from … [eight years from the date of entry "
                    "into force of this Regulation]; and\n"
                    "(vii) Article 26(4) shall apply from … [10 years from the entry into force "
                    "of the Regulation].\n"
                    "This Regulation shall be binding in its entirety and directly applicable "
                    "in all Member States."
                ),
                "prazo": None,
                "prazo_src": None,
                "prazos": ["2a"],
            },
        ],
    },
]


# ---------------------------------------------------------------------------
# Auxiliares XML / python-docx
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn("w:shd")):
        tcPr.remove(s)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_width(cell, cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:tcW")):
        tcPr.remove(el)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm / 2.54 * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_cell_valign(cell, val="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(el)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)


def set_cell_margins(cell, top=40, bottom=40, left=72, right=72):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_paragraph_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn("w:spacing")):
        pPr.remove(s)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    pPr.append(sp)


def clear_cell(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    p.clear()
    return p


def disable_autofit(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)


def set_table_width(table, cm):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(cm / 2.54 * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), C_BORDER)
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Construtores de linhas
# ---------------------------------------------------------------------------

def build_header_row(row):
    cells = row.cells
    c = cells[0]
    set_cell_shading(c, C_HEADER_BG)
    set_cell_width(c, COL1_CM)
    set_cell_valign(c, "center")
    set_cell_margins(c)
    p = clear_cell(c)
    set_paragraph_spacing(p, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Norma / Disposição")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, col in enumerate(COLUNAS):
        c = cells[i + 1]
        set_cell_shading(c, C_HEADER_BG)
        set_cell_width(c, PRAZO_CM)
        set_cell_valign(c, "center")
        set_cell_margins(c, top=40, bottom=40, left=15, right=15)
        p = clear_cell(c)
        set_paragraph_spacing(p, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = col["label"].split("\n")
        r = p.add_run(parts[0])
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if len(parts) > 1:
            br = OxmlElement("w:br")
            r._r.append(br)
            r2 = p.add_run(parts[1])
            r2.bold = True
            r2.font.size = Pt(7.5)
            r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def build_group_row(row, titulo):
    cells = row.cells
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)
    set_cell_shading(merged, C_GROUP_BG)
    set_cell_valign(merged, "center")
    set_cell_margins(merged, top=50, bottom=50, left=100, right=80)
    p = clear_cell(merged)
    set_paragraph_spacing(p, 0, 0)
    run = p.add_run(titulo)
    run.bold = True
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_text_paragraphs(cell, text, font_size=7.0, color=None, italic=False,
                         spacing_after_last=0):
    """
    Adiciona o texto ao cell dividindo em parágrafos onde há '\n'.
    O primeiro parágrafo é obtido com clear_cell(); os seguintes com add_paragraph().
    Devolve (primeiro_p, ultimo_p).
    """
    if color is None:
        color = RGBColor(0x30, 0x30, 0x30)
    lines = text.split("\n")
    first_p = last_p = None
    for idx, line in enumerate(lines):
        if idx == 0:
            p = clear_cell(cell)
        else:
            p = cell.add_paragraph()
        after = spacing_after_last if idx == len(lines) - 1 else 0
        set_paragraph_spacing(p, 0, after)
        if line:
            r = p.add_run(line)
            r.font.size = Pt(font_size)
            r.font.color.rgb = color
            r.italic = italic
        if idx == 0:
            first_p = p
        last_p = p
    return first_p, last_p


def build_norma_row(row, norma, even):
    bg = C_ROW_EVEN if even else C_ROW_ODD
    cells = row.cells

    # ---- Coluna 1: referência + [NORMA] + [PRAZO] ----
    c = cells[0]
    set_cell_shading(c, bg)
    set_cell_width(c, COL1_CM)
    set_cell_valign(c, "top")
    set_cell_margins(c, top=50, bottom=50, left=80, right=60)

    # Parágrafo 1: referência
    p_ref = clear_cell(c)
    set_paragraph_spacing(p_ref, 0, 40)
    ref_parts = norma["ref"].split("\n")
    r_ref = p_ref.add_run(ref_parts[0])
    r_ref.bold = True
    r_ref.font.size = Pt(8)
    r_ref.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    for part in ref_parts[1:]:
        br = OxmlElement("w:br")
        r_ref._r.append(br)
        r_via = p_ref.add_run(part)
        r_via.italic = True
        r_via.font.size = Pt(6.5)
        r_via.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Parágrafos 2+: texto substantivo da norma (splitado por \n)
    has_prazo_block = norma.get("prazo") is not None
    norma_lines = norma["norma"].split("\n")
    for idx, line in enumerate(norma_lines):
        p = c.add_paragraph()
        is_last = (idx == len(norma_lines) - 1)
        set_paragraph_spacing(p, 0, 60 if (is_last and not has_prazo_block) else 0)
        if line:
            r = p.add_run(line)
            r.font.size = Pt(7)
            r.font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    # Bloco [PRAZO] — apenas se o prazo for externo ao texto da norma
    if has_prazo_block:
        # Label separador
        p_sep = c.add_paragraph()
        set_paragraph_spacing(p_sep, 50, 20)
        r_label = p_sep.add_run(f"▸ Prazo [{norma['prazo_src']}]:")
        r_label.bold = True
        r_label.font.size = Pt(6.5)
        r_label.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        # Texto verbatim do prazo
        p_prazo = c.add_paragraph()
        set_paragraph_spacing(p_prazo, 0, 0)
        r_prazo = p_prazo.add_run(norma["prazo"])
        r_prazo.italic = True
        r_prazo.font.size = Pt(7)
        r_prazo.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ---- Colunas de prazo ----
    for i, col in enumerate(COLUNAS):
        cell = cells[i + 1]
        set_cell_shading(cell, bg)
        set_cell_width(cell, PRAZO_CM)
        set_cell_valign(cell, "center")
        set_cell_margins(cell, top=40, bottom=40, left=8, right=8)
        p = clear_cell(cell)
        set_paragraph_spacing(p, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col["key"] in norma["prazos"]:
            r = p.add_run("✓")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


# ---------------------------------------------------------------------------
# Documento principal
# ---------------------------------------------------------------------------

def gerar():
    doc = Document()

    # Página landscape A4
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width  = int(29.7 / 2.54 * 914400)
    sec.page_height = int(21.0 / 2.54 * 914400)
    sec.left_margin   = int(MARGIN_CM / 2.54 * 914400)
    sec.right_margin  = int(MARGIN_CM / 2.54 * 914400)
    sec.top_margin    = int(1.5 / 2.54 * 914400)
    sec.bottom_margin = int(1.5 / 2.54 * 914400)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # Título
    h = doc.add_paragraph()
    set_paragraph_spacing(h, 0, 80)
    r = h.add_run("Prazos de Implementação — Regulamento 2023/0447 (cães e gatos) — Versão Extensa")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    set_paragraph_spacing(sub, 0, 160)
    r2 = sub.add_run(
        "Cada linha cita ipsis verbis o texto substantivo do artigo [NORMA] e a cláusula de prazo [PRAZO]. "
        "Organização temática. Prazos contados a partir da entrada em vigor (vigésimo dia após publicação no JOUE). "
        "Datas de 1 Jul. 2030 e 1 Jul. 2036 são absolutas (Art. 33, ponto iii)."
    )
    r2.italic = True
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    total_normas = sum(len(g["normas"]) for g in GRUPOS)
    total_rows = 1 + len(GRUPOS) + total_normas

    table = doc.add_table(rows=total_rows, cols=1 + N_PRAZO_COLS)
    disable_autofit(table)
    set_table_width(table, USABLE_CM)
    set_table_borders(table)

    build_header_row(table.rows[0])

    row_idx = 1
    norma_counter = 0
    for grupo in GRUPOS:
        build_group_row(table.rows[row_idx], grupo["titulo"])
        row_idx += 1
        norma_counter = 0
        for norma in grupo["normas"]:
            build_norma_row(table.rows[row_idx], norma, even=(norma_counter % 2 == 0))
            row_idx += 1
            norma_counter += 1

    # Nota de rodapé
    nota = doc.add_paragraph()
    set_paragraph_spacing(nota, 200, 0)
    rn = nota.add_run(
        "Nota: Texto verbatim em inglês conforme Regulamento 2023/0447 (versão de 11.12.2025). "
        "Os prazos em anos são contados a partir da entrada em vigor (vigésimo dia após publicação no JOUE). "
        "As datas 1 Jul. 2030 e 1 Jul. 2036 são absolutas (Art. 33, ponto iii). "
        "Quando há um bloco [PRAZO] separado, o prazo vem de Art. 33 (disposições gerais de aplicação); "
        "quando não há, o prazo está embutido no próprio artigo."
    )
    rn.italic = True
    rn.font.size = Pt(7.5)
    rn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out = "tabela_prazos_regulamento_extenso.docx"
    doc.save(out)
    print(f"Documento gerado: {out}")
    print(f"  Linhas na tabela: {total_rows} ({len(GRUPOS)} grupos, {total_normas} normas)")
    print(f"  Colunas: 1 (norma) + {N_PRAZO_COLS} (prazos) = {1 + N_PRAZO_COLS}")


if __name__ == "__main__":
    gerar()
