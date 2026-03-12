#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gerar_traducao_regulamento.py
Gera Regulamento_traduzido_PT.docx a partir do documento EN fonte,
usando traduções PT-PT extraídas de:
  - comparativo_reuniao_exemplo.docx (Art 5-33 + Anexos)
  - Regulamento - Primeira Versão portuguesa.docx (Art 1-4 + considerandos)
  - 11.12.2025 Regulamento cães e gatos - votação com tradução-ocr.docx (considerandos)
"""

import re
import shutil
import sys
from pathlib import Path
from docx import Document

# ---------------------------------------------------------------------------
# Ficheiros
# ---------------------------------------------------------------------------
BASE = Path(__file__).parent
EN_SOURCE = BASE / "11.12.2025 Regulamento cães e gatos-ocr - sem rasuras(atualizado) com índice automático - sem quebras.docx"
COMPARATIVO = BASE / "comparativo_reuniao_exemplo.docx"
PT1_FILE    = BASE / "Regulamento - Primeira Versão portuguesa.docx"
VOTACAO     = BASE / "11.12.2025 Regulamento cães e gatos - votação com tradução-ocr.docx"
OUTPUT      = BASE / "Regulamento_traduzido_PT.docx"

# ---------------------------------------------------------------------------
# Dicionário de termos EN→PT (fallback para texto sem tradução exacta)
# ---------------------------------------------------------------------------
TERMOS = [
    ("Member States", "Estados-Membros"),
    ("Member State", "Estado-Membro"),
    ("European Commission", "Comissão Europeia"),
    ("the Commission", "a Comissão"),
    ("competent authorities", "autoridades competentes"),
    ("competent authority", "autoridade competente"),
    ("dogs and cats", "cães e gatos"),
    ("dog or cat", "cão ou gato"),
    ("dogs or cats", "cães ou gatos"),
    ("dog and cat", "cão e gato"),
    ("welfare", "bem-estar"),
    ("traceability", "rastreabilidade"),
    ("breeding establishment", "estabelecimento de criação"),
    ("breeding establishments", "estabelecimentos de criação"),
    ("placing on the market", "colocação no mercado"),
    ("placed on the market", "colocados no mercado"),
    ("Union market", "mercado da União"),
    ("internal market", "mercado interno"),
    ("transponder", "transpondedor"),
    ("transponders", "transpondedores"),
    ("microchip", "microchip"),
    ("microchips", "microchips"),
    ("operator", "operador"),
    ("operators", "operadores"),
    ("pet", "animal de companhia"),
    ("pets", "animais de companhia"),
    ("animal welfare", "bem-estar animal"),
    ("third countries", "países terceiros"),
    ("third country", "país terceiro"),
    ("Union", "União"),
    ("Regulation", "Regulamento"),
    ("regulation", "regulamento"),
    ("Directive", "Diretiva"),
    ("directive", "diretiva"),
    ("Article", "Artigo"),
    ("article", "artigo"),
    ("Annex", "Anexo"),
    ("annex", "anexo"),
    ("paragraph", "n.º"),
    ("point", "alínea"),
    ("entry into force", "entrada em vigor"),
    ("date of application", "data de aplicação"),
    ("delegated acts", "atos delegados"),
    ("implementing acts", "atos de execução"),
    ("official controls", "controlos oficiais"),
    ("non-commercial movement", "circulação não comercial"),
    ("commercial movement", "circulação comercial"),
    ("identification", "identificação"),
    ("registration", "registo"),
    ("database", "base de dados"),
    ("databases", "bases de dados"),
    ("stray animals", "animais errantes"),
    ("stray animal", "animal errante"),
    ("shelter", "abrigo"),
    ("shelters", "abrigos"),
    ("rescue organisation", "organização de resgate"),
    ("rescue organisations", "organizações de resgate"),
    ("breeder", "criador"),
    ("breeders", "criadores"),
    ("litter", "ninhada"),
    ("litters", "ninhadas"),
    ("dam", "progenitora"),
    ("sire", "progenitor"),
    ("in utero", "in utero"),
    ("health certificate", "certificado de saúde"),
    ("health certificates", "certificados de saúde"),
    ("veterinarian", "médico veterinário"),
    ("veterinarians", "médicos veterinários"),
    ("veterinary", "veterinário"),
    ("neutering", "esterilização"),
    ("castration", "castração"),
    ("spaying", "ovariohisterectomia"),
    ("sterilisation", "esterilização"),
    ("enforcement", "aplicação"),
    ("penalties", "sanções"),
    ("penalty", "sanção"),
    ("infringement", "infração"),
    ("infringements", "infrações"),
    ("proportionate", "proporcionadas"),
    ("proportional", "proporcional"),
    ("effective", "eficazes"),
    ("dissuasive", "dissuasoras"),
    ("natural person", "pessoa singular"),
    ("legal person", "pessoa coletiva"),
    ("natural persons", "pessoas singulares"),
    ("legal persons", "pessoas coletivas"),
    ("phased-in", "faseada"),
    ("transitional", "transitório"),
    ("transitional period", "período transitório"),
    ("derogation", "derrogação"),
    ("derogations", "derrogações"),
    ("pursuant to", "nos termos de"),
    ("in accordance with", "em conformidade com"),
    ("by way of derogation", "em derrogação"),
    ("without prejudice to", "sem prejuízo de"),
    ("inter alia", "nomeadamente"),
    ("i.e.", "ou seja"),
    ("e.g.", "por exemplo"),
    ("etc.", "etc."),
    ("TFEU", "TFUE"),
    ("TEU", "TUE"),
]

# ---------------------------------------------------------------------------
# Dicionário de preâmbulo / fórmulas fixas
# ---------------------------------------------------------------------------
PREAMBLE_EXACT = {
    "OF THE EUROPEAN PARLIAMENT AND OF THE COUNCIL": "DO PARLAMENTO EUROPEU E DO CONSELHO",
    "on the welfare of dogs and cats and their traceability": "relativo ao bem-estar dos cães e gatos e à respetiva rastreabilidade",
    "THE EUROPEAN PARLIAMENT AND THE COUNCIL OF THE EUROPEAN UNION,": "O PARLAMENTO EUROPEU E O CONSELHO DA UNIÃO EUROPEIA,",
    "Having regard to the proposal from the European Commission,": "Tendo em conta a proposta da Comissão Europeia,",
    "Having regard to the opinion of the European Economic and Social Committee,": "Tendo em conta o parecer do Comité Económico e Social Europeu,",
    "After consulting the Committee of the Regions,": "Após consulta do Comité das Regiões,",
    "Acting in accordance with the ordinary legislative procedure,": "Deliberando de acordo com o processo legislativo ordinário,",
    "Whereas:": "Considerando o seguinte:",
    "HAVE ADOPTED THIS REGULATION:": "ADOTARAM O PRESENTE REGULAMENTO:",
    "Done at": "Feito em",
    "For the European Parliament": "Pelo Parlamento Europeu",
    "For the Council": "Pelo Conselho",
    "The President": "O Presidente",
    "The President of the European Parliament": "O Presidente do Parlamento Europeu",
    "The President of the Council": "O Presidente do Conselho",
}

# Prefixos de preâmbulo para matching parcial
PREAMBLE_STARTS = [
    ("Having regard to the Treaty on the Functioning", "Tendo em conta o Tratado sobre o Funcionamento da União Europeia, nomeadamente"),
    ("Having regard to the Treaty", "Tendo em conta o Tratado"),
    ("After transmission of the draft legislative act", "Após transmissão do projeto de ato legislativo"),
    ("After transmission", "Após transmissão"),
    ("Acting in ordinary legislative procedure", "Deliberando de acordo com o processo legislativo ordinário,"),
]

# Dicionário de capítulos e títulos
HEADING_DICT = {
    "CHAPTER I": "CAPÍTULO I",
    "CHAPTER II": "CAPÍTULO II",
    "CHAPTER III": "CAPÍTULO III",
    "CHAPTER IV": "CAPÍTULO IV",
    "CHAPTER V": "CAPÍTULO V",
    "CHAPTER VI": "CAPÍTULO VI",
    "CHAPTER VII": "CAPÍTULO VII",
    "CHAPTER VIII": "CAPÍTULO VIII",
    "CHAPTER IX": "CAPÍTULO IX",
    "CHAPTER X": "CAPÍTULO X",
    "GENERAL PROVISIONS": "DISPOSIÇÕES GERAIS",
    "SUBJECT MATTER, SCOPE AND DEFINITIONS": "OBJETO, ÂMBITO DE APLICAÇÃO E DEFINIÇÕES",
    "ANIMAL WELFARE REQUIREMENTS": "REQUISITOS DE BEM-ESTAR ANIMAL",
    "TRACEABILITY": "RASTREABILIDADE",
    "MOVEMENT": "CIRCULAÇÃO",
    "OFFICIAL CONTROLS AND ENFORCEMENT": "CONTROLOS OFICIAIS E APLICAÇÃO",
    "DELEGATED AND IMPLEMENTING POWERS": "PODERES DELEGADOS E DE EXECUÇÃO",
    "FINAL PROVISIONS": "DISPOSIÇÕES FINAIS",
    "ANNEXES": "ANEXOS",
    "ANNEX I": "ANEXO I",
    "ANNEX II": "ANEXO II",
    "ANNEX III": "ANEXO III",
}

# Títulos dos artigos EN→PT
ARTICLE_TITLES = {
    "1": "Objeto",
    "2": "Âmbito de aplicação",
    "3": "Definições",
    "4": "Derrogações",
    "5": "Deveres gerais dos operadores",
    "6": "Requisitos de bem-estar",
    "7": "Responsabilidades dos operadores",
    "8": "Estabelecimentos de criação",
    "9": "Identificação e registo",
    "10": "Base de dados nacional",
    "11": "Intercâmbio de informações",
    "12": "Circulação de cães e gatos",
    "13": "Exceções à identificação",
    "14": "Requisitos de identificação",
    "15": "Requisitos de registo",
    "16": "Operadores de alojamento",
    "17": "Colocação no mercado",
    "18": "Publicidade",
    "19": "Circulação não comercial",
    "20": "Circulação comercial",
    "21": "Importação",
    "22": "Exportação",
    "23": "Controlos oficiais",
    "24": "Autoridades competentes",
    "25": "Poderes de execução",
    "26": "Sanções",
    "27": "Exercício da delegação",
    "28": "Atos de execução",
    "29": "Procedimento de comité",
    "30": "Alterações",
    "31": "Revogações",
    "32": "Disposições transitórias",
    "33": "Entrada em vigor e data de aplicação",
}


def apply_terms(text):
    """Aplica substituição de termos EN→PT ao texto."""
    result = text
    for en, pt in TERMOS:
        result = result.replace(en, pt)
    return result


# ---------------------------------------------------------------------------
# Fase 1a — Extrair traduções do comparativo (Art 5-33 + Anexos)
# ---------------------------------------------------------------------------
def extract_comparativo_translations():
    """Extrai traduções PT-PT das tabelas do comparativo."""
    translations = {}
    try:
        doc = Document(COMPARATIVO)
    except Exception as e:
        print(f"AVISO: Não foi possível abrir comparativo: {e}", file=sys.stderr)
        return translations

    for table in doc.tables:
        try:
            if not table.rows:
                continue
            header_cell = table.rows[0].cells[0].text.strip()
            if 'Tradução PT-PT' not in header_cell:
                continue
            # Extrair número do artigo/anexo do cabeçalho
            # Padrões: "Art.º 5.º do Regulamento", "ANEXO I ponto 1", etc.
            content = ""
            if len(table.rows) > 1:
                content = table.rows[1].cells[0].text.strip()

            # Determinar chave
            key = None
            # Anexos
            m_annex = re.search(r'ANEX[OI]+\s+(I+|II|III)(?:[,\s]+[Pp]onto\s+(\d+))?', header_cell, re.IGNORECASE)
            if m_annex:
                annex_num = m_annex.group(1).upper()
                ponto = m_annex.group(2)
                if ponto:
                    key = f"ANEXO {annex_num} ponto {ponto}"
                else:
                    key = f"ANEXO {annex_num}"
            else:
                # Artigos: "Art.º 5.º" ou "Art. 5"
                m_art = re.search(r'Art(?:\.º|igo)?\s*(\d+)\.?º?', header_cell, re.IGNORECASE)
                if m_art:
                    num = m_art.group(1)
                    key = f"{num}.º"

            if key and content:
                translations[key] = content
        except Exception:
            continue

    print(f"Comparativo: {len(translations)} entradas extraídas: {sorted(translations.keys())[:10]}...")
    return translations


# ---------------------------------------------------------------------------
# Fase 1b — Extrair traduções PT1 (Art 1-4 + considerandos)
# ---------------------------------------------------------------------------
def extract_pt1_data():
    """Extrai artigos 1-4 e considerandos do PT1."""
    art_blocks = {}
    recitals = []
    try:
        doc = Document(PT1_FILE)
    except Exception as e:
        print(f"AVISO: Não foi possível abrir PT1: {e}", file=sys.stderr)
        return art_blocks, recitals

    paras = doc.paragraphs

    # Extrair considerandos (style Considerant)
    for p in paras:
        style_name = p.style.name if p.style else ''
        text = p.text.strip()
        if not text:
            continue
        if 'onsider' in style_name or 'Recital' in style_name:
            recitals.append(text)

    # Extrair artigos após "ADOTARAM"
    in_articles = False
    current_art = None
    current_block = []

    for p in paras:
        style_name = p.style.name if p.style else ''
        text = p.text.strip()

        if not in_articles:
            if 'ADOTARAM' in text or 'ADOPTED' in text:
                in_articles = True
            continue

        if 'Titre article' in style_name or 'titre article' in style_name.lower() or \
           ('Heading' in style_name and re.match(r'Artigo\s+\d+', text, re.IGNORECASE)):
            # Extrair número (ex: "Artigo 1.º"); ignorar sub-títulos sem número (ex: "Objeto")
            m = re.search(r'Artigo\s+(\d+)\.?', text, re.IGNORECASE)
            if m:
                # Guardar artigo anterior antes de começar novo
                if current_art is not None and current_block:
                    art_blocks[current_art] = '\n'.join(current_block)
                current_art = int(m.group(1))
                current_block = []
            # Se não há número (sub-título), não resetar current_art
        elif current_art is not None and text:
            current_block.append(text)

    if current_art is not None and current_block:
        art_blocks[current_art] = '\n'.join(current_block)

    print(f"PT1: {len(recitals)} considerandos, artigos {sorted(art_blocks.keys())}")
    return art_blocks, recitals


# ---------------------------------------------------------------------------
# Fase 1c — Extrair considerandos da votação
# ---------------------------------------------------------------------------
def extract_votacao_recitals():
    """Extrai considerandos PT das tabelas da votação."""
    recitals = {}
    try:
        doc = Document(VOTACAO)
    except Exception as e:
        print(f"AVISO: Não foi possível abrir votação: {e}", file=sys.stderr)
        return recitals

    for table in doc.tables:
        try:
            if not table.rows:
                continue
            # Percorrer linhas para encontrar considerandos
            for row in table.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                # Primeira célula — texto PT da proposta da Comissão
                cell0 = cells[0].text.strip()
                # Última célula — emenda do Parlamento
                cell_last = cells[-1].text.strip()

                # Verificar se começa com número de considerando
                m = re.match(r'^\((\d+)\)\s+(.+)', cell0, re.DOTALL)
                if m:
                    num = int(m.group(1))
                    text = m.group(2).strip()
                    recitals[num] = text

                # Também verificar emendas
                m2 = re.match(r'^\((\d+)\)\s+(.+)', cell_last, re.DOTALL)
                if m2:
                    num2 = int(m2.group(1))
                    text2 = m2.group(2).strip()
                    if num2 not in recitals:
                        recitals[num2] = text2
        except Exception:
            continue

    # Também verificar parágrafos directos
    for p in doc.paragraphs:
        text = p.text.strip()
        m = re.match(r'^\((\d+)\)\s+(.+)', text)
        if m:
            num = int(m.group(1))
            content = m.group(2).strip()
            if num not in recitals and len(content) > 20:
                recitals[num] = content

    print(f"Votação: {len(recitals)} considerandos com numeração da proposta")
    return recitals


# ---------------------------------------------------------------------------
# Construir dicionário de considerandos EN → PT por matching de conteúdo
# ---------------------------------------------------------------------------
def build_recital_dict(en_doc, votacao_recitals, pt1_recitals):
    """
    Faz o matching entre considerandos EN (numeração final) e PT disponível.
    Estratégia:
    1. Tentar matching por conteúdo com votação (primeiras palavras)
    2. Se não encontrar, usar PT1 sequencial
    3. Fallback: substituição de termos
    """
    recital_dict = {}

    # Recolher considerandos EN do documento fonte
    en_recitals = []
    for p in en_doc.paragraphs:
        style_name = p.style.name if p.style else ''
        if 'Consid' in style_name or 'onsid' in style_name:
            text = p.text.strip()
            m = re.match(r'^\((\d+)\)\s*\t?\s*(.+)', text, re.DOTALL)
            if m:
                num = int(m.group(1))
                content = m.group(2).strip()
                en_recitals.append((num, content))

    print(f"EN: {len(en_recitals)} considerandos encontrados")

    # Construir lista de PT disponível para matching
    # votacao_recitals: {num_proposta: text_pt}
    # pt1_recitals: [text_pt, ...] (sequential)
    pt_pool = list(votacao_recitals.values()) + pt1_recitals

    # Para cada considerando EN, tentar matching por primeiras palavras
    pt1_index = 0
    used_votacao = set()

    for en_num, en_content in en_recitals:
        # Extrair primeiras 6 palavras do EN para comparação
        en_words = re.findall(r'\w+', en_content.lower())[:6]

        best_match = None
        best_score = 0

        # Tentar votação primeiro
        for v_num, v_text in votacao_recitals.items():
            if v_num in used_votacao:
                continue
            pt_words = re.findall(r'\w+', v_text.lower())[:6]
            # Score: número de palavras comuns nas primeiras posições
            # (não fazemos matching directo EN-PT pois são línguas diferentes)
            # Usamos apenas para verificar se o número bate com algum critério
            # Estratégia simplificada: usar PT1 sequencialmente para a maioria

        # Usar PT1 sequencial como fonte principal de considerandos
        if pt1_index < len(pt1_recitals):
            best_match = pt1_recitals[pt1_index]
            pt1_index += 1
        else:
            # Fallback: substituição de termos
            best_match = apply_terms(en_content)

        recital_dict[en_num] = best_match

    print(f"Dicionário de considerandos: {len(recital_dict)} entradas")
    return recital_dict


# ---------------------------------------------------------------------------
# Matching parágrafo EN → PT dentro de um artigo
# ---------------------------------------------------------------------------
def split_pt_block(pt_block):
    """Divide bloco PT em parágrafos por linha, filtrando vazios."""
    lines = [l.strip() for l in pt_block.split('\n') if l.strip()]
    return lines


def get_prefix(text):
    """Extrai prefixo numérico ou de alínea do texto."""
    # Número: "1.", "2.", "1.\t", etc.
    m = re.match(r'^(\d+)[.\t\s]', text)
    if m:
        return ('num', m.group(1))
    # Alínea: "(a)", "(b)", etc.
    m = re.match(r'^\(([a-z]+)\)', text)
    if m:
        return ('alinea', m.group(1))
    # Sub-alínea: "(i)", "(ii)", etc.
    m = re.match(r'^\(([ivxlcdm]+)\)', text)
    if m:
        return ('sub', m.group(1))
    # Alínea portuguesa: "a)", "b)"
    m = re.match(r'^([a-z])\)', text)
    if m:
        return ('alinea_pt', m.group(1))
    return None


def find_pt_for_en_para(en_text, pt_lines, used_indices):
    """
    Encontra a linha PT correspondente ao parágrafo EN.
    Tenta matching por prefixo, depois por posição ordinal.
    """
    en_prefix = get_prefix(en_text)

    if en_prefix:
        ptype, pval = en_prefix
        for i, pt_line in enumerate(pt_lines):
            if i in used_indices:
                continue
            pt_prefix = get_prefix(pt_line)
            if pt_prefix and pt_prefix[0] in ('num', 'alinea', 'alinea_pt', 'sub'):
                if pt_prefix[1] == pval:
                    used_indices.add(i)
                    return pt_line

    # Fallback: primeiro disponível por ordem
    for i, pt_line in enumerate(pt_lines):
        if i not in used_indices:
            used_indices.add(i)
            return pt_line

    return None


# ---------------------------------------------------------------------------
# Preservação de formatação dos runs
# ---------------------------------------------------------------------------
def replace_para_text(para, new_text):
    """Substitui texto do parágrafo preservando formatação do primeiro run."""
    if not new_text:
        return
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)


# ---------------------------------------------------------------------------
# Tradução de cabeçalhos
# ---------------------------------------------------------------------------
def translate_heading(text):
    """Traduz cabeçalho de capítulo ou artigo."""
    # Capítulos: "CHAPTER I — GENERAL PROVISIONS"
    m_chap = re.match(r'^(CHAPTER\s+[IVXLCDM]+)\s*[—–-]?\s*(.*)', text, re.IGNORECASE)
    if m_chap:
        chap_num = m_chap.group(1).upper()
        chap_title = m_chap.group(2).strip()
        pt_chap = HEADING_DICT.get(chap_num, chap_num.replace('CHAPTER', 'CAPÍTULO'))
        # Traduzir título
        pt_title = chap_title
        for en_key, pt_val in HEADING_DICT.items():
            if en_key.upper() == chap_title.upper():
                pt_title = pt_val
                break
        if pt_title == chap_title:
            pt_title = apply_terms(chap_title)
        return f"{pt_chap} — {pt_title}" if pt_title else pt_chap

    # Artigos: "Article 5 — Operators' general obligations" ou "Article 5 - title"
    m_art = re.match(r'^Article\s+(\d+)\s*[—–-]\s*(.*)', text, re.IGNORECASE)
    if m_art:
        num = m_art.group(1)
        en_title = m_art.group(2).strip()
        pt_title = ARTICLE_TITLES.get(num, apply_terms(en_title))
        return f"Artigo {num}.º — {pt_title}"

    # Sem separador
    m_art2 = re.match(r'^Article\s+(\d+)$', text.strip(), re.IGNORECASE)
    if m_art2:
        num = m_art2.group(1)
        pt_title = ARTICLE_TITLES.get(num, '')
        if pt_title:
            return f"Artigo {num}.º — {pt_title}"
        return f"Artigo {num}.º"

    # Anexos
    m_annex = re.match(r'^(ANNEX\s+[IVXLCDM]+)(.*)', text, re.IGNORECASE)
    if m_annex:
        ann = m_annex.group(1).upper()
        rest = m_annex.group(2)
        pt_ann = ann.replace('ANNEX', 'ANEXO')
        return pt_ann + apply_terms(rest) if rest else pt_ann

    # Fallback
    for en_key, pt_val in HEADING_DICT.items():
        if en_key.upper() in text.upper():
            return text.upper().replace(en_key.upper(), pt_val)

    return apply_terms(text)


# ---------------------------------------------------------------------------
# Função principal de tradução de parágrafo
# ---------------------------------------------------------------------------
def translate_paragraph(para, style_name, text,
                         art_translations, recital_dict,
                         current_article, art_pt_lines, art_used_indices,
                         current_annex, annex_pt_lines, annex_used_indices):
    """
    Devolve (new_text, updated_state) para o parágrafo.
    Modifica in-place os dicionários de estado do artigo corrente.
    """

    if not text:
        return text

    # --- Cabeçalhos de capítulo (Heading 1) ---
    if 'Heading 1' in style_name:
        return translate_heading(text)

    # --- Cabeçalhos de artigo (Heading 3) ---
    if 'Heading 3' in style_name:
        return translate_heading(text)

    # --- Considerandos ---
    if 'Consid' in style_name or 'onsid' in style_name:
        m = re.match(r'^\((\d+)\)\s*\t?\s*(.*)', text, re.DOTALL)
        if m:
            num = int(m.group(1))
            en_content = m.group(2).strip()
            pt_content = recital_dict.get(num)
            if pt_content:
                # Preservar tab se existir no original
                sep = '\t' if '\t' in text else ' '
                return f"({num}){sep}{pt_content}"
            else:
                # Fallback: substituição de termos
                return f"({num}){sep if 'sep' in dir() else ' '}{apply_terms(en_content)}"
        # Cabeçalho "Whereas:" etc.
        return translate_preamble(text)

    # --- Preâmbulo/fórmulas (Normale, Application directe) ---
    if 'Normale' in style_name or 'Application' in style_name or 'Statut' in style_name:
        return translate_preamble(text)

    # --- Corpo de artigo ou Anexo (Normal) ---
    if current_annex and annex_pt_lines:
        pt_line = find_pt_for_en_para(text, annex_pt_lines, annex_used_indices)
        if pt_line:
            return pt_line

    if current_article and art_pt_lines:
        pt_line = find_pt_for_en_para(text, art_pt_lines, art_used_indices)
        if pt_line:
            return pt_line

    # Fallback: substituição de termos
    return apply_terms(text)


def translate_preamble(text):
    """Traduz texto do preâmbulo usando dicionário fixo + termos."""
    # Exact match
    if text in PREAMBLE_EXACT:
        return PREAMBLE_EXACT[text]

    # Strip e tentar novamente
    t = text.strip()
    if t in PREAMBLE_EXACT:
        return PREAMBLE_EXACT[t]

    # Partial starts
    for en_start, pt_start in PREAMBLE_STARTS:
        if text.startswith(en_start):
            return pt_start

    # REGULATION (EU) pattern
    m = re.match(r'^(REGULATION\s*\(EU\)\s*[\d/]+)', text, re.IGNORECASE)
    if m:
        return text.replace('REGULATION', 'REGULAMENTO').replace('(EU)', '(UE)')

    # Datas e assinaturas
    if re.match(r'^Done at', text):
        return re.sub(r'^Done at', 'Feito em', text)

    return apply_terms(text)


# ---------------------------------------------------------------------------
# Detecção do artigo/anexo corrente
# ---------------------------------------------------------------------------
def detect_article(text, style_name):
    """Detecta se o parágrafo é um cabeçalho de artigo e devolve o número."""
    if 'Heading 3' in style_name:
        m = re.match(r'^Article\s+(\d+)', text, re.IGNORECASE)
        if m:
            return int(m.group(1))
    return None


def detect_annex(text, style_name):
    """Detecta se o parágrafo é um cabeçalho de anexo."""
    if 'Heading 1' in style_name or 'Heading 3' in style_name:
        m = re.match(r'^ANNEX\s+([IVXLCDM]+)(?:\s+(.*))?', text, re.IGNORECASE)
        if m:
            ann_num = m.group(1).upper()
            ann_sub = m.group(2) if m.group(2) else ''
            return ann_num, ann_sub
    return None, None


def get_annex_key(ann_num, ann_sub, art_translations):
    """Determina a chave do dicionário para um anexo."""
    # Tentar com ponto número
    m = re.search(r'(\d+)', ann_sub) if ann_sub else None
    if m:
        key = f"ANEXO {ann_num} ponto {m.group(1)}"
        if key in art_translations:
            return key
    # Tentar sem ponto
    key = f"ANEXO {ann_num}"
    if key in art_translations:
        return key
    # Tentar variações
    for k in art_translations:
        if f"ANEXO {ann_num}" in k:
            return k
    return None


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main():
    print("=== Gerador de Tradução PT do Regulamento ===")

    # Verificar ficheiros
    for f in [EN_SOURCE, COMPARATIVO, PT1_FILE, VOTACAO]:
        if not f.exists():
            print(f"ERRO: Ficheiro não encontrado: {f}", file=sys.stderr)
            sys.exit(1)

    # Fase 1: Extrair traduções
    print("\n--- Fase 1: Extracção de traduções ---")
    art_translations = extract_comparativo_translations()
    pt1_art_blocks, pt1_recitals = extract_pt1_data()
    votacao_recitals = extract_votacao_recitals()

    # Adicionar artigos 1-4 do PT1 ao dicionário principal
    for art_num, block in pt1_art_blocks.items():
        key = f"{art_num}.º"
        if key not in art_translations:
            art_translations[key] = block
            print(f"  Art {key} adicionado do PT1")

    # Carregar documento EN para construir dicionário de considerandos
    print("\n--- Carregando documento EN ---")
    en_doc = Document(EN_SOURCE)

    # Construir dicionário de considerandos
    recital_dict = build_recital_dict(en_doc, votacao_recitals, pt1_recitals)

    # Fase 2+3: Copiar EN e substituir texto
    print(f"\n--- Fase 2+3: Geração de {OUTPUT.name} ---")
    shutil.copy(EN_SOURCE, OUTPUT)
    out_doc = Document(OUTPUT)

    # Estado do processamento
    current_article = None       # número do artigo corrente
    art_pt_lines = []            # linhas PT do artigo corrente
    art_used_indices = set()     # índices já usados

    current_annex = None         # nome do anexo corrente
    current_annex_point = 1      # ponto do anexo
    annex_pt_lines = []
    annex_used_indices = set()

    stats = {'translated': 0, 'fallback': 0, 'empty': 0, 'total': 0}

    for para in out_doc.paragraphs:
        stats['total'] += 1
        text = para.text.strip()
        style_name = para.style.name if para.style else 'Normal'

        if not text:
            stats['empty'] += 1
            continue

        # Detectar mudança de artigo
        new_art = detect_article(text, style_name)
        if new_art is not None:
            current_article = new_art
            current_annex = None
            annex_pt_lines = []
            annex_used_indices = set()
            art_key = f"{new_art}.º"
            pt_block = art_translations.get(art_key, '')
            art_pt_lines = split_pt_block(pt_block) if pt_block else []
            art_used_indices = set()

        # Detectar mudança de anexo
        ann_num, ann_sub = detect_annex(text, style_name)
        if ann_num:
            # Verificar se é início de ponto do Anexo I
            m_ponto = re.search(r'(\d+)', ann_sub) if ann_sub else None
            if m_ponto:
                ann_key = f"ANEXO {ann_num} ponto {m_ponto.group(1)}"
            else:
                ann_key = get_annex_key(ann_num, ann_sub, art_translations)

            if ann_key and ann_key in art_translations:
                current_annex = ann_key
                current_article = None
                art_pt_lines = []
                art_used_indices = set()
                pt_block = art_translations[ann_key]
                annex_pt_lines = split_pt_block(pt_block)
                annex_used_indices = set()
            else:
                current_annex = None
                annex_pt_lines = []
                annex_used_indices = set()

        # Traduzir
        new_text = translate_paragraph(
            para, style_name, text,
            art_translations, recital_dict,
            current_article, art_pt_lines, art_used_indices,
            current_annex, annex_pt_lines, annex_used_indices
        )

        if new_text != text:
            replace_para_text(para, new_text)
            stats['translated'] += 1
        else:
            stats['fallback'] += 1

    # Processar tabelas (se houver no documento EN)
    for table in out_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    stats['total'] += 1
                    text = para.text.strip()
                    if not text:
                        stats['empty'] += 1
                        continue
                    style_name = para.style.name if para.style else 'Normal'
                    new_text = translate_preamble(text)
                    if new_text != text:
                        replace_para_text(para, new_text)
                        stats['translated'] += 1

    out_doc.save(OUTPUT)

    print(f"\n=== Concluído ===")
    print(f"Output: {OUTPUT}")
    print(f"Total parágrafos: {stats['total']}")
    print(f"Traduzidos: {stats['translated']}")
    print(f"Fallback/sem alteração: {stats['fallback']}")
    print(f"Vazios: {stats['empty']}")
    print(f"\nCobertura: {stats['translated'] / max(stats['total'] - stats['empty'], 1) * 100:.1f}%")


if __name__ == '__main__':
    main()
