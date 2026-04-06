#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Análise de disposições totalmente novas no Regulamento 2023/0447
Compara com: @codigo, @rgbeac, DL 276/2001, DL 82/2019, Lei 27/2016
"""

import os
import csv
import re
from docx import Document
from collections import defaultdict
from pathlib import Path

# Caminhos dos ficheiros
BASE_PATH = "/home/user/Legislacao"
FILES = {
    "regulamento": "11.12.2025 Regulamento cães e gatos-ocr - sem rasuras.docx",
    "codigo": "Código do Animal DL214.2013_OCR.docx.docx",
    "rgbeac": "RGBEAC_junh_2025 Original com Índice.docx",
    "dl_276_2001": "@legislacao Decreto-Lei n.º 276-2001, de 17 de outubro v2.docx",
    "dl_82_2019": "@legislacao DL n. 82_2019, de 27 de Junho_ocred.docx",
    "lei_27_2016": "@legislacao Lei n.º 27_2016, de 23 de agosto - Aprova medidas para a criação de uma rede de centros de recolha oficial de animais e estabelece a proibição do abate de animais errantes como forma de controlo da população_ocred.docx",
}

def read_docx(filepath):
    """Extrai todo o texto de um ficheiro DOCX."""
    try:
        doc = Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs])
        # Adiciona texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        return text
    except Exception as e:
        print(f"Erro ao ler {filepath}: {e}")
        return ""

def extract_articles(text, source_name):
    """
    Extrai artigos usando padrões regex.
    Retorna dict com: {artigo_num: {titulo, texto, full_text}}
    """
    articles = {}

    # Padrão: "Article X." ou "Artigo X.º" ou "Art. X"
    patterns = [
        r"Article\s+(\d+)[.\s]*\n\s*([^\n]*)\n(.*?)(?=Article\s+\d+|$)",  # EN
        r"Artigo\s+(\d+)[.º]*\s*[–\-]\s*([^\n]*)\n(.*?)(?=Artigo\s+\d+|$)",  # PT
        r"Art\.\s+(\d+)[.º]*\s*[–\-]\s*([^\n]*)\n(.*?)(?=Art\.\s+\d+|$)",  # PT
    ]

    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
        for match in matches:
            art_num = int(match.group(1))
            titulo = match.group(2).strip()[:100]  # Primeiras 100 chars do título
            conteudo = match.group(3).strip()[:500]  # Primeiros 500 chars

            if art_num not in articles:
                articles[art_num] = {
                    "titulo": titulo,
                    "conteudo": conteudo,
                    "fonte": source_name
                }

    return articles

def calculate_similarity(text1, text2, threshold=0.3):
    """
    Calcula similaridade simples baseada em palavras-chave comuns.
    Retorna True se há correspondência potencial.
    """
    if not text1 or not text2:
        return False

    # Normaliza e extrai palavras significativas
    def extract_keywords(text):
        # Remove stopwords português/inglês comuns
        stopwords = {'o', 'a', 'de', 'do', 'da', 'e', 'ou', 'um', 'uma', 'the', 'of', 'and', 'or', 'a', 'is', 'be', 'to', 'that', 'this'}
        words = set(re.findall(r'\b\w{3,}\b', text.lower()))
        return words - stopwords

    kw1 = extract_keywords(text1)
    kw2 = extract_keywords(text2)

    if not kw1 or not kw2:
        return False

    overlap = len(kw1 & kw2)
    similarity = overlap / max(len(kw1), len(kw2))

    return similarity >= threshold

def find_correspondences(reg_article, other_texts):
    """
    Procura correspondência de um artigo do regulamento noutros diplomas.
    Retorna lista de diplomas com potencial correspondência.
    """
    reg_text = f"{reg_article['titulo']} {reg_article['conteudo']}"
    matches = []

    for source, text in other_texts.items():
        if text and calculate_similarity(reg_text, text, threshold=0.25):
            matches.append(source)

    return matches

def main():
    print("=" * 80)
    print("ANÁLISE DE DISPOSIÇÕES NOVAS - Regulamento 2023/0447")
    print("=" * 80)

    # 1. Lê todos os ficheiros
    print("\n[1/4] Lendo ficheiros...")
    all_texts = {}
    for source, filename in FILES.items():
        filepath = os.path.join(BASE_PATH, filename)
        if os.path.exists(filepath):
            all_texts[source] = read_docx(filepath)
            print(f"  ✓ {source}: {len(all_texts[source])} caracteres")
        else:
            print(f"  ✗ {source}: FICHEIRO NÃO ENCONTRADO")

    # 2. Extrai artigos do Regulamento
    print("\n[2/4] Extraindo artigos do Regulamento 2023/0447...")
    reg_articles = extract_articles(all_texts.get("regulamento", ""), "regulamento")
    print(f"  Encontrados {len(reg_articles)} artigos")

    # 3. Extrai artigos dos outros diplomas
    print("\n[3/4] Extraindo artigos dos diplomas portugueses...")
    other_articles = {}
    for source in ["codigo", "rgbeac", "dl_276_2001", "dl_82_2019", "lei_27_2016"]:
        articles = extract_articles(all_texts.get(source, ""), source)
        other_articles[source] = articles
        print(f"  {source}: {len(articles)} artigos")

    # 4. Identifica disposições novas
    print("\n[4/4] Identificando disposições novas...")

    results = []
    articles_sem_correspondencia = []

    for art_num in sorted(reg_articles.keys()):
        reg_art = reg_articles[art_num]

        # Procura correspondência em todos os diplomas
        correspondencias = []
        for source in ["codigo", "rgbeac", "dl_276_2001", "dl_82_2019", "lei_27_2016"]:
            if source in other_articles:
                # Procura em todo o texto (não só por número de artigo)
                source_text = all_texts.get(source, "")
                if find_correspondences(reg_art, {source: source_text}):
                    correspondencias.append(source)

        # Classifica como "novo" se não tem correspondência
        if not correspondencias:
            articles_sem_correspondencia.append({
                "artigo": art_num,
                "titulo": reg_art["titulo"],
                "conteudo": reg_art["conteudo"],
                "correspondencias": correspondencias
            })

        results.append({
            "artigo": art_num,
            "titulo": reg_art["titulo"],
            "conteudo_resumo": reg_art["conteudo"][:300],
            "tem_correspondencia": len(correspondencias) > 0,
            "fontes_correspondentes": ", ".join(correspondencias) if correspondencias else "NENHUMA"
        })

    # 5. Salva CSV
    print("\n[5/5] Gerando relatório CSV...")

    csv_path = os.path.join(BASE_PATH, "disposicoes_novas_regulamento.csv")
    with open(csv_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=[
            "artigo", "titulo", "conteudo_resumo",
            "tem_correspondencia", "fontes_correspondentes"
        ])
        writer.writeheader()
        writer.writerows(results)

    print(f"  ✓ Gravado: {csv_path}")

    # Sumário
    print("\n" + "=" * 80)
    print("SUMÁRIO")
    print("=" * 80)
    print(f"Total de artigos no Regulamento: {len(reg_articles)}")
    print(f"Artigos COM correspondência: {len([r for r in results if r['tem_correspondencia']])}")
    print(f"Artigos SEM correspondência (NOVOS): {len(articles_sem_correspondencia)}")

    if articles_sem_correspondencia:
        print("\n" + "-" * 80)
        print("ARTIGOS TOTALMENTE NOVOS (sem correspondência):")
        print("-" * 80)
        for art in articles_sem_correspondencia:
            print(f"\n  Artigo {art['artigo']}: {art['titulo']}")
            print(f"  Resumo: {art['conteudo'][:200]}...")

if __name__ == "__main__":
    main()
