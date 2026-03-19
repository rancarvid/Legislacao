from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'start', 'bottom', 'end'):
        tag = f'w:{edge}'
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_header_row(table, col_headers):
    row = table.rows[0]
    for i, hdr in enumerate(col_headers):
        cell = row.cells[i]
        cell.text = hdr
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'


def add_data_row(table, values):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        run.font.name = 'Arial'


def add_section_title(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def create_table(doc, col_widths_cm, header_row):
    table = doc.add_table(rows=1, cols=len(col_widths_cm))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, w in enumerate(col_widths_cm):
        table.columns[i].width = Cm(w)
    add_header_row(table, header_row)
    return table


doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

title = doc.add_paragraph()
run = title.add_run('Quadro Comparativo — Artigo 21.º do Regulamento (UE) 2023/0447')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Arial'
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
run2 = sub.add_run(
    'Requisitos de publicidade em linha e colocação no mercado — '
    'correspondência com @rgbeac (jun. 2025), @codigo (DL n.º 214/2013) e legislação vigente'
)
run2.font.size = Pt(9)
run2.italic = True
run2.font.name = 'Arial'
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(8)

HDRS = [
    'Critério / Determinação do Regulamento',
    'Regulamento (UE) 2023/0447\n(texto PT-PT)',
    '@rgbeac\n(proposta jun. 2025)',
    '@codigo\n(proposta DL n.º 214/2013)',
]
COL_W = [5.5, 4.5, 4.5, 4.5]

add_section_title(doc, 'Artigo 21.º — Requisitos de publicidade em linha e colocação no mercado')

intro_p = doc.add_paragraph()
run_intro = intro_p.add_run(
    'O art.º 21.º regula os anúncios online de cães e gatos com vista à colocação no mercado da União, '
    'impondo avisos obrigatórios, informação ao adquirente, um sistema de token único de verificação '
    'e obrigações para plataformas digitais. Analisa-se cada n.º de forma desagregada.'
)
run_intro.font.size = Pt(8.5)
run_intro.font.name = 'Arial'
intro_p.paragraph_format.space_after = Pt(4)

t = create_table(doc, COL_W, HDRS)

rows = [
    # n.º 1 — aviso obrigatório nos anúncios online (operadores)
    (
        'n.º 1 — Aviso nos anúncios online\n(operadores)\nCaracteres visíveis e negrito:\n"Um animal não é um brinquedo..."',
        'Quando operadores anunciem online um cão ou gato com vista ao seu colocamento no mercado '
        'da União, devem assegurar a apresentação do seguinte aviso em caracteres claramente visíveis '
        'e em negrito:\n"Um animal não é um brinquedo. Ter um é uma decisão que muda a vida. '
        'É seu dever assegurar a sua saúde e bem-estar e não o abandonar."\n(n.º 1 do art.º 21.º)',
        'Art.º 92.º — requisitos de validade do anúncio de transmissão: sem exigência de aviso '
        'de detenção responsável com texto predefinido.\n'
        'Art.º 91.º, n.º 2 — obrigação do vendedor de aferir adequação do estilo de vida do '
        'comprador: dever pré-contratual de informação, não equivalente ao aviso em negrito.\n'
        'Não coberto.',
        'Art.º 80.º, n.º 9 — obrigação de disponibilizar informação escrita sobre cuidados gerais '
        'e obrigações legais do detentor (no momento da venda, não no anúncio).\n'
        'Sem aviso pré-definido em anúncio online.\nNão coberto.',
    ),
    # n.º 2 — aviso nos anúncios online (não operadores)
    (
        'n.º 2 — Aviso nos anúncios online\n(não operadores)\nAviso de detenção responsável\n(formulação própria ou equivalente)',
        'Quando pessoas singulares ou coletivas que não sejam operadores anunciem online um cão ou '
        'gato com vista ao seu colocamento no mercado da União, devem assegurar a apresentação de '
        'um aviso sobre detenção responsável utilizando a redação referida no n.º 1 ou uma redação '
        'diferente com significado equivalente.\n(n.º 2 do art.º 21.º)',
        'Sem disposição equivalente para particulares/não operadores que anunciem online.\n'
        'O art.º 92.º dirige-se apenas a quem anuncia transmissões (sem distinguir operadores '
        'de particulares), mas não exige aviso de detenção responsável.\nNão coberto.',
        'Art.º 80.º, n.ºs 1 a 3 — comercialização restrita a alojamentos autorizados; '
        'particulares sem fins lucrativos excetuados.\n'
        'Sem aviso de detenção responsável para anúncios de particulares.\nNão coberto.',
    ),
    # n.º 3 (1/2) — informação obrigatória ao adquirente: prova de identificação e dados do animal
    (
        'n.º 3 (1/2) — Informação ao adquirente\nProva de identificação/registo\n(art.º 20.º) + espécie, sexo,\ndata/país de nascimento, raça',
        'Ao colocar um cão ou gato no mercado na União, a pessoa singular ou coletiva deve '
        'fornecer ao adquirente:\n(a) prova de identificação e registo em conformidade com '
        'o art.º 20.º;\n(b) espécie; sexo; data e país de nascimento; raça (quando relevante).\n'
        '(n.º 3 do art.º 21.º)',
        'Art.º 94.º, n.º 1 — documentos obrigatórios na transmissão: contrato, documento de '
        'identificação do animal (art.º 26.º, n.º 6), declaração veterinária de boa saúde '
        '(15 dias), historial clínico, licença de detenção (animais perigosos), prova de '
        'legalidade.\n'
        'Art.º 92.º, n.º 1 — conteúdo do anúncio: data de nascimento, país de origem, '
        'fotografia, n.º de identificação eletrónica, raça (pura ou indeterminada).\n'
        'Parcialmente coberto (país de origem e data de nascimento no anúncio; '
        'prova de identificação na transmissão; sexo não referenciado de forma autónoma).',
        'Art.º 80.º, n.ºs 5 e 6 — documentos obrigatórios: boletim sanitário/passaporte com '
        'etiqueta de microchip, comprovativo de profilaxia, documento de registo com n.º de '
        'microchip como documento de transferência de propriedade.\n'
        'Art.º 80.º, n.º 7 — comprovativo de profilaxia obrigatória para cães e gatos > 3 meses.\n'
        'Parcialmente coberto (identificação e registo cobertos; sexo, país de nascimento '
        'não referidos de forma autónoma).',
    ),
    # n.º 3 (2/2) — token único de verificação nos anúncios online
    (
        'n.º 3 (2/2) — Token único de verificação\nnos anúncios online\n(gerado pelo sistema do n.º 5;\nweblink incluído no anúncio)',
        'Caso uma pessoa singular ou coletiva anuncie um cão ou gato online com vista à colocação '
        'no mercado da União, deve utilizar o sistema referido no n.º 5 para gerar um token único '
        'de verificação e incluir esse token no anúncio, juntamente com uma ligação web para o '
        'sistema referido no n.º 5.\nO sistema deve permitir aos adquirentes verificar a '
        'autenticidade da identificação, do registo e da propriedade de cães ou gatos anunciados '
        'online.\n(n.º 3 do art.º 21.º)',
        'Art.º 93.º — plataformas obrigadas a implementar sistema de controlo e validação da '
        'veracidade dos dados do anunciante.\n'
        'Sem sistema de token único gerado pela autoridade/Comissão para cada anúncio.\n'
        'Não coberto.',
        'Sem disposição sobre token único de verificação ou sistema de autenticação '
        'de anúncios online.\nNão coberto.',
    ),
    # n.º 4 — obrigações das plataformas online
    (
        'n.º 4 — Plataformas online\nFacilitar conformidade dos anunciantes;\ninformar adquirentes do sistema\nde verificação; sem monitorização geral',
        'Os fornecedores de plataformas online devem assegurar que a sua interface é concebida '
        'de forma a facilitar o cumprimento das obrigações dos n.ºs 1 a 3, e devem informar os '
        'adquirentes, de forma visível, da possibilidade de verificar a autenticidade da '
        'identificação, do registo e da propriedade no sistema referido no n.º 5.\n'
        'Apenas a pessoa que coloca o animal no mercado é responsável pela exatidão das '
        'informações. Sem obrigação geral de monitorização para a plataforma '
        '(art.º 8.º do Regulamento (UE) 2022/2065).\n(n.º 4 do art.º 21.º)',
        'Art.º 93.º, n.º 1 — plataformas apenas podem publicitar anúncios que cumpram os '
        'requisitos do art.º 92.º.\n'
        'Art.º 93.º, n.º 2 — dever de implementar e manter sistema de controlo e validação '
        'para verificar a veracidade dos dados introduzidos pelo anunciante.\n'
        'Parcialmente coberto (obrigação de validação mais exigente que o Regulamento; '
        'sem distinção entre responsabilidade da plataforma e do anunciante; '
        'sem obrigação de informar adquirentes do sistema de verificação).',
        'Sem disposições sobre obrigações de plataformas digitais de anúncios.\nNão coberto.',
    ),
    # n.º 5 — sistema de verificação da Comissão
    (
        'n.º 5 — Sistema de verificação\nda Comissão\nGratuito, automatizado, online;\ntoken único; proteção de dados;\npossibilidade de delegação',
        'A Comissão garante que um sistema de verificação automatizado da autenticidade da '
        'identificação, do registo e da propriedade de cães ou gatos anunciados online, '
        'utilizando a base de dados referida no art.º 23.º, está publicamente disponível online '
        'e a título gratuito, e gera o token único de verificação referido no n.º 3.\n'
        'A Comissão pode confiar o desenvolvimento e manutenção a uma entidade independente.\n'
        'O sistema deve assegurar: verificação fiável da autenticidade utilizando bases de dados '
        'nacionais; conformidade com proteção de dados (Regulamentos (UE) 2018/1725 e 2016/679).\n'
        '(n.º 5 do art.º 21.º)',
        'Art.º 20.º (SIAC) — sistema nacional de informação de animais de companhia, '
        'gerido pela DGAV, com dados de rastreabilidade, identificação e titularidade.\n'
        'O SIAC é a base de dados que alimentaria o sistema de verificação, mas o @rgbeac '
        'não prevê um sistema de verificação online público com token único.\nNão coberto.',
        'Art.º 53.º-55.º — base de dados nacional (SIAC) para identificação e registo, '
        'gerida pela DGAV.\nSem sistema de verificação online público com token único.\n'
        'Não coberto.',
    ),
    # n.º 6 — atos de execução
    (
        'n.º 6 — Atos de execução\nInformação de prova de ID/registo;\ninformação ao sistema de verificação;\ncaracterísticas técnicas do sistema',
        'A Comissão adota atos de execução que estabelecem:\n(a) informações a fornecer como '
        'prova de identificação e registo (n.º 3, al. (a));\n(b) informações a fornecer ao '
        'sistema de verificação do n.º 5 para demonstrar autenticidade;\n(c) características '
        'do sistema do n.º 5: funções essenciais, requisitos técnicos/eletrónicos/criptográficos, '
        'passos procedimentais para o colocador e para o adquirente.\n'
        'Prazos: al. (a) até 2 anos; als. (b) e (c) até 3 anos após entrada em vigor.\n'
        '(n.º 6 do art.º 21.º)',
        'Sem equivalência a atos de execução da Comissão Europeia.\n'
        'O @rgbeac prevê portarias nacionais para regulamentação de matérias específicas '
        '(ex.: art.º 58.º — requisitos de alojamento), mas não relativas a sistemas de '
        'verificação online ou interoperabilidade europeia.\nNão aplicável.',
        'Sem equivalência a atos de execução da Comissão Europeia.\nNão aplicável.',
    ),
]

for r in rows:
    add_data_row(t, r)

output_path = '/home/user/Legislacao/tabela_art21_comparativo.docx'
doc.save(output_path)
print(f'Documento gerado: {output_path}')
